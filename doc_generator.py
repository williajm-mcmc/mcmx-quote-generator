# doc_generator.py
import os
import sys
import re
import math as _math
from html.parser import HTMLParser
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm, Pt, RGBColor
from docx.oxml.ns import qn as _qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# markupsafe is a required dependency of Jinja2 (and therefore docxtpl).
# Wrapping pre-escaped strings in Markup tells Jinja2 the value is already
# XML-safe, preventing double-escaping when docxtpl's autoescape is on.
try:
    from markupsafe import Markup as _Markup
except ImportError:
    _Markup = None   # fall back to plain str; escaping still applied


def _xml_safe(v):
    """
    Escape XML special characters in a user-supplied string, then mark it
    as already-safe so docxtpl / Jinja2 will not escape it a second time.

    Without this, a project name like "AT&T" produces a bare & in the Word
    XML string that docxtpl passes to lxml, causing:
        xmlParseEntityRef: no name, line 1, column N (<string>, line 1)
    """
    if not isinstance(v, str):
        return v
    escaped = (v.replace('&',  '&amp;')
                .replace('<',  '&lt;')
                .replace('>',  '&gt;')
                .replace('"',  '&quot;'))
    return _Markup(escaped) if _Markup is not None else escaped

_HERE = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))

BOM_HEADERS = ["Part Number", "Qty", "Price Each", "Total Price"]

def _style_cell(cell, text, bold=False, font_name="Aptos", font_size=11,
                align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold; run.font.name = font_name; run.font.size = Pt(font_size)
    if color: run.font.color.rgb = color
    cell.paragraphs[0].alignment = align

def _shade_row(row, fill_hex: str):
    """Apply a solid background fill to every cell in a table row."""
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(_qn("w:val"), "clear")
        shd.set(_qn("w:color"), "auto")
        shd.set(_qn("w:fill"), fill_hex)
        tcPr.append(shd)



def _set_table_full_width(table):
    """Full page width + column alignment: left | center | right."""
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _E
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # python-docx CT_Tbl may not have get_or_add_tblPr — use OxmlElement fallback
    tbl_el = table._tbl
    tbl_pr = tbl_el.find(_qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = _E('w:tblPr')
        tbl_el.insert(0, tbl_pr)
    for ex in tbl_pr.findall(_qn('w:tblW')): tbl_pr.remove(ex)
    tblW = _E('w:tblW')
    tblW.set(_qn('w:type'), 'pct')
    tblW.set(_qn('w:w'), '5000')   # 100%
    tbl_pr.append(tblW)

    n_cols = len(table.columns)
    for row in table.rows:
        for c_i, cell in enumerate(row.cells):
            if c_i == 0:
                align = WD_ALIGN_PARAGRAPH.LEFT
            elif c_i == n_cols - 1:
                align = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                align = WD_ALIGN_PARAGRAPH.CENTER
            for para in cell.paragraphs:
                para.alignment = align


def _build_bom_table(doc, rows):
    """
    rows: each row is [part_num, qty, cost_each, marked_up, total]
    The doc always shows the marked-up price (index 3) as the export price.
    If marked_up is empty (margin disabled), falls back to cost_each * qty.
    """
    HEADER_BG    = "D9D9D9"          # light gray fill
    HEADER_COLOR = RGBColor(0x1a, 0x1a, 0x2e)   # dark text on gray
    TOTAL_COLOR  = RGBColor(0x1a, 0x1a, 0x2e)

    table = doc.add_table(rows=1 + len(rows) + 1, cols=4)
    table.style = "Table Grid"

    # Header row — gray background, bold dark text
    _shade_row(table.rows[0], HEADER_BG)
    for i, heading in enumerate(BOM_HEADERS):
        _style_cell(table.rows[0].cells[i], heading, bold=True,
                    font_size=11, color=HEADER_COLOR,
                    align=WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT)

    grand_total = 0.0
    for r_i, row in enumerate(rows):
        cells = table.rows[r_i + 1].cells
        part      = row[0] if len(row) > 0 else ""
        qty       = row[1] if len(row) > 1 else ""
        sale_unit = row[2] if len(row) > 2 else ""  # sale price per unit
        part_html = row[5] if len(row) > 5 else ""

        # Strip $ and commas before float conversion
        def _clean(s): return s.replace('$','').replace(',','').strip() if s else ''
        try:
            unit_price = float(_clean(sale_unit)) if _clean(sale_unit) else 0.0
        except (ValueError, AttributeError):
            unit_price = 0.0
        try:
            # Handle '5 (Hours)' — strip non-numeric suffix
            import re as _re
            qty_f = float(_re.sub(r'[^\d.]', '', qty.split()[0])) if qty else 0.0
        except (ValueError, TypeError, IndexError):
            qty_f = 0.0

        line_total  = qty_f * unit_price
        grand_total += line_total

        # Use rich HTML if available (bullets, bold etc from cell editor).
        # Render into a fresh temporary Document first (same approach as
        # paragraph sections) so all python-docx APIs work on a real Document
        # object, then deep-copy the resulting paragraph XML into the cell.
        if part_html and '<' in part_html:
            try:
                import copy as _copy
                from docx import Document as _TmpDoc
                from docx.oxml.ns import qn as _qnc

                tmp = _TmpDoc()
                _html_to_docx_paragraphs(tmp, part_html)

                # Post-process: bold every run in the first non-empty paragraph
                # so the part-name line is always bold regardless of cell editor
                # formatting (custom items, service rows, etc.)
                from docx.shared import Pt as _Pt2
                _all_paras = [p for p in tmp.paragraphs if p.runs]
                if _all_paras:
                    for _r in _all_paras[0].runs:
                        _r.bold = True

                # Copy bullet numbering definitions into the rendered doc so
                # any numId references inside the cell resolve correctly.
                try:
                    tmp_np = tmp.part.numbering_part._element
                    try:
                        ren_np = doc.part.numbering_part._element
                    except Exception:
                        from docx.parts.numbering import NumberingPart as _NP
                        _ren_part = _NP.new()
                        doc.part._add_relationship(
                            'http://schemas.openxmlformats.org/officeDocument/'
                            '2006/relationships/numbering', _ren_part)
                        ren_np = _ren_part._element
                    _ex_abs = {x.get(_qnc('w:abstractNumId'))
                               for x in ren_np.findall(_qnc('w:abstractNum'))}
                    _ex_num = {x.get(_qnc('w:numId'))
                               for x in ren_np.findall(_qnc('w:num'))}
                    for _ae in tmp_np.findall(_qnc('w:abstractNum')):
                        if _ae.get(_qnc('w:abstractNumId')) not in _ex_abs:
                            ren_np.append(_copy.deepcopy(_ae))
                    for _ne in tmp_np.findall(_qnc('w:num')):
                        if _ne.get(_qnc('w:numId')) not in _ex_num:
                            ren_np.append(_copy.deepcopy(_ne))
                except Exception:
                    pass

                # Swap cell content: remove old paragraphs, inject new ones.
                for p_el in list(cells[0]._tc.findall(_qnc('w:p'))):
                    cells[0]._tc.remove(p_el)
                for elem in list(tmp.element.body):
                    if elem.tag.endswith('}sectPr'):
                        continue
                    cells[0]._tc.append(_copy.deepcopy(elem))
                # Ensure cell always has at least one paragraph (Word requires it)
                if not cells[0]._tc.findall(_qnc('w:p')):
                    from docx.oxml import OxmlElement as _Ep
                    cells[0]._tc.append(_Ep('w:p'))
            except Exception:
                import traceback; traceback.print_exc()
                _style_cell(cells[0], part, bold=True, font_size=11)
        else:
            _style_cell(cells[0], part, bold=True, font_size=11)
        _style_cell(cells[1], qty,  bold=True, font_size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
        _style_cell(cells[2], f"${_math.ceil(unit_price):,}", bold=True, font_size=11, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _style_cell(cells[3], f"${_math.ceil(line_total):,}", bold=True, font_size=11, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Apply column-width and default alignment first, then stamp total row
    # alignment afterwards so _set_table_full_width can't clobber it.
    _set_table_full_width(table)
    total_row = table.rows[-1].cells
    # Merge the first three cells so only "Total Price:" + value remain
    merged = total_row[0].merge(total_row[2])
    _style_cell(merged, "Total Price:", bold=True, font_size=11,
                color=TOTAL_COLOR, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _style_cell(total_row[3], f"${_math.ceil(grand_total):,}", bold=True, font_size=11,
                color=TOTAL_COLOR, align=WD_ALIGN_PARAGRAPH.RIGHT)
    return table

def _replace_placeholder_with_table(rendered_doc, key, rows):
    for para in rendered_doc.paragraphs:
        if key in para.text:
            table = _build_bom_table(rendered_doc, rows)
            para._element.addnext(table._element)
            para._element.getparent().remove(para._element)
            return
    print(f"Warning: placeholder '{key}' not found.")


def _html_to_docx_paragraphs(doc_or_cell, html: str, _doc_obj=None):
    """
    Convert Qt HTML into docx paragraphs.
    doc_or_cell: a Document OR a Cell.
    _doc_obj: the Document object (required when doc_or_cell is a Cell,
              so numbering lookups work correctly).
    """
    is_cell = hasattr(doc_or_cell, '_tc')
    if is_cell:
        cell     = doc_or_cell
        real_doc = _doc_obj   # Document passed in explicitly

        class _CellWrap:
            def add_paragraph(self_, style=None):
                from docx.oxml import OxmlElement as _Ec
                from docx.text.paragraph import Paragraph
                p_el = _Ec('w:p')
                cell._tc.append(p_el)
                p = Paragraph(p_el, cell._tc)
                if style and real_doc is not None:
                    try:
                        from docx.oxml.ns import qn as _qns
                        from docx.oxml import OxmlElement as _Es
                        pPr = p._p.get_or_add_pPr()
                        pStyle = _Es('w:pStyle')
                        pStyle.set(_qns('w:val'), style)
                        pPr.insert(0, pStyle)
                    except Exception:
                        pass
                return p
            @property
            def part(self_):
                # Return the real document part so numbering lookups work
                return real_doc.part if real_doc is not None else None

        _html_to_docx_paragraphs_inner(_CellWrap(), html)
    else:
        _html_to_docx_paragraphs_inner(doc_or_cell, html)


# Module-level cache: persists across all calls within one process run
_bullet_numid_cache = {}

def _html_to_docx_paragraphs_inner(doc, html: str):
    """
    Convert Qt-generated rich-text HTML into python-docx content.

    Qt's QTextEdit HTML quirks this handles:
      - Bullet lists: <ul style="..."><li style="-qt-list-indent:N"><p>text</p></li></ul>
        AND the flattened form where Qt emits <p style="-qt-list-indent:N ...">
      - Tables: <table><thead/tbody><tr><td><p>text</p></td></tr></table>
        but also nested inside a <p style="-qt-paragraph-type:empty"> wrapper
      - Bold/italic via inline style font-weight/font-style as well as <b>/<i> tags
      - Indented bullets: -qt-list-indent value maps to Word indent level
    """
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # ── pull body HTML ────────────────────────────────────────────────────────
    body_m = re.search(r'<body[^>]*>(.*?)</body>', html, re.S | re.I)
    body_html = body_m.group(1) if body_m else html

    # ── pre-extract data URIs before the HTML parser sees them ────────────────
    # Python's html.parser can corrupt very long attribute values (e.g. base64
    # images).  Pull every data: URI out of src attributes up front, decode the
    # bytes into BytesIO objects keyed by a short placeholder token, and put the
    # token in the HTML instead.  _walk's img handler then looks up the token.
    import base64 as _b64_pre, io as _io_pre
    _data_uri_store = {}   # token -> BytesIO

    def _strip_data_uris(h):
        def _repl(m):
            q, uri = m.group(1), m.group(2)
            if not uri.startswith('data:'):
                return m.group(0)
            try:
                hdr, raw_b64 = uri.split(',', 1)
                mime = (hdr.split(';')[0].split(':')[1]
                        if ':' in hdr else 'image/png')
                ext = mime.split('/')[-1].lower()
                if ext not in ('png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'tif'):
                    ext = 'png'
                clean_b64 = ''.join(raw_b64.split())
                img_bytes  = _b64_pre.b64decode(clean_b64 + '==')
                token = f'__IMG_{len(_data_uri_store)}__.{ext}'
                _data_uri_store[token] = _io_pre.BytesIO(img_bytes)
                return f'src={q}{token}{q}'
            except Exception:
                return m.group(0)
        return re.sub(r'src=(["\'])([^"\']*)\1', _repl, h)

    body_html = _strip_data_uris(body_html)

    # ── minimal DOM builder ───────────────────────────────────────────────────
    class _N:
        __slots__ = ('tag','attrs','children','text')
        def __init__(self, tag, attrs=None):
            self.tag      = (tag or '').lower()
            self.attrs    = attrs or {}
            self.children = []
            self.text     = ''
        def attr(self, k, default=''):
            return self.attrs.get(k, default)
        def style(self):
            return self.attr('style')
        def has_style(self, fragment):
            return fragment.lower() in self.style().lower()

    class _Builder(HTMLParser):
        def __init__(self):
            super().__init__()
            self.root = _N('root')
            self._stack = [self.root]
        def handle_starttag(self, tag, attrs):
            n = _N(tag, dict(attrs))
            self._stack[-1].children.append(n)
            self._stack.append(n)
        def handle_endtag(self, _):
            if len(self._stack) > 1: self._stack.pop()
        def handle_data(self, data):
            self._stack[-1].text += data
        def handle_entityref(self, name):
            mapping = {'amp':'&','lt':'<','gt':'>','quot':'"','nbsp':' ','apos':"'"}
            self._stack[-1].text += mapping.get(name, '')
        def handle_charref(self, name):
            try:
                c = chr(int(name[1:], 16) if name.startswith('x') else int(name))
                self._stack[-1].text += c
            except Exception:
                pass

    b = _Builder()
    b.feed(body_html)

    # ── run extraction ────────────────────────────────────────────────────────
    def _bold(node):
        if node.tag in ('b', 'strong'): return True
        s = node.style()
        return ('font-weight:600' in s or 'font-weight:700' in s or
                'font-weight: 600' in s or 'font-weight: 700' in s or
                'font-weight:bold' in s.lower())

    def _italic(node):
        if node.tag in ('i', 'em'): return True
        s = node.style()
        return 'font-style:italic' in s or 'font-style: italic' in s

    # 'p' and 'li' removed so their inline text IS collected as run content
    SKIP_TEXT_TAGS = frozenset(
        ('td','th','tr','table','thead','tbody','tfoot',
         'ul','ol','body','html','root','head','style','script'))

    def _font_size(node):
        """Return pt size from inline style font-size, or None."""
        m = re.search(r'font-size\s*:\s*([\d.]+)pt', node.style())
        return float(m.group(1)) if m else None

    def _font_family(node):
        """Return font-family value from inline style, or None."""
        m = re.search(r"font-family\s*:\s*'?([^;',]+)", node.style())
        return m.group(1).strip().strip("'\"") if m else None

    def _font_color(node):
        """Return hex color string (no #) from color:#RRGGBB in style, or None."""
        m = re.search(r'color\s*:\s*#([0-9a-fA-F]{6})', node.style())
        return m.group(1) if m else None

    def _runs(node, bold=False, italic=False, size=None, family=None, color=None,
              _stop_lists=False):
        # When called from a bullet context, don't descend into nested lists —
        # they will be walked separately so their items get their own paragraphs.
        if _stop_lists and node.tag in ('ul', 'ol'):
            return
        b2      = bold   or _bold(node)
        i2      = italic or _italic(node)
        size2   = _font_size(node)   or size
        family2 = _font_family(node) or family
        color2  = _font_color(node)  or color
        if node.text and (node.text.strip() or node.tag not in SKIP_TEXT_TAGS):
            yield (node.text, b2, i2, size2, family2, color2)
        for child in node.children:
            yield from _runs(child, b2, i2, size2, family2, color2,
                             _stop_lists=_stop_lists)

    def _fill_para(p, node, stop_lists=False):
        from docx.shared import Pt, RGBColor as _RGB
        for text, b, i, size, family, color in _runs(node, _stop_lists=stop_lists):
            # Skip pure-whitespace runs that carry no formatting
            if not text.strip() and not color and not size and not family:
                continue
            run = p.add_run(text)
            run.bold = b; run.italic = i
            if size:   run.font.size = Pt(size)
            run.font.name = family or "Aptos"
            if color:
                try:
                    run.font.color.rgb = _RGB(
                        int(color[0:2], 16),
                        int(color[2:4], 16),
                        int(color[4:6], 16))
                except Exception:
                    pass

    # ── list helpers ──────────────────────────────────────────────────────────
    def _list_indent(node):
        """Return Qt list indent level (1-based) from -qt-list-indent style."""
        m = re.search(r'-qt-list-indent\s*:\s*(\d+)', node.style())
        return int(m.group(1)) if m else 1

    def _is_bullet(node):
        s = node.style()
        return ('-qt-list-indent' in s or
                'list-style-type' in s.lower() or
                node.tag == 'li')

    def _list_style_type(node):
        """Detect ordered vs unordered from parent ul/ol or style."""
        s = node.style().lower()
        if 'decimal' in s or 'lower-alpha' in s or 'upper-alpha' in s:
            return 'ordered'
        return 'bullet'

    def _ensure_bullet_numid(doc, ordered=False):
        """
        Create (once per doc) an abstractNum defining bullet • on all 9 levels,
        and return the numId string to reference it.
        """
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _E

        # Use id of the document PART (not the wrapper) so _CellWrap
        # and the real Document share the same cache entry
        try:
            _part = doc.part
        except Exception:
            _part = None
        key = (id(_part), ordered)
        if key in _bullet_numid_cache:  # module-level cache
            return _bullet_numid_cache[key]

        try:
            numbering_part = doc.part.numbering_part
        except Exception:
            from docx.parts.numbering import NumberingPart
            numbering_part = NumberingPart.new()
            doc.part._add_relationship(
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering',
                numbering_part)

        np = numbering_part._element

        # Pick IDs that won't clash — use 200+ to stay clear of template defs
        abs_id = '200' if not ordered else '201'
        num_id = '200' if not ordered else '201'

        # Remove any previous attempt with same ID (idempotent)
        for old in np.findall(_qn('w:abstractNum')):
            if old.get(_qn('w:abstractNumId')) == abs_id:
                np.remove(old)
        for old in np.findall(_qn('w:num')):
            if old.get(_qn('w:numId')) == num_id:
                np.remove(old)

        # Build abstractNum with 9 levels
        absNum = _E('w:abstractNum')
        absNum.set(_qn('w:abstractNumId'), abs_id)
        for lvl_idx in range(9):
            lvl = _E('w:lvl')
            lvl.set(_qn('w:ilvl'), str(lvl_idx))
            start = _E('w:start'); start.set(_qn('w:val'), '1')
            lvl.append(start)
            fmt = _E('w:numFmt')
            fmt.set(_qn('w:val'), 'decimal' if ordered else 'bullet')
            lvl.append(fmt)
            txt = _E('w:lvlText')
            txt.set(_qn('w:val'), ('%' + '1.') if ordered else '\u2022')
            lvl.append(txt)
            jc = _E('w:jc'); jc.set(_qn('w:val'), 'left')
            lvl.append(jc)
            # Per-level indentation
            pPr_lvl = _E('w:pPr')
            ind = _E('w:ind')
            ind.set(_qn('w:left'),    str(360 * (lvl_idx + 1)))
            ind.set(_qn('w:hanging'), '360')
            pPr_lvl.append(ind); lvl.append(pPr_lvl)
            absNum.append(lvl)
        np.append(absNum)

        num = _E('w:num')
        num.set(_qn('w:numId'), num_id)
        absRef = _E('w:abstractNumId'); absRef.set(_qn('w:val'), abs_id)
        num.append(absRef)
        np.append(num)

        _bullet_numid_cache[key] = num_id
        return num_id

    def _add_bullet_para(doc, node, indent_level=1, ordered=False):
        """
        Build a bullet paragraph with a self-contained numbering definition
        so every indent level reliably shows a bullet character.

        When Qt nests a <ul> inside a <li> (the structure it uses for
        double-indented bullets), the outer <li> carries no visible text —
        it is a structural container only.  In that case we skip creating a
        paragraph for it (which would produce a spurious empty bullet line)
        and instead only walk the nested list to produce real paragraphs.
        """
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _E

        # ── Nested-list walker ────────────────────────────────────────────────
        # Defined before the content check so it is available in both branches.
        # _walk is assigned later in the same closure scope; Python's late
        # binding guarantees it is resolved at call time, not definition time.
        def _walk_nested_lists(n):
            for child in n.children:
                if child.tag in ('ul', 'ol'):
                    _walk(child)
                else:
                    _walk_nested_lists(child)

        # ── Content check ─────────────────────────────────────────────────────
        # If this <li> has no actual text (only nested <ul>/<ol> children),
        # skip creating an empty bullet paragraph and just walk the children.
        has_content = any(
            text.strip()
            for text, *_ in _runs(node, _stop_lists=True)
        )
        if not has_content:
            _walk_nested_lists(node)
            return None

        # ── Build the bullet paragraph ────────────────────────────────────────
        p   = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()

        num_id = _ensure_bullet_numid(doc, ordered)

        numPr = _E('w:numPr')
        ilvl  = _E('w:ilvl'); ilvl.set(_qn('w:val'), str(max(0, indent_level - 1)))
        numId = _E('w:numId'); numId.set(_qn('w:val'), num_id)
        numPr.append(ilvl); numPr.append(numId)
        pPr.append(numPr)

        spacing = _E('w:spacing')
        spacing.set(_qn('w:before'), '0')
        spacing.set(_qn('w:after'),  '0')
        pPr.append(spacing)

        # Explicit indent — overrides any Normal style inheritance
        left = 360 * indent_level
        ind = _E('w:ind')
        ind.set(_qn('w:left'),    str(left))
        ind.set(_qn('w:hanging'), '360')
        pPr.append(ind)

        # stop_lists=True so text from nested <ul>/<ol> doesn't bleed into
        # this bullet paragraph — nested lists are walked separately below.
        _fill_para(p, node, stop_lists=True)
        _walk_nested_lists(node)

        return p

    # ── table builder ─────────────────────────────────────────────────────────
    def _build_table(doc, node):
        """Build a python-docx table from a <table> node."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        # Collect all <tr> nodes through any tbody/thead/tfoot wrappers
        sections = [c for c in node.children
                    if c.tag in ('tbody','thead','tfoot')]
        if not sections:
            sections = [node]
        all_trs = []
        for sec in sections:
            for ch in sec.children:
                if ch.tag == 'tr':
                    all_trs.append(ch)
        if not all_trs:
            return

        max_cols = max(
            sum(1 for c in tr.children if c.tag in ('td','th'))
            for tr in all_trs
        ) if all_trs else 1

        tbl = doc.add_table(rows=len(all_trs), cols=max(max_cols, 1))
        tbl.style = 'Table Grid'

        from docx.oxml import OxmlElement as _E
        from docx.oxml.ns import qn as _qn

        # Collect numeric values per column for total row calculation
        col_values = [[] for _ in range(max_cols)]

        for r_i, tr in enumerate(all_trs):
            cell_nodes = [c for c in tr.children if c.tag in ('td','th')]
            is_header_row = (r_i == 0)
            # Detect total row: last row whose first cell text is 'Total'
            first_text = ''.join(
                n.text for n in cell_nodes[0].children
                if n.text) if cell_nodes else ''
            is_total_row = (r_i == len(all_trs) - 1 and
                            first_text.strip().lower() == 'total')

            for c_i, cn in enumerate(cell_nodes[:max_cols]):
                cell = tbl.rows[r_i].cells[c_i]
                cell.text = ''
                p = cell.paragraphs[0]

                from docx.shared import Pt as _Pt, RGBColor as _TRGB

                # For total row sum columns, compute and insert the sum
                if is_total_row and c_i > 0 and col_values[c_i]:
                    total_val = sum(col_values[c_i])
                    # Format: int if whole number, else 2dp
                    text_val = (f"{int(total_val):,}" if total_val == int(total_val)
                                else f"{total_val:,.2f}")
                    run = p.add_run(text_val)
                    run.bold = True
                    run.font.name = "Aptos"
                elif is_total_row and c_i == 0:
                    run = p.add_run('Total')
                    run.bold = True
                    run.font.name = "Aptos"
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    for text, bold, italic, size, family, color in _runs(cn):
                        run = p.add_run(text)
                        run.bold   = bold or is_header_row
                        run.italic = italic
                        if size:   run.font.size = _Pt(size)
                        run.font.name = family or "Aptos"
                        if color:
                            try:
                                run.font.color.rgb = _TRGB(
                                    int(color[0:2],16), int(color[2:4],16),
                                    int(color[4:6],16))
                            except Exception: pass
                        # Collect numeric values for sum columns
                        if not is_header_row and not is_total_row and text.strip():
                            try:
                                col_values[c_i].append(
                                    float(text.replace(',','').replace('$','')))
                            except ValueError:
                                pass

                # Shade header and total rows
                fill = 'D9D9D9' if is_header_row else ('C8CEDD' if is_total_row else None)
                if fill:
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = _E('w:shd')
                    shd.set(_qn('w:val'),   'clear')
                    shd.set(_qn('w:color'), 'auto')
                    shd.set(_qn('w:fill'),  fill)
                    tcPr.append(shd)

        _set_table_full_width(tbl)

    # ── main walker ───────────────────────────────────────────────────────────
    def _walk(node):
        tag = node.tag

        # Transparent containers — recurse
        if tag in ('html','head','body','root','style','script','meta','title'):
            for child in node.children:
                _walk(child)
            return

        # Qt wraps everything in <p> — check if it actually contains a table
        # or is a list item before treating it as a plain paragraph
        if tag in ('div', 'p'):
            # Qt emits <p style="-qt-paragraph-type:empty;…"><br/></p> as a
            # structural spacer between <ul> blocks of different indent levels.
            # These carry no real content — skip them to avoid blank lines between
            # indented bullet groups in the Word output.
            if '-qt-paragraph-type:empty' in node.style().lower():
                return

            # Also skip any <p> that carries no text content AND has no
            # table/list children.  Qt can emit these without the explicit
            # empty-type flag (e.g. at the 2→3 indent-level transition), and
            # they would otherwise produce a spurious blank paragraph in Word.
            # We check _runs() so that inline formatting nodes inside the <p>
            # are also considered — only truly contentless nodes are dropped.
            # Strip ASCII whitespace AND non-breaking space (\xa0) — Qt emits
            # &#160; in spacer paragraphs between list levels, which Python's
            # plain .strip() does not remove.
            _p_text   = ''.join(t for t, *_ in _runs(node)).strip(' \t\n\r\xa0')
            _p_struct = any(c.tag in ('table', 'ul', 'ol', 'img') for c in node.children)
            if not _p_text and not _p_struct:
                return

            # Check for nested table
            for child in node.children:
                if child.tag == 'table':
                    # Any text before the table
                    if node.text and node.text.strip():
                        p = doc.add_paragraph()
                        r = p.add_run(node.text)
                        r.font.name = "Aptos"
                    _build_table(doc, child)
                    return

            # Qt bullet list paragraph: <p style="-qt-list-indent:N ...">
            if _is_bullet(node):
                indent = _list_indent(node)
                ordered = _list_style_type(node) == 'ordered'
                _add_bullet_para(doc, node, indent_level=indent, ordered=ordered)
                return

            # Plain paragraph — collect all child content
            # If the <p> contains only <img> children (no text), skip creating a
            # blank paragraph and let _walk handle each image directly.
            _img_children = [c for c in node.children if c.tag == 'img']
            if _img_children and not _p_text:
                for _ic in _img_children:
                    _walk(_ic)
                return

            p = doc.add_paragraph()
            _fill_para(p, node)
            # Qt emits alignment as style='text-align:center' OR align='center'
            _align_val = None
            _am = re.search(r'text-align\s*:\s*(\w+)', node.style())
            if _am:
                _align_val = _am.group(1).lower()
            elif node.attrs.get('align'):
                _align_val = node.attrs['align'].lower()
            if _align_val:
                _amap = {'left': WD_ALIGN_PARAGRAPH.LEFT,
                         'center': WD_ALIGN_PARAGRAPH.CENTER,
                         'right': WD_ALIGN_PARAGRAPH.RIGHT,
                         'justify': WD_ALIGN_PARAGRAPH.JUSTIFY}
                _al = _amap.get(_align_val)
                if _al is not None: p.alignment = _al
            # Walk any inline <img> children mixed with text
            for _ic in _img_children:
                _walk(_ic)
            return

        if tag == 'span':
            # Span outside a paragraph — wrap in one
            p = doc.add_paragraph()
            _fill_para(p, node)
            return

        if tag in ('ul', 'ol'):
            ordered = (tag == 'ol')
            # -qt-list-indent lives on <ul>, not on each <li>
            ul_indent = _list_indent(node)
            for child in node.children:
                if child.tag == 'li':
                    # Use ul-level indent; child may override if it has its own
                    li_indent = _list_indent(child)
                    indent = li_indent if li_indent != 1 or ul_indent == 1 else ul_indent
                    _add_bullet_para(doc, child, indent_level=indent, ordered=ordered)
                else:
                    _walk(child)
            return

        if tag == 'li':
            _add_bullet_para(doc, node, indent_level=_list_indent(node))
            return

        if tag == 'table':
            _build_table(doc, node)
            return

        if tag == 'br':
            doc.add_paragraph()
            return

        if tag == 'img':
            # Don't call add_picture on the tmp document — image relationships
            # can't be safely remapped when copying XML to rendered_doc.
            # Instead write a unique text placeholder; _replace_placeholder_with_rich
            # will find it and add the image directly to the target document.
            src_attr = node.attrs.get('src', '')
            w_attr   = node.attrs.get('width', '')
            align_attr = re.search(r'text-align\s*:\s*(\w+)',
                                   node.attrs.get('style', ''))
            align_val = align_attr.group(1).lower() if align_attr else ''
            try:
                inches = max(1.0, min(6.0, int(w_attr) / 96.0)) if w_attr else 4.0
            except (ValueError, TypeError):
                inches = 4.0

            # Resolve data URI store key -> store key string so the placeholder
            # carries enough info for later insertion.
            if src_attr in _data_uri_store:
                img_key = src_attr           # token like "__IMG_0__.png"
            elif src_attr and not src_attr.startswith('data:'):
                img_key = src_attr           # file path
            else:
                return                       # unknown src, skip

            token = f'[[DOCIMG::{img_key}::{inches:.4f}::{align_val}::DOCIMGEND]]'
            ip = doc.add_paragraph()
            ip.add_run(token)
            return

        # Fallback: recurse into anything else
        for child in node.children:
            _walk(child)

    _walk(b.root)

    # Post-process: remove blank paragraphs sandwiched between two bullet paragraphs.
    # Qt HTML for nested lists emits empty <p> elements between list items that
    # produce spurious blank lines in the Word output.
    _W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    _W_P  = f'{{{_W_NS}}}p'
    _W_T  = f'{{{_W_NS}}}t'
    _numPr_tag = _qn('w:numPr')
    body_paras = [e for e in doc.element.body if e.tag == _W_P]
    for i in range(1, len(body_paras) - 1):
        curr = body_paras[i]
        curr_text = ''.join((t.text or '') for t in curr.iter(_W_T))
        if curr_text.strip():
            continue
        prev_bullet = body_paras[i - 1].find('.//' + _numPr_tag) is not None
        next_bullet = body_paras[i + 1].find('.//' + _numPr_tag) is not None
        if prev_bullet and next_bullet:
            curr.getparent().remove(curr)

# ── Margin configuration widget ───────────────────────────────────────────────


def _replace_placeholder_with_rich(rendered_doc, key, html):
    """Find placeholder paragraph and replace with rich-text content."""
    for i, para in enumerate(rendered_doc.paragraphs):
        if key in para.text:
            anchor = para._element
            import copy
            from docx import Document as _Doc
            from docx.oxml.ns import qn as _qn
            tmp = _Doc()
            _html_to_docx_paragraphs(tmp, html)

            # Copy numbering definitions from tmp into rendered_doc so bullet
            # paragraphs (which reference numId '200'/'201') resolve correctly.
            try:
                tmp_np = tmp.part.numbering_part._element
                try:
                    ren_np = rendered_doc.part.numbering_part._element
                except Exception:
                    from docx.parts.numbering import NumberingPart
                    ren_part = NumberingPart.new()
                    rendered_doc.part._add_relationship(
                        'http://schemas.openxmlformats.org/officeDocument/2006/'
                        'relationships/numbering', ren_part)
                    ren_np = ren_part._element
                existing_abs = {x.get(_qn('w:abstractNumId'))
                                for x in ren_np.findall(_qn('w:abstractNum'))}
                existing_num = {x.get(_qn('w:numId'))
                                for x in ren_np.findall(_qn('w:num'))}
                for abs_el in tmp_np.findall(_qn('w:abstractNum')):
                    if abs_el.get(_qn('w:abstractNumId')) not in existing_abs:
                        ren_np.append(copy.deepcopy(abs_el))
                for num_el in tmp_np.findall(_qn('w:num')):
                    if num_el.get(_qn('w:numId')) not in existing_num:
                        ren_np.append(copy.deepcopy(num_el))
            except Exception as _ne:
                print(f'Numbering copy failed: {_ne}')

            prev = anchor
            for elem in list(tmp.element.body):
                if elem.tag.endswith("}sectPr"):
                    continue
                imported = copy.deepcopy(elem)
                prev.addnext(imported)
                prev = imported
            anchor.getparent().remove(anchor)

            # ── resolve image placeholders ────────────────────────────────
            # Images were not added to tmp; instead _walk wrote a text token
            # ([[DOCIMG::<src>::<inches>::<align>::DOCIMGEND]]).
            # Now add them directly to rendered_doc so there are no cross-
            # document relationship issues.
            import io as _io
            from docx.shared import Inches as _docx_In
            from docx.enum.text import WD_ALIGN_PARAGRAPH as _WA
            _ALIGN_MAP = {'center': _WA.CENTER, 'right': _WA.RIGHT,
                          'left': _WA.LEFT}
            _TOKEN_RE = re.compile(
                r'\[\[DOCIMG::(.+?)::([\d.]+)::(.*?)::DOCIMGEND\]\]')

            for para in list(rendered_doc.paragraphs):
                text = para.text
                m = _TOKEN_RE.search(text)
                if not m:
                    continue
                img_key   = m.group(1)
                inches    = float(m.group(2))
                align_val = m.group(3)

                # Determine the image source
                if img_key in _data_uri_store:
                    _data_uri_store[img_key].seek(0)
                    img_src = _data_uri_store[img_key]
                elif os.path.exists(img_key):
                    img_src = img_key
                else:
                    # Can't resolve — remove the placeholder paragraph
                    para._element.getparent().remove(para._element)
                    continue

                # Clear the placeholder run and insert the picture
                p_el = para._p
                for r_el in list(p_el.findall(_qn('w:r'))):
                    p_el.remove(r_el)
                try:
                    run = para.add_run()
                    run.add_picture(img_src, width=_docx_In(inches))
                    if align_val in _ALIGN_MAP:
                        para.alignment = _ALIGN_MAP[align_val]
                except Exception as _ie:
                    print(f'Image insert failed: {_ie}')
                    para._element.getparent().remove(para._element)

            return
    print(f"Warning: rich placeholder '{key}' not found.")


def _extract_subheaders(html):
    """Return list of sub-header texts (H-button dark-red spans) from Qt HTML.
    H-button color is #9E1B32 as set by RichTextEditor._HDR_COLOR.
    Qt serialises it as lowercase #9e1b32 in span style attributes.
    """
    import re
    from html import unescape
    if not html:
        return []
    results = []
    # Match colour in hex or rgb() form to handle any Qt serialisation variant.
    _COLOR_RE = re.compile(
        r'color\s*:\s*(?:#?9[Ee]1[Bb]32|rgb\s*\(\s*158\s*,\s*27\s*,\s*50\s*\))',
        re.I)
    # Use re.S so .*? spans newlines; capture span attributes separately from
    # content so colour is only checked against the opening tag's style string.
    span_re = re.compile(r'<span(\b[^>]*)>(.*?)</span>', re.I | re.S)
    for m in span_re.finditer(html):
        attrs, content = m.group(1), m.group(2)
        if not _COLOR_RE.search(attrs):
            continue
        # Strip any inner HTML tags to get plain text, then unescape entities.
        text = unescape(re.sub(r'<[^>]+>', '', content)).strip()
        if text:
            results.append(text)
    return results


def _stamp_bookmark(p_el, bm_name, bm_id, qn):
    """Insert a bookmark wrapping all content in a paragraph element."""
    from docx.oxml import OxmlElement as _E
    bm_s = _E('w:bookmarkStart')
    bm_s.set(qn('w:id'), str(bm_id))
    bm_s.set(qn('w:name'), bm_name)
    bm_e = _E('w:bookmarkEnd')
    bm_e.set(qn('w:id'), str(bm_id))
    first_r = p_el.find(qn('w:r'))
    if first_r is not None:
        first_r.addprevious(bm_s)
    else:
        p_el.insert(0, bm_s)
    p_el.append(bm_e)


def _inject_section_numbers(doc, sections):
    """Prepend 'N.  ' to paragraphs matching each section header.

    - Paragraph sections: normalised-whitespace exact match so minor spacing
      differences between the form and the rendered template don't break matching.
    - Table sections: fuzzy contains-match; ALL existing runs are cleared and
      the paragraph is rebuilt as 'N.  Bill of Material'.

    Returns a dict mapping sec_num (str) → bookmark_name so the caller can
    build PAGEREF-based TOC page numbers.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    _XML = 'http://www.w3.org/XML/1998/namespace'

    def _hdr_run_xml(text, size_hp=32):
        """Return a dark-red 16pt bold Aptos <w:r> element."""
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        fonts = OxmlElement('w:rFonts')
        fonts.set(qn('w:ascii'), 'Aptos'); fonts.set(qn('w:hAnsi'), 'Aptos')
        rPr.append(fonts)
        rPr.append(OxmlElement('w:b'))
        sz = OxmlElement('w:sz');   sz.set(qn('w:val'), str(size_hp))
        szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), str(size_hp))
        rPr.append(sz); rPr.append(szCs)
        col = OxmlElement('w:color'); col.set(qn('w:val'), '9E1B32')
        rPr.append(col)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(f'{{{_XML}}}space', 'preserve')
        t.text = text
        r.append(t)
        return r

    _norm = lambda s: ' '.join(s.split()).lower()
    bookmarks    = {}   # sec_num (str) -> bookmark_name (str)
    _stamped_ids = set()  # id(para._p) of already-stamped paragraphs

    for sec in sections:
        hdr      = sec.get("header", "").strip()
        num      = str(sec.get("section_number", ""))
        is_table = sec.get("type") == "table"
        if not hdr or not num:
            continue
        for para in doc.paragraphs:
            # Skip paragraphs already stamped by an earlier section so that
            # multiple BOM sections with the same header ("Bill of Material")
            # each match a distinct paragraph instead of all piling onto the first.
            if id(para._p) in _stamped_ids:
                continue
            text = para.text.strip()
            # Normalised-whitespace match so minor spacing differences between
            # the form and the rendered template don't break matching.
            matched = (_norm(hdr) in _norm(text)) if is_table \
                      else (_norm(text) == _norm(hdr))
            if not matched or not para.runs:
                continue
            if is_table:
                p_el = para._p
                for r_el in list(p_el.findall(qn('w:r'))):
                    p_el.remove(r_el)
                p_el.append(_hdr_run_xml(f"{num}.  "))
                p_el.append(_hdr_run_xml(hdr))
            else:
                para.runs[0]._r.addprevious(_hdr_run_xml(f"{num}.  "))
            # Explicitly set outlineLvl=0 so the TOC \u field picks this up
            # (guarantees correct page numbers even if _set_header_outline_levels misses it)
            _pPr = para._p.get_or_add_pPr()
            for _old in _pPr.findall(qn('w:outlineLvl')):
                _pPr.remove(_old)
            _ol = OxmlElement('w:outlineLvl')
            _ol.set(qn('w:val'), '0')
            _pPr.append(_ol)
            # Add bookmark so the TOC can use PAGEREF for accurate page numbers
            bm_name = f'_tocs{num}'
            _stamp_bookmark(para._p, bm_name, 1000 + len(bookmarks), qn)
            bookmarks[num] = bm_name
            _stamped_ids.add(id(para._p))
            break

    return bookmarks


def _inject_subheader_numbers(doc, sections=None, subheader_counts=None):
    """Inject decimal sub-header numbers (1.1.  1.2.  …) into the document body.

    Must run AFTER _inject_section_numbers so section-header paragraphs are
    already stamped with their 'N.  ' prefix (used to detect section boundaries
    and derive the current section number).

    Detection rules (applied to every paragraph in document order):
      • Has at least one run with 16pt+ dark-red font (same as section headers).
      • First non-empty run text matches '^\\d+\\.\\s' → section header (already
        stamped); reset sub-counter and record current section number, skip.
      • First non-empty run text matches '^\\d+\\.\\d+' → already numbered
        (idempotent re-run guard); skip.
      • Text (case-insensitive) matches a known section header from `sections`
        but has no stamp yet → injection missed it; treat as section boundary,
        reset counter, skip.  Also catches static template headers (T&C, etc.)
        that live outside the app's sections list when their text is passed.
      • Anything else with the colour/size → sub-header; inject 'N.M.  ' label.

    `sections` — the same list passed to _inject_section_numbers; used to build
    a case-insensitive set of known top-level header texts so that static
    template sections (Terms & Conditions, Copyright…) are never mis-numbered
    as sub-headers even if _inject_section_numbers couldn't stamp them.
    """
    import re
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    _XML        = 'http://www.w3.org/XML/1998/namespace'
    _HDR_COLORS = {'9E1B32', '9e1b32', '8B0000', '8b0000'}
    _SEC_RE     = re.compile(r'^\d+\.\s')    # "1.  "  ← section header marker
    _SUB_RE     = re.compile(r'^\d+\.\d+')   # "1.1"   ← already has sub-number

    # Build a case-insensitive set of known top-level header texts.
    # Any dark-red paragraph whose text matches one of these (and hasn't been
    # stamped with "N.  " yet) is treated as a section boundary, not a sub-header.
    _known_hdrs = {sec.get("header", "").strip().lower()
                   for sec in (sections or [])
                   if sec.get("header", "").strip()}

    def _is_hdr_run(run):
        if not run.text.strip():
            return False
        rPr = run._r.find(qn('w:rPr'))
        if rPr is None:
            return False
        col_el = rPr.find(qn('w:color'))
        sz_el  = rPr.find(qn('w:sz'))
        if col_el is None or sz_el is None:
            return False
        return (col_el.get(qn('w:val'), '') in _HDR_COLORS and
                int(sz_el.get(qn('w:val'), '0')) >= 28)

    current_section = None
    sub_count       = 0

    for para in doc.paragraphs:
        runs = para.runs
        if not runs:
            continue
        if not any(_is_hdr_run(r) for r in runs):
            continue

        first_text = next((r.text for r in runs if r.text.strip()), '')

        if _SEC_RE.match(first_text):
            # Already stamped section header — record section, reset counter.
            m = re.match(r'^(\d+)\.', first_text)
            if m:
                current_section = m.group(1)
                sub_count       = 0
            continue

        if _SUB_RE.match(first_text):   # idempotency guard
            continue

        if current_section is None:
            continue

        # Check whether this unstamped dark-red paragraph is a known top-level
        # header (or any top-level boundary like a static T&C / Copyright section).
        full_text = para.text.strip().lower()
        if full_text in _known_hdrs:
            current_section = None
            sub_count       = 0
            continue

        # Hard cap: if we have already stamped as many sub-headers as
        # _extract_subheaders found in the current section's HTML, every
        # further dark-red paragraph must be a top-level boundary (e.g. a
        # static T&C / Copyright section that lives outside form_data).
        if subheader_counts is not None:
            expected = subheader_counts.get(current_section, 0)
            if sub_count >= expected:
                current_section = None
                sub_count       = 0
                continue

        sub_count += 1
        label = f"{current_section}.{sub_count}.  "

        # Build label run with matching dark-red 16pt bold Aptos formatting
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        fonts = OxmlElement('w:rFonts')
        fonts.set(qn('w:ascii'), 'Aptos'); fonts.set(qn('w:hAnsi'), 'Aptos')
        rPr.append(fonts)
        rPr.append(OxmlElement('w:b'))
        sz = OxmlElement('w:sz');   sz.set(qn('w:val'), '32')
        szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), '32')
        rPr.append(sz); rPr.append(szCs)
        col = OxmlElement('w:color'); col.set(qn('w:val'), '9E1B32')
        rPr.append(col)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(f'{{{_XML}}}space', 'preserve')
        t.text = label
        r.append(t)
        runs[0]._r.addprevious(r)


def _insert_toc(doc, entries):
    """
    Insert a Word { TOC \\u \\h \\z } field after the TABLE OF CONTENTS anchor.
    Word rebuilds all entries and page numbers from outline-level paragraphs
    (set by _set_header_outline_levels / _detect_and_stamp_static_headers)
    when the document is opened with updateFields=1 in settings.xml.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    _XML = 'http://www.w3.org/XML/1998/namespace'

    toc_para = None
    for para in doc.paragraphs:
        if para.text.strip().upper() in ("TABLE OF CONTENTS", "TOC",
                                         "TABLE OF CONTENTS:"):
            toc_para = para
            break
    if toc_para is None:
        return

    anchor = toc_para._element

    # Remove any existing TOC placeholder paragraphs (TOC1/TOC2/TOCHeading style)
    _W_NS_STR = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    _TOC_STYLES_RM = {'TOC1', 'TOC2', 'TOCHeading', 'toc 1', 'toc 2', 'toc heading'}
    _sib = anchor.getnext()
    _rm = []
    while _sib is not None:
        _tag = _sib.tag
        if _tag == f'{{{_W_NS_STR}}}p':
            _pPr = _sib.find(f'{{{_W_NS_STR}}}pPr')
            _pSt = _pPr.find(f'{{{_W_NS_STR}}}pStyle') if _pPr is not None else None
            _sv  = (_pSt.get(f'{{{_W_NS_STR}}}val', '') if _pSt is not None else '').lower()
            if _sv in _TOC_STYLES_RM:
                _rm.append(_sib)
                _sib = _sib.getnext()
                continue
        break
    for _e in _rm:
        _e.getparent().remove(_e)

    # Build a single paragraph containing the { TOC \o "1-9" \u \h \z } field.
    # \o "1-9" = include built-in heading styles 1–9 (standard Word ribbon default)
    # \u       = also include paragraphs with an outline level set via w:outlineLvl
    #            (our section headers use outlineLvl=0, not a Heading style)
    # \h       = make entries hyperlinks
    # \z       = hide tab/page-number in web layout
    # dirty="true" on fldChar:begin + updateFields=1 in settings.xml forces Word
    # to regenerate on first open.
    p = OxmlElement('w:p')

    r_begin = OxmlElement('w:r')
    fc_b = OxmlElement('w:fldChar')
    fc_b.set(qn('w:fldCharType'), 'begin')
    fc_b.set(qn('w:dirty'), 'true')
    r_begin.append(fc_b)
    p.append(r_begin)

    r_instr = OxmlElement('w:r')
    it = OxmlElement('w:instrText')
    it.set(f'{{{_XML}}}space', 'preserve')
    it.text = ' TOC \\o "1-9" \\u \\h \\z '
    r_instr.append(it)
    p.append(r_instr)

    r_sep = OxmlElement('w:r')
    fc_s = OxmlElement('w:fldChar')
    fc_s.set(qn('w:fldCharType'), 'separate')
    r_sep.append(fc_s)
    p.append(r_sep)

    r_end = OxmlElement('w:r')
    fc_e = OxmlElement('w:fldChar')
    fc_e.set(qn('w:fldCharType'), 'end')
    r_end.append(fc_e)
    p.append(r_end)

    anchor.addnext(p)

    # Page break after the TOC field
    p_break = OxmlElement('w:p')
    r_break = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r_break.append(br)
    p_break.append(r_break)
    p.addnext(p_break)


def _set_header_outline_levels(doc):
    """
    Scan all paragraphs and set Word outline levels so the TOC field can
    auto-populate via the \\u flag when 'Update Fields' is run in Word.

    Must run AFTER _inject_section_numbers and _inject_subheader_numbers so
    the injected number prefixes can be used to distinguish levels:
      • First run text matches '^\\d+\\.\\s'  (e.g. '1.  ')  → outlineLvl 0 (Heading 1)
      • First run text matches '^\\d+\\.\\d+' (e.g. '1.1.') → outlineLvl 1 (Heading 2)
    """
    import re
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    _HDR_COLORS = {'9E1B32', '9e1b32', '8B0000', '8b0000'}
    _SEC_RE     = re.compile(r'^\d+\.\s')
    _SUB_RE     = re.compile(r'^\d+\.\d+')

    for para in doc.paragraphs:
        hdr_run_found = False
        for run in para.runs:
            if not run.text.strip():
                continue
            rPr = run._r.find(qn('w:rPr'))
            if rPr is None:
                continue
            col_el = rPr.find(qn('w:color'))
            sz_el  = rPr.find(qn('w:sz'))
            if col_el is None or sz_el is None:
                continue
            if (col_el.get(qn('w:val'), '') in _HDR_COLORS and
                    int(sz_el.get(qn('w:val'), '0')) >= 28):
                hdr_run_found = True
                break

        if not hdr_run_found:
            continue

        first_text = next((r.text for r in para.runs if r.text.strip()), '')
        if _SUB_RE.match(first_text):
            level = '1'   # sub-header  → Heading 2
        elif _SEC_RE.match(first_text):
            level = '0'   # section header → Heading 1
        else:
            # Paragraph was not stamped by _inject_section_numbers or
            # _inject_subheader_numbers (e.g. static template sections that
            # were not in the form data).  Do NOT assign any outline level —
            # this prevents them from appearing in the auto-updated TOC.
            continue

        pPr = para._p.get_or_add_pPr()
        for old in pPr.findall(qn('w:outlineLvl')):
            pPr.remove(old)
        ol = OxmlElement('w:outlineLvl')
        ol.set(qn('w:val'), level)
        pPr.append(ol)


def _detect_and_stamp_static_headers(doc, sections):
    """
    Find dark-red 16pt paragraphs that were NOT stamped by _inject_section_numbers
    (i.e. they still have no 'N.  ' prefix) and stamp them with the next available
    section numbers so they appear in the TOC.

    This handles static template sections such as Terms & Conditions and
    Copyright & Intellectual Property that live in the .docx template but are
    not part of form_data["sections"].

    Returns a list of (level, label, text) TOC entries for these sections so
    the caller can append them to toc_entries before calling _insert_toc.
    """
    import re
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    _XML        = 'http://www.w3.org/XML/1998/namespace'
    _HDR_COLORS = {'9E1B32', '9e1b32', '8B0000', '8b0000'}
    _SEC_RE     = re.compile(r'^\d+\.\s')
    _SUB_RE     = re.compile(r'^\d+\.\d+')

    def _is_hdr_run(run):
        if not run.text.strip():
            return False
        rPr = run._r.find(qn('w:rPr'))
        if rPr is None:
            return False
        col_el = rPr.find(qn('w:color'))
        sz_el  = rPr.find(qn('w:sz'))
        if col_el is None or sz_el is None:
            return False
        return (col_el.get(qn('w:val'), '') in _HDR_COLORS and
                int(sz_el.get(qn('w:val'), '0')) >= 28)

    def _hdr_run_xml(text):
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        fonts = OxmlElement('w:rFonts')
        fonts.set(qn('w:ascii'), 'Aptos')
        fonts.set(qn('w:hAnsi'), 'Aptos')
        rPr.append(fonts)
        rPr.append(OxmlElement('w:b'))
        sz = OxmlElement('w:sz');    sz.set(qn('w:val'),   '32')
        szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), '32')
        rPr.append(sz); rPr.append(szCs)
        col = OxmlElement('w:color'); col.set(qn('w:val'), '9E1B32')
        rPr.append(col)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(f'{{{_XML}}}space', 'preserve')
        t.text = text
        r.append(t)
        return r

    # Find the highest section number already assigned by _inject_section_numbers
    max_sec = 0
    for sec in (sections or []):
        try:
            n = int(str(sec.get("section_number", 0)))
            if n > max_sec:
                max_sec = n
        except (ValueError, TypeError):
            pass

    extra_entries = []
    next_sec = max_sec + 1

    # Paragraphs that must never be auto-stamped — _insert_toc owns the
    # "Table of Contents" heading and sets its outlineLvl itself.
    _PROTECTED = {'table of contents', 'toc', 'table of contents:'}

    for para in doc.paragraphs:
        if not para.runs:
            continue
        if not any(_is_hdr_run(r) for r in para.runs):
            continue

        first_text = next((r.text for r in para.runs if r.text.strip()), '')

        # Skip already-stamped section headers and sub-headers
        if _SEC_RE.match(first_text) or _SUB_RE.match(first_text):
            continue

        # Skip protected headings (e.g. the TOC anchor paragraph)
        if para.text.strip().lower() in _PROTECTED:
            continue

        # This is an unstamped dark-red paragraph — stamp it as a new section
        hdr_text = para.text.strip()
        label    = str(next_sec)

        # Insert number run before the first run
        para.runs[0]._r.addprevious(_hdr_run_xml(f"{label}.  "))

        # Set outlineLvl=0 so _set_header_outline_levels + Word's \u flag pick it up
        pPr = para._p.get_or_add_pPr()
        for old in pPr.findall(qn('w:outlineLvl')):
            pPr.remove(old)
        ol = OxmlElement('w:outlineLvl')
        ol.set(qn('w:val'), '0')
        pPr.append(ol)

        bm_name = f'_tocss{label}'
        _stamp_bookmark(para._p, bm_name, 2000 + len(extra_entries), qn)
        extra_entries.append((1, label, hdr_text, bm_name))
        next_sec += 1

    return extra_entries


def generate_doc(form_data, template_path=None, output_path=None):
    # Clear numbering cache so each document gets fresh numIds
    _bullet_numid_cache.clear()

    if template_path is None:
        template_path = os.path.join(_HERE, "template.docx")
    if output_path is None:
        output_path = os.path.join(_HERE, "output.docx")

    doc = DocxTemplate(template_path)
    sections = []
    table_data  = {}   # key -> rows
    rich_data   = {}   # key -> html

    toc_entries      = []   # (level, label, text)
    subheader_counts = {}   # sec_num -> int  (used by _inject_subheader_numbers)
    for i, section in enumerate(form_data.get("sections", [])):
        sec_num = str(section.get("section_number", i + 1))
        if section["type"] == "paragraph":
            key = f"RICH_SECTION_{i}"
            hdr = section.get("header", "")
            toc_entries.append((1, sec_num, hdr or f"Section {sec_num}"))
            html_body = section.get("html", "")
            sub_hdrs  = _extract_subheaders(html_body)
            # Record expected sub-header count so _inject_subheader_numbers can
            # stop numbering once all legitimate sub-headers have been stamped,
            # preventing static sections (T&C, Copyright…) from being mis-labelled.
            subheader_counts[sec_num] = len(sub_hdrs)
            for sub_i, sub_txt in enumerate(sub_hdrs, 1):
                toc_entries.append((2, f"{sec_num}.{sub_i}", sub_txt))
            sections.append({"type": "paragraph",
                             "header": hdr,
                             "section_number": sec_num,
                             "text":   section.get("text",""),
                             "placeholder": key})
            rich_data[key] = html_body or section.get("text","")
        elif section["type"] == "table":
            key = f"BOM_TABLE_{i}"
            bom_hdr = section.get("header", "").strip() or "Bill of Material"
            toc_entries.append((1, sec_num, bom_hdr))
            sections.append({"type": "table", "placeholder": key,
                             "section_number": sec_num,
                             "header": bom_hdr})
            rows = section.get("rows", [])
            table_data[key] = rows

    customer_image = None
    if form_data.get("customer_picture"):
        customer_image = InlineImage(doc, form_data["customer_picture"], width=Mm(50))

    def _safe_section(sec: dict) -> dict:
        """Return a copy of a section dict with all string values XML-escaped."""
        return {k: (_xml_safe(v) if isinstance(v, str) else v)
                for k, v in sec.items()}

    context = {
        "project_name":      _xml_safe(form_data.get("project_name","")),
        "customer_name":     _xml_safe(form_data.get("customer_name","")),
        "customer_location": _xml_safe(form_data.get("customer_location","")),
        "contact_info":      _xml_safe(form_data.get("contact_info","")),
        "proposal_number":   _xml_safe(form_data.get("proposal_number","")),
        "customer_picture":  customer_image,
        "today_date":        _xml_safe(form_data.get("today_date","")),
        "sections":          [_safe_section(sec) for sec in sections],
    }

    doc.render(context)
    doc.save(output_path)

    from docx import Document
    rendered = Document(output_path)

    import traceback as _tb
    for key, rows in table_data.items():
        try:
            _replace_placeholder_with_table(rendered, key, rows)
        except Exception:
            print(f"ERROR replacing table {key}:")
            _tb.print_exc()

    for key, html in rich_data.items():
        try:
            _replace_placeholder_with_rich(rendered, key, html)
        except Exception:
            print(f"ERROR replacing rich {key}:")
            _tb.print_exc()

    try:
        section_bm_map = _inject_section_numbers(rendered, sections)
        _inject_subheader_numbers(rendered, sections=sections,
                                  subheader_counts=subheader_counts)
        # Stamp static template sections (T&C, Copyright, etc.) that live
        # in the Word template but are not part of form_data["sections"].
        static_entries = _detect_and_stamp_static_headers(rendered, sections)
        toc_entries.extend(static_entries)
        # Upgrade toc_entries to 4-tuples (level, label, text, bookmark_name)
        # so _insert_toc can use PAGEREF fields for accurate page numbers.
        toc_entries = [
            (e[0], e[1], e[2], e[3] if len(e) > 3
             else (section_bm_map.get(e[1]) if e[0] == 1 else None))
            for e in toc_entries
        ]
        _set_header_outline_levels(rendered)
        _insert_toc(rendered, toc_entries)
    except Exception:
        import traceback as _tb2; _tb2.print_exc()

    rendered.save(output_path)

    # Post-save: force <w:updateFields w:val="1"/> into word/settings.xml so Word
    # recalculates the TOC page numbers on first open regardless of prior state.
    import zipfile as _zf, os as _os2, re as _re2
    _tmp_path = output_path + '.__settingstmp'
    try:
        with _zf.ZipFile(output_path, 'r') as _zin:
            with _zf.ZipFile(_tmp_path, 'w', _zf.ZIP_DEFLATED) as _zout:
                for _item in _zin.infolist():
                    _data = _zin.read(_item.filename)
                    if _item.filename == 'word/settings.xml':
                        _s = _data.decode('utf-8')
                        # Remove any existing updateFields element.
                        # Handle both the self-closing form (<w:updateFields .../>) and
                        # the paired form (<w:updateFields ...>...</w:updateFields>) so
                        # we don't end up with two conflicting elements (which causes
                        # Word to honour the first one, often val="0", and skip updates).
                        _s = _re2.sub(r'<w:updateFields[^>]*/>', '', _s)
                        _s = _re2.sub(r'<w:updateFields[^>]*>.*?</w:updateFields>',
                                      '', _s, flags=_re2.DOTALL)
                        # Insert with val="1" before closing tag
                        _s = _re2.sub(
                            r'</w:settings>',
                            '<w:updateFields w:val="1"/></w:settings>',
                            _s, count=1)
                        _data = _s.encode('utf-8')
                    _zout.writestr(_item, _data)
        _os2.replace(_tmp_path, output_path)
    except Exception:
        import traceback as _tb3; _tb3.print_exc()
        try: _os2.remove(_tmp_path)
        except Exception: pass

    return output_path
