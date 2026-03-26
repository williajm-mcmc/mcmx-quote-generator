import re
import base64
import tempfile
from html.parser import HTMLParser

# Module-level cache for bullet numId definitions — persists across calls
_fw_bullet_numid_cache = {}

from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QPushButton,
    QTableWidget, QTableWidgetItem, QTextEdit, QLineEdit, QLabel,
    QHeaderView, QSizePolicy, QGraphicsDropShadowEffect, QStyledItemDelegate,
    QDialog, QSpinBox, QDialogButtonBox, QFormLayout, QScrollArea,
    QDoubleSpinBox, QCheckBox, QToolBar, QFrame
)
from PyQt6.QtCore import Qt, QSize, QUrl, QRegularExpression
from PyQt6.QtGui import (
    QPixmap, QImage, QColor, QRegularExpressionValidator,
    QTextCursor, QTextTableFormat, QTextCharFormat, QTextBlockFormat,
    QTextListFormat, QFont, QAction
)
from PyQt6.QtNetwork import QNetworkAccessManager, QNetworkRequest, QNetworkReply


# ── Style tokens ──────────────────────────────────────────────────────────────
CARD_STYLE = """
    QWidget#card {
        background: #ffffff;
        border: 1px solid #dde1e7;
        border-radius: 8px;
    }
"""
HEADER_LABEL_STYLE = """
    QLabel {
        font-size: 12px; font-weight: 600;
        color: #1a1a2e; background: transparent; border: none;
    }
"""
REMOVE_BTN_STYLE = """
    QPushButton {
        background: transparent; color: #e05252;
        border: 1px solid #e05252; border-radius: 4px;
        padding: 4px 10px; font-size: 11px;
    }
    QPushButton:hover  { background: #fdf2f2; }
    QPushButton:pressed { background: #f9e0e0; }
"""
ACTION_BTN_STYLE = """
    QPushButton {
        background: #f4f6fa; color: #3a3a5c;
        border: 1px solid #dde1e7; border-radius: 4px;
        padding: 4px 10px; font-size: 11px;
    }
    QPushButton:hover  { background: #e8ecf5; border-color: #b0b8d0; }
    QPushButton:pressed { background: #dce2ef; }
"""
MARGIN_BTN_ACTIVE = """
    QPushButton {
        background: #e8f5e9; color: #2e7d32;
        border: 1px solid #81c784; border-radius: 4px;
        padding: 4px 10px; font-size: 11px; font-weight: 600;
    }
    QPushButton:hover { background: #c8e6c9; }
"""
TABLE_STYLE = """
    QTableWidget {
        border: 1px solid #dde1e7; border-radius: 4px;
        gridline-color: #edf0f5; background: #ffffff;
        color: #1a1a2e; font-size: 12px;
        selection-background-color: #e8f0fe;
        selection-color: #1a1a2e;
    }
    QTableWidget::item { padding: 6px 8px; color: #1a1a2e; }
    QTableWidget::item:alternate { background: #f8f9fc; }
    QHeaderView::section {
        background: #f4f6fa; color: #3a3a5c;
        font-weight: 600; font-size: 11px;
        padding: 6px 8px; border: none;
        border-right: 1px solid #dde1e7;
        border-bottom: 1px solid #dde1e7;
    }
"""
FIELD_STYLE = """
    QLineEdit {
        border: 1px solid #dde1e7; border-radius: 4px;
        padding: 6px 10px; font-size: 12px;
        color: #1a1a2e; background: #fafbfd;
    }
    QLineEdit:focus { border-color: #6a8fd8; background: #ffffff; }
"""
TEXTEDIT_STYLE = (
    "QTextEdit {"
    " border: 1px solid #dde1e7;"
    " border-radius: 4px;"
    " padding-left: 8px;"
    " padding-top: 4px;"
    " padding-bottom: 4px;"
    " font-size: 12px;"
    " color: #1a1a2e;"
    " background: #fafbfd;"
    "}"
    "QTextEdit:focus { border-color: #6a8fd8; background: #ffffff; }"
)
EDITOR_STYLE = (
    "QLineEdit { color: #1a1a2e; background: #ffffff; "
    "border: 1px solid #6a8fd8; border-radius: 0px; "
    "padding: 0px; margin: 0px; font-size: 12px; }"
)


def _add_shadow(widget):
    s = QGraphicsDropShadowEffect()
    s.setBlurRadius(12); s.setOffset(0, 2)
    s.setColor(QColor(0, 0, 0, 28))
    widget.setGraphicsEffect(s)


import math as _math

def _apply_margin(cost: float, margin_pct: float) -> float:
    """margin_pct is 0-100. e.g. 20% → cost / 0.80, rounded up to nearest dollar."""
    if margin_pct <= 0:
        return cost
    if margin_pct >= 100:
        return cost
    return _math.ceil(cost / (1.0 - margin_pct / 100.0))


# ── Delegates ─────────────────────────────────────────────────────────────────
class PlainTextDelegate(QStyledItemDelegate):
    def createEditor(self, parent, option, index):
        e = QLineEdit(parent); e.setStyleSheet(EDITOR_STYLE); return e
    def updateEditorGeometry(self, editor, option, index):
        editor.setGeometry(option.rect)


class RichTextDelegate(QStyledItemDelegate):
    """Delegate that renders stored HTML (UserRole) in the cell instead of plain text."""
    def paint(self, painter, option, index):
        html = index.data(Qt.ItemDataRole.UserRole)
        if not html:
            super().paint(painter, option, index)
            return
        from PyQt6.QtWidgets import QStyleOptionViewItem, QApplication, QStyle
        from PyQt6.QtGui import QTextDocument, QAbstractTextDocumentLayout
        from PyQt6.QtCore import QSizeF
        # Draw the base background / selection highlight
        opt = QStyleOptionViewItem(option)
        self.initStyleOption(opt, index)
        style = opt.widget.style() if opt.widget else QApplication.style()
        style.drawPrimitive(QStyle.PrimitiveElement.PE_PanelItemViewItem, opt, painter, opt.widget)
        # Render HTML — use a fixed Aptos 11pt font so bullets/formatting
        # are not affected by whatever font the table widget inherits.
        doc = QTextDocument()
        from PyQt6.QtGui import QFont as _QFont
        doc.setDefaultFont(_QFont('Aptos', 11))
        doc.setIndentWidth(20)   # match CellEditorDialog indent so bullets render correctly
        doc.setHtml(html)
        doc.setTextWidth(option.rect.width() - 8)
        painter.save()
        painter.translate(option.rect.left() + 4, option.rect.top() + 4)
        ctx = QAbstractTextDocumentLayout.PaintContext()
        doc.documentLayout().draw(painter, ctx)
        painter.restore()

    def sizeHint(self, option, index):
        html = index.data(Qt.ItemDataRole.UserRole)
        if not html:
            return super().sizeHint(option, index)
        from PyQt6.QtGui import QTextDocument
        # Use the widget's actual column width so measurement is accurate
        # even when called before the row has been laid out
        w = option.rect.width()
        if w < 10 and option.widget:
            try:
                col = index.column()
                w = option.widget.columnWidth(col)
            except Exception:
                w = 200
        if w < 10:
            w = 200
        doc = QTextDocument()
        from PyQt6.QtGui import QFont as _QFont
        doc.setDefaultFont(_QFont('Aptos', 11))
        doc.setIndentWidth(20)
        doc.setHtml(html)
        doc.setTextWidth(w - 8)
        h = int(doc.size().height()) + 12
        return QSize(super().sizeHint(option, index).width(), max(32, h))

class NumericDelegate(QStyledItemDelegate):
    def createEditor(self, parent, option, index):
        e = QLineEdit(parent); e.setStyleSheet(EDITOR_STYLE)
        e.setValidator(QRegularExpressionValidator(
            QRegularExpression(r"^\d*\.?\d*$"), e)); return e
    def updateEditorGeometry(self, editor, option, index):
        editor.setGeometry(option.rect)



class CellEditorDialog(QDialog):
    """
    Pop-up editor for a BOM Part Number cell.
    Shows a mini toolbar (Bold, Bullet, →Indent, ←Outdent) above a QTextEdit
    that supports Enter=newline, Tab=indent, Shift+Tab=outdent.
    """
    def __init__(self, html="", plain="", parent=None):
        super().__init__(parent)
        self.setWindowTitle("Edit Item")
        self.setMinimumWidth(480)
        self.setMinimumHeight(200)
        self.resize(520, 300)
        self.setSizeGripEnabled(True)
        self.setStyleSheet(
            "QDialog { background:#ffffff; }"
            "QToolBar { background:#f4f6fa; border:1px solid #dde1e7; spacing:2px; }"
            "QPushButton { background:transparent; border:1px solid transparent;"
            "  border-radius:3px; font-size:12px; color:#3a3a5c;"
            "  padding:2px 6px; min-width:24px; min-height:22px; }"
            "QPushButton:hover { background:#e0e4ef; border-color:#c0c8d8; }"
            "QPushButton:pressed { background:#d0d8ef; }"
            "QTextEdit { background:#ffffff; color:#1a1a2e; border:1px solid #dde1e7;"
            "  border-radius:4px; padding:6px; font-size:12px; }")

        layout = QVBoxLayout(self)
        layout.setContentsMargins(8, 8, 8, 8)
        layout.setSpacing(4)

        # Toolbar
        tb = QWidget()
        tb.setStyleSheet("QWidget { background:#f4f6fa; border:1px solid #dde1e7;"
                         " border-radius:4px; }")
        tb_row = QHBoxLayout(tb)
        tb_row.setContentsMargins(4, 2, 4, 2); tb_row.setSpacing(2)

        def _btn(text, tip):
            b = QPushButton(text); b.setToolTip(tip)
            b.setFocusPolicy(Qt.FocusPolicy.NoFocus)
            b.setFixedSize(28, 24)
            return b

        _active = ("QPushButton:checked{background:#f5d0da;"
                   "border-color:#920d2e;color:#920d2e;}")

        self._btn_bold   = _btn("B",   "Bold (Ctrl+B)")
        self._btn_bold.setStyleSheet(self._btn_bold.styleSheet() +
                                     "QPushButton{font-weight:700;}" + _active)
        self._btn_italic = _btn("I",   "Italic (Ctrl+I)")
        self._btn_italic.setStyleSheet(self._btn_italic.styleSheet() +
                                       "QPushButton{font-style:italic;}" + _active)
        self._btn_bullet  = _btn("•≡", "Toggle bullet list")

        _sep1 = QFrame(); _sep1.setFrameShape(QFrame.Shape.VLine)
        _sep1.setStyleSheet("color:#dde1e7;")
        _sep2 = QFrame(); _sep2.setFrameShape(QFrame.Shape.VLine)
        _sep2.setStyleSheet("color:#dde1e7;")

        self._btn_align_l = _btn("⬤≡",  "Align left")
        self._btn_align_c = _btn("≡⬤≡", "Align center")
        self._btn_align_r = _btn("≡⬤",  "Align right")
        self._btn_align_c.setFixedSize(32, 24)

        for w in (self._btn_bold, self._btn_italic,
                  _sep1,
                  self._btn_bullet,
                  _sep2,
                  self._btn_align_l, self._btn_align_c, self._btn_align_r):
            tb_row.addWidget(w)
        tb_row.addStretch()
        layout.addWidget(tb)

        # Editor
        self.editor = _CellTextEdit()
        # Reduce per-level indent from Qt's default 40 px to something that
        # matches typical Word bullet indentation visually.
        self.editor.document().setIndentWidth(20)
        self.editor.document().contentsChanged.connect(self._on_contents_changed)
        layout.addWidget(self.editor)

        # OK / Cancel
        btns = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok |
            QDialogButtonBox.StandardButton.Cancel)
        btns.setStyleSheet(
            "QPushButton { background:#f4f6fa; color:#1a1a2e;"
            "  border:1px solid #c8cedd; border-radius:4px;"
            "  padding:5px 16px; min-width:64px; }"
            "QPushButton:hover { background:#e0e6f0; }"
            "QPushButton:pressed { background:#d0d8ef; }")
        btns.accepted.connect(self.accept)
        btns.rejected.connect(self.reject)
        layout.addWidget(btns)

        # Default character format: Aptos 11pt black, not bold
        _def = QTextCharFormat()
        _def.setFontFamily('Aptos')
        _def.setFontPointSize(11.0)
        _def.setFontWeight(QFont.Weight.Normal)
        _def.setForeground(QColor(0x1a, 0x1a, 0x2e))
        _df = QFont('Aptos', 11)
        _df.setWeight(QFont.Weight.Normal)
        self.editor.document().setDefaultFont(_df)
        # Load content
        if html:
            self.editor.setHtml(html)
        else:
            self.editor.setPlainText(plain)
        # setCurrentCharFormat must come AFTER loading — setPlainText/setHtml
        # resets the cursor, discarding any format set before loading.
        self.editor.setCurrentCharFormat(_def)

        # Wire toolbar
        self._btn_bold.setCheckable(True)
        self._btn_italic.setCheckable(True)
        self._btn_bold.clicked.connect(self._toggle_bold)
        self._btn_italic.clicked.connect(self._toggle_italic)
        self._btn_bullet.clicked.connect(self._toggle_bullet)
        self._btn_align_l.clicked.connect(
            lambda: self._set_alignment(Qt.AlignmentFlag.AlignLeft))
        self._btn_align_c.clicked.connect(
            lambda: self._set_alignment(Qt.AlignmentFlag.AlignHCenter))
        self._btn_align_r.clicked.connect(
            lambda: self._set_alignment(Qt.AlignmentFlag.AlignRight))
        self.editor.currentCharFormatChanged.connect(self._sync_fmt_btns)

    def _on_contents_changed(self):
        """Grow the dialog vertically as content is added."""
        doc_h = int(self.editor.document().size().height())
        target = min(500, max(160, doc_h + 120))
        if self.height() < target:
            self.resize(self.width(), target)

    def _sync_fmt_btns(self, fmt):
        self._btn_bold.setChecked(fmt.fontWeight() == QFont.Weight.Bold)
        self._btn_italic.setChecked(fmt.fontItalic())

    def _toggle_bold(self):
        fmt = QTextCharFormat()
        c = self.editor.textCursor()
        fmt.setFontWeight(
            QFont.Weight.Normal if c.charFormat().fontWeight() == QFont.Weight.Bold
            else QFont.Weight.Bold)
        c.mergeCharFormat(fmt)
        self.editor.setTextCursor(c)

    def _toggle_italic(self):
        fmt = QTextCharFormat()
        c = self.editor.textCursor()
        fmt.setFontItalic(not c.charFormat().fontItalic())
        c.mergeCharFormat(fmt)
        self.editor.setTextCursor(c)

    def _set_alignment(self, alignment):
        self.editor.setAlignment(alignment)

    def _toggle_bullet(self):
        cursor = self.editor.textCursor()
        lst = cursor.currentList()
        if lst:
            # Detach from list then clear block indent
            block = cursor.block()
            lst.remove(block)
            bf = QTextBlockFormat()
            bf.setIndent(0)
            cursor.setBlockFormat(bf)
        else:
            fmt = QTextListFormat()
            fmt.setStyle(QTextListFormat.Style.ListDisc)
            fmt.setIndent(1)
            cursor.createList(fmt)
        self.editor.setTextCursor(cursor)
        self.editor.setFocus()

    def _indent(self):
        cursor = self.editor.textCursor()
        if cursor.currentList():
            _safe_indent(cursor, self.editor)

    def _outdent(self):
        cursor = self.editor.textCursor()
        lst = cursor.currentList()
        if lst:
            _safe_outdent(cursor, lst, lst.format(), self.editor)

    def html(self):  return self.editor.toHtml()
    def plain(self): return self.editor.toPlainText()


class _CellTextEdit(QTextEdit):
    """QTextEdit with bullet/tab keyboard shortcuts for CellEditorDialog."""
    _PASTE_FAMILY = 'Aptos'
    _PASTE_SIZE   = 11.0
    _PASTE_COLOR  = (0x1a, 0x1a, 0x2e)

    def insertFromMimeData(self, source):
        start = self.textCursor().position()
        super().insertFromMimeData(source)
        # Normalise pasted text to Aptos 11pt black
        doc = self.document()
        end = self.textCursor().position()
        if end <= start:
            return
        r, g, b = self._PASTE_COLOR
        fmt = QTextCharFormat()
        fmt.setFontFamily(self._PASTE_FAMILY)
        fmt.setFontPointSize(self._PASTE_SIZE)
        fmt.setForeground(QColor(r, g, b))
        block = doc.findBlock(start)
        while block.isValid() and block.position() <= end:
            it = block.begin()
            while not it.atEnd():
                frag = it.fragment()
                if frag.isValid():
                    fs, fe = frag.position(), frag.position() + frag.length()
                    if fe > start and fs < end:
                        c = QTextCursor(doc)
                        c.setPosition(max(start, fs))
                        c.setPosition(min(end, fe), QTextCursor.MoveMode.KeepAnchor)
                        c.mergeCharFormat(fmt)
                it += 1
            block = block.next()

    def keyPressEvent(self, event):
        key = event.key()
        if key == Qt.Key.Key_Backtab:
            cursor = self.textCursor()
            lst = cursor.currentList()
            if lst:
                _safe_outdent(cursor, lst, lst.format(), self)
            return
        if key == Qt.Key.Key_Tab:
            cursor = self.textCursor()
            if not cursor.currentList():
                # No list yet — create a level-1 bullet and stop.
                # Do NOT call _safe_indent afterwards or the new bullet would
                # immediately jump to level 2 and appear to vanish.
                fmt = QTextListFormat()
                fmt.setStyle(QTextListFormat.Style.ListDisc)
                fmt.setIndent(1)
                cursor.createList(fmt)
                return
            # Already in a list — indent to the next level.
            _safe_indent(cursor, self)
            return
        if key in (Qt.Key.Key_Return, Qt.Key.Key_Enter):
            cursor = self.textCursor()
            lst = cursor.currentList()
            if lst:
                try:
                    if not cursor.block().text().strip():
                        cursor.beginEditBlock()
                        bf = QTextBlockFormat(); bf.setIndent(0)
                        cursor.setBlockFormat(bf)
                        live = cursor.block().textList()
                        if live:
                            live.remove(cursor.block())
                        cursor.endEditBlock()
                        self.setTextCursor(cursor)
                    else:
                        saved = QTextListFormat(lst.format())
                        cursor.insertBlock()
                        if cursor.block().textList() is None:
                            cursor.createList(saved)
                        self.setTextCursor(cursor)
                except RuntimeError:
                    super().keyPressEvent(event)
                return
        super().keyPressEvent(event)


# ── DragDropLabel ─────────────────────────────────────────────────────────────
class DragDropLabel(QLabel):
    IDLE_STYLE = "QLabel { background:#f8f9fc; border:2px dashed #c0c8d8; border-radius:8px; color:#8892a4; font-size:11px; }"
    HOVER_STYLE = "QLabel { background:#eef2fb; border:2px dashed #6a8fd8; border-radius:8px; color:#3a5bd9; font-size:11px; }"
    FILLED_STYLE = "QLabel { background:#ffffff; border:2px solid #c0c8d8; border-radius:8px; }"

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAcceptDrops(True)
        self.setFocusPolicy(Qt.FocusPolicy.StrongFocus)
        self.setScaledContents(False)
        self.setFixedSize(150, 150)
        self.image_path = None
        self._nam = QNetworkAccessManager(self)
        self._reset_label()
        # Ensure all ancestor widgets pass drag events through
        self._enable_ancestor_drops()

    def _enable_ancestor_drops(self):
        """Walk up parent chain and enable acceptDrops on every ancestor."""
        p = self.parent()
        while p is not None:
            try:
                p.setAcceptDrops(True)
                p = p.parent()
            except Exception:
                break

    def showEvent(self, event):
        """Re-run ancestor drop setup after widget is inserted into layout."""
        super().showEvent(event)
        self._enable_ancestor_drops()

    def _reset_label(self):
        self.clear(); self.setText("📎  Drop image\nor Ctrl+V")
        self.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.setStyleSheet(self.IDLE_STYLE)

    @staticmethod
    def _to_qurl(raw):
        return raw if isinstance(raw, QUrl) else QUrl(str(raw))

    @staticmethod
    def _src_from_html(html):
        m = re.search(r'src=["\']((https?|data):[^"\']+)["\']', html, re.I)
        return m.group(1) if m else None

    @staticmethod
    def _qimage_from_data_url(data_url):
        try:
            _, b64 = data_url.split(",", 1)
            raw = base64.b64decode(b64)
            img = QImage(); img.loadFromData(bytes(raw))
            return img if not img.isNull() else None
        except Exception:
            return None

    def _is_image_drag(self, mime):
        if mime.hasImage(): return True
        if mime.hasUrls():
            url = self._to_qurl(mime.urls()[0])
            local = url.toLocalFile()
            if local and local.lower().endswith((".png",".jpg",".jpeg",".bmp",".gif",".webp")): return True
            if url.scheme() in ("http","https","data"): return True
        if mime.hasHtml() and self._src_from_html(mime.html()): return True
        if mime.hasText() and mime.text().strip().startswith("data:image/"): return True
        return False

    def dragEnterEvent(self, event):
        if self._is_image_drag(event.mimeData()):
            self.setStyleSheet(self.HOVER_STYLE); event.acceptProposedAction()
        else: event.ignore()

    def dragLeaveEvent(self, event):
        if not self.image_path: self.setStyleSheet(self.IDLE_STYLE)

    def dropEvent(self, event):
        mime = event.mimeData()
        if mime.hasImage():
            qi = mime.imageData()
            if qi and not qi.isNull():
                self.set_image_from_qimage(qi); event.acceptProposedAction(); return
        # HTML src= checked FIRST — Chrome puts the image URL here,
        # while hasUrls() contains the page URL, not the image URL
        if mime.hasHtml():
            img_src = self._src_from_html(mime.html())
            if img_src:
                if img_src.startswith("data:image/"):
                    img = self._qimage_from_data_url(img_src)
                    if img: self.set_image_from_qimage(img); event.acceptProposedAction(); return
                else:
                    self._download_and_set(QUrl(img_src))
                    event.acceptProposedAction(); return
        # Local file drag from Explorer
        if mime.hasUrls():
            url = self._to_qurl(mime.urls()[0])
            local = url.toLocalFile()
            if local and local.lower().endswith(
                    (".png",".jpg",".jpeg",".bmp",".gif",".webp",".svg")):
                self.set_image(local); event.acceptProposedAction(); return
            if url.scheme() == "data":
                img = self._qimage_from_data_url(url.toString())
                if img: self.set_image_from_qimage(img); event.acceptProposedAction(); return
        if mime.hasText():
            t = mime.text().strip()
            if t.startswith("data:image/"):
                img = self._qimage_from_data_url(t)
                if img: self.set_image_from_qimage(img); event.acceptProposedAction(); return
        event.ignore()

    def _download_and_set(self, url):
        self.setText("Loading…")
        req = QNetworkRequest(url)
        req.setRawHeader(b"User-Agent",
            b"Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            b"AppleWebKit/537.36 (KHTML, like Gecko) "
            b"Chrome/120.0.0.0 Safari/537.36")
        req.setRawHeader(b"Accept",
            b"image/avif,image/webp,image/apng,image/jpeg,image/*,*/*;q=0.8")
        req.setRawHeader(b"Accept-Language", b"en-US,en;q=0.9")
        referer = (url.scheme() + "://" + url.host()).encode()
        req.setRawHeader(b"Referer", referer)
        req.setAttribute(
            QNetworkRequest.Attribute.RedirectPolicyAttribute, 3)
        reply = self._nam.get(req)
        reply.finished.connect(lambda: self._on_download_finished(reply))

    def _on_download_finished(self, reply):
        from PyQt6.QtNetwork import QNetworkReply
        if reply.error() != QNetworkReply.NetworkError.NoError:
            self._reset_label(); self.setText("Download\nfailed")
            reply.deleteLater(); return
        data = reply.readAll()
        raw  = bytes(data)
        url_lower = reply.url().toString().lower()
        ct = bytes(reply.rawHeader(b'Content-Type')).decode(errors='replace').lower()
        # SVG: render via Qt's SVG module
        is_svg = ('svg' in ct or url_lower.endswith('.svg')
                  or raw[:500].lower().find(b'<svg') != -1)
        if is_svg:
            try:
                from PyQt6.QtSvg import QSvgRenderer
                from PyQt6.QtGui import QPainter
                renderer = QSvgRenderer(data)
                if renderer.isValid():
                    sz = renderer.defaultSize()
                    w = sz.width()  if sz.width()  > 0 else 400
                    h = sz.height() if sz.height() > 0 else 400
                    img = QImage(w, h, QImage.Format.Format_ARGB32)
                    img.fill(0)
                    painter = QPainter(img)
                    renderer.render(painter)
                    painter.end()
                    if not img.isNull():
                        self.set_image_from_qimage(img)
                        reply.deleteLater(); return
            except Exception:
                pass
        # Raster: QImage → QPixmap → Pillow
        img = QImage()
        if img.loadFromData(data) and not img.isNull():
            self.set_image_from_qimage(img); reply.deleteLater(); return
        px = QPixmap()
        if px.loadFromData(data) and not px.isNull():
            self._display_pixmap(px)
            import tempfile
            tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
            tmp.close(); px.save(tmp.name, 'PNG')
            self.image_path = tmp.name; reply.deleteLater(); return
        try:
            from PIL import Image as _PIL
            import io, tempfile
            pil = _PIL.open(io.BytesIO(raw)).convert('RGBA')
            tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
            pil.save(tmp.name, 'PNG'); tmp.close()
            self.set_image(tmp.name); reply.deleteLater(); return
        except Exception:
            pass
        self._reset_label(); self.setText("Not an image")
        reply.deleteLater()

    def keyPressEvent(self, event):
        if event.key() == Qt.Key.Key_V and event.modifiers() == Qt.KeyboardModifier.ControlModifier:
            from PyQt6.QtWidgets import QApplication
            img = QApplication.clipboard().image()
            if not img.isNull(): self.set_image_from_qimage(img)
        else: super().keyPressEvent(event)

    def set_image(self, path):
        px = QPixmap(path)
        if px.isNull(): return
        self._display_pixmap(px); self.image_path = path

    def set_image_from_qimage(self, qi):
        px = QPixmap.fromImage(qi)
        if px.isNull(): return
        self._display_pixmap(px)
        tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        tmp.close(); qi.save(tmp.name, "PNG"); self.image_path = tmp.name

    def clear_image(self):
        self.image_path = None; self._reset_label()

    def _display_pixmap(self, px):
        self.setPixmap(px.scaled(self.width(), self.height(),
            Qt.AspectRatioMode.KeepAspectRatio,
            Qt.TransformationMode.SmoothTransformation))
        self.setText(""); self.setStyleSheet(self.FILLED_STYLE)

    def mousePressEvent(self, event):
        self.setFocus(); super().mousePressEvent(event)


# ── Auto-height table ─────────────────────────────────────────────────────────
class AutoHeightTable(QTableWidget):
    def __init__(self, cols):
        super().__init__(0, cols)
        self.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum)
        self.setStyleSheet(TABLE_STYLE); self.setAlternatingRowColors(True)
        self.verticalHeader().setDefaultSectionSize(56)

    def sizeHint(self):
        h = self.horizontalHeader().height()
        for i in range(self.rowCount()): h += self.rowHeight(i)
        return QSize(super().sizeHint().width(), h + self.frameWidth() * 2)

    def minimumSizeHint(self): return self.sizeHint()


# ── Collapsible card base ─────────────────────────────────────────────────────
class CollapsibleCard(QWidget):
    def __init__(self, title, parent=None):
        super().__init__(parent)
        self.setObjectName("card"); self.setStyleSheet(CARD_STYLE)
        _add_shadow(self)
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum)

        outer = QVBoxLayout(self)
        outer.setContentsMargins(0,0,0,0); outer.setSpacing(0)

        self._header_bar = QWidget(); self._header_bar.setObjectName("cardHeader")
        self._header_bar.setStyleSheet("""
            QWidget#cardHeader {
                background:#f4f6fa; border-bottom:1px solid #dde1e7;
                border-top-left-radius:8px; border-top-right-radius:8px;
            }""")
        self._header_bar.setCursor(Qt.CursorShape.PointingHandCursor)
        self._header_bar.setFixedHeight(42)
        hr = QHBoxLayout(self._header_bar)
        hr.setContentsMargins(12,7,12,7); hr.setSpacing(8)

        self._toggle_btn = QPushButton("▼"); self._toggle_btn.setFixedSize(20,20)
        self._toggle_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self._toggle_btn.setStyleSheet(
            "QPushButton{background:transparent;border:none;font-size:9px;color:#6672a0;padding:0}"
            "QPushButton:hover{color:#1a1a2e}")
        self._toggle_btn.clicked.connect(self.toggle)

        self._title_label = QLabel(title); self._title_label.setStyleSheet(HEADER_LABEL_STYLE)
        hr.addWidget(self._toggle_btn); hr.addWidget(self._title_label); hr.addStretch()
        self._header_extras = QHBoxLayout(); self._header_extras.setSpacing(6)
        # Up/Down reorder buttons
        _arrow_style = (
            "QPushButton{background:transparent;border:1px solid transparent;"
            "border-radius:3px;font-size:11px;color:#920d2e;padding:0 3px;}"
            "QPushButton:hover{background:#f5d0da;border-color:#920d2e;}")
        self._btn_up = QPushButton("▲")
        self._btn_up.setFixedSize(22, 22)
        self._btn_up.setToolTip("Move section up")
        self._btn_up.setStyleSheet(_arrow_style)
        self._btn_down = QPushButton("▼")
        self._btn_down.setFixedSize(22, 22)
        self._btn_down.setToolTip("Move section down")
        self._btn_down.setStyleSheet(_arrow_style)
        self._header_extras.addWidget(self._btn_up)
        self._header_extras.addWidget(self._btn_down)
        hr.addLayout(self._header_extras)
        outer.addWidget(self._header_bar)

        self._body = QWidget()
        self._body.setStyleSheet("background:transparent;border:none;")
        self._body_layout = QVBoxLayout(self._body)
        self._body_layout.setContentsMargins(14,12,14,12); self._body_layout.setSpacing(8)
        outer.addWidget(self._body)
        self._collapsed = False

    def toggle(self):
        self._collapsed = not self._collapsed
        self._body.setVisible(not self._collapsed)
        self._toggle_btn.setText("▶" if self._collapsed else "▼")
        self.updateGeometry()

    def set_title(self, text): self._title_label.setText(text)

    def mousePressEvent(self, event):
        if self._header_bar.geometry().contains(event.pos()): self.toggle()
        super().mousePressEvent(event)


# ── Custom table dialog ───────────────────────────────────────────────────────
class InsertTableDialog(QDialog):
    """
    Configure a table: rows, columns, optional total row with summed columns.
    Row 0 = header (auto-styled). Optional last row = totals.
    """
    _STYLE = """
        QDialog  { background:#ffffff; color:#1a1a2e; }
        QLabel   { color:#1a1a2e; background:transparent; }
        QSpinBox { background:#ffffff; color:#1a1a2e;
                   border:1px solid #dde1e7; border-radius:4px; padding:4px 8px; }
        QCheckBox { color:#1a1a2e; }
        QPushButton { background:#f4f6fa; color:#1a1a2e;
                      border:1px solid #c8cedd; border-radius:4px;
                      padding:5px 16px; min-width:64px; }
        QPushButton:hover   { background:#e0e6f0; }
        QPushButton:pressed { background:#d0d8ef; }
    """

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Insert Table")
        self.setMinimumWidth(340)
        self.setSizeGripEnabled(True)
        self.setStyleSheet(self._STYLE)

        layout = QVBoxLayout(self)
        layout.setSpacing(10)

        note = QLabel("Row 1 is automatically styled as the header row.")
        note.setStyleSheet(
            "color:#6672a0; font-size:11px; background:#f4f6fa; "
            "border:1px solid #dde1e7; border-radius:4px; padding:6px 10px;")
        note.setWordWrap(True)
        layout.addWidget(note)

        form = QFormLayout()
        form.setSpacing(8)

        self.rows_spin = QSpinBox()
        self.rows_spin.setRange(2, 50)
        self.rows_spin.setValue(4)
        self.cols_spin = QSpinBox()
        self.cols_spin.setRange(1, 20)
        self.cols_spin.setValue(3)
        self.cols_spin.valueChanged.connect(self._on_cols_changed)

        form.addRow("Rows (incl. header):", self.rows_spin)
        form.addRow("Columns:", self.cols_spin)
        layout.addLayout(form)

        # ── Total row section ─────────────────────────────────────────────────
        sep = QFrame(); sep.setFrameShape(QFrame.Shape.HLine)
        sep.setStyleSheet("color:#dde1e7;")
        layout.addWidget(sep)

        self.total_cb = QCheckBox("Add a Total row at the bottom")
        self.total_cb.setStyleSheet("QCheckBox { font-weight:600; color:#1a1a2e; }")
        self.total_cb.stateChanged.connect(self._on_total_toggled)
        layout.addWidget(self.total_cb)

        # Column checkboxes — shown only when total row is checked
        self._sum_frame = QWidget()
        self._sum_frame.setVisible(False)
        sf_layout = QVBoxLayout(self._sum_frame)
        sf_layout.setContentsMargins(16, 4, 0, 0)
        sf_layout.setSpacing(4)
        lbl = QLabel("Columns to sum:")
        lbl.setStyleSheet("color:#3a3a5c; font-size:11px;")
        sf_layout.addWidget(lbl)
        self._col_checks_layout = QVBoxLayout()
        self._col_checks_layout.setSpacing(2)
        sf_layout.addLayout(self._col_checks_layout)
        layout.addWidget(self._sum_frame)

        self._col_checkboxes = []
        self._rebuild_col_checks(self.cols_spin.value())

        btns = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok |
            QDialogButtonBox.StandardButton.Cancel)
        btns.accepted.connect(self.accept)
        btns.rejected.connect(self.reject)
        layout.addWidget(btns)

    def _on_cols_changed(self, n):
        self._rebuild_col_checks(n)

    def _rebuild_col_checks(self, n):
        # Remove old checkboxes
        for cb in self._col_checkboxes:
            cb.deleteLater()
        self._col_checkboxes.clear()
        for i in range(n):
            cb = QCheckBox(f"Column {i + 1}")
            cb.setStyleSheet("QCheckBox { color:#1a1a2e; font-size:11px; }")
            self._col_checks_layout.addWidget(cb)
            self._col_checkboxes.append(cb)

    def _on_total_toggled(self, state):
        self._sum_frame.setVisible(
            state == Qt.CheckState.Checked.value or
            state == int(Qt.CheckState.Checked.value))
        self.adjustSize()

    def get_config(self):
        """
        Returns (rows, cols, has_total, sum_cols)
        sum_cols: list of 0-based column indices to sum (empty if no total row)
        """
        has_total = self.total_cb.isChecked()
        sum_cols  = [i for i, cb in enumerate(self._col_checkboxes)
                     if cb.isChecked()] if has_total else []
        return self.rows_spin.value(), self.cols_spin.value(), has_total, sum_cols


# ── Rich-text toolbar + editor ────────────────────────────────────────────────
def _style_qt_table_header_row(table, cols: int):
    """
    Apply bold + light gray background to all cells in row 0 of a QTextTable.
    Qt's QTextEdit has no direct cell background API so we use a char format
    with a background color applied to the entire cell content.
    """
    gray = QColor("#d9d9d9")
    for c_i in range(cols):
        try:
            cell = table.cellAt(0, c_i)
            cc = cell.firstCursorPosition()
            # Select the whole cell content
            end = cell.lastCursorPosition()
            cc.setPosition(cell.firstPosition())
            cc.setPosition(cell.lastPosition(),
                           QTextCursor.MoveMode.KeepAnchor)
            fmt = QTextCharFormat()
            fmt.setFontWeight(QFont.Weight.Bold)
            fmt.setBackground(gray)
            cc.mergeCharFormat(fmt)
        except Exception:
            pass


def _parse_html_table_to_qt(html: str, cursor, editor):
    """
    Parse an HTML <table> (e.g. pasted from Word) and insert it as a
    QTextTable into the editor at the given cursor position.
    Styles the first row as a header automatically.
    """
    from html.parser import HTMLParser as _HP

    class _N:
        __slots__ = ('tag', 'attrs', 'children', 'text')
        def __init__(self, tag, attrs=None):
            self.tag = (tag or '').lower()
            self.attrs = dict(attrs or {})
            self.children = []
            self.text = ''

    class _B(_HP):
        def __init__(self):
            super().__init__()
            self.root = _N('root')
            self._s = [self.root]
        def handle_starttag(self, tag, attrs):
            n = _N(tag, attrs)
            self._s[-1].children.append(n)
            self._s.append(n)
        def handle_endtag(self, _):
            if len(self._s) > 1: self._s.pop()
        def handle_data(self, d):
            self._s[-1].text += d
        def handle_entityref(self, name):
            m = {'amp':'&','lt':'<','gt':'>','quot':'"','nbsp':' ','apos':"'"}
            self._s[-1].text += m.get(name, '')

    b = _B(); b.feed(html)

    # Find the first <table> node anywhere in the tree
    def _find_table(node):
        if node.tag == 'table': return node
        for ch in node.children:
            r = _find_table(ch)
            if r: return r
        return None

    tbl_node = _find_table(b.root)
    if not tbl_node:
        return False  # no table found — caller should fall back to default paste

    # Collect rows
    sections = [c for c in tbl_node.children if c.tag in ('tbody','thead','tfoot')]
    if not sections: sections = [tbl_node]
    all_trs = []
    for sec in sections:
        for ch in sec.children:
            if ch.tag == 'tr': all_trs.append(ch)
    if not all_trs: return False

    max_cols = max(
        sum(1 for c in tr.children if c.tag in ('td','th'))
        for tr in all_trs
    ) or 1

    def _cell_text(node):
        """Recursively collect all text from a cell node."""
        parts = []
        if node.text: parts.append(node.text)
        for ch in node.children: parts.append(_cell_text(ch))
        return ''.join(parts)

    def _cell_bold(node):
        if node.tag in ('b','strong'): return True
        s = node.attrs.get('style','').lower()
        return 'font-weight:bold' in s or 'font-weight:700' in s or 'font-weight:600' in s

    from PyQt6.QtGui import QTextLength as _QTL
    fmt = QTextTableFormat()
    fmt.setCellPadding(4); fmt.setCellSpacing(0)
    fmt.setBorderStyle(QTextTableFormat.BorderStyle.BorderStyle_Solid)
    fmt.setBorder(1)
    fmt.setWidth(_QTL(_QTL.Type.PercentageLength, 100))

    table = cursor.insertTable(len(all_trs), max_cols, fmt)

    gray = QColor("#d9d9d9")
    for r_i, tr in enumerate(all_trs):
        cell_nodes = [c for c in tr.children if c.tag in ('td','th')]
        is_header = r_i == 0 or all(c.tag == 'th' for c in cell_nodes)
        for c_i, cn in enumerate(cell_nodes[:max_cols]):
            try:
                cell = table.cellAt(r_i, c_i)
                cc = cell.firstCursorPosition()
                text = _cell_text(cn).strip()
                cf = QTextCharFormat()
                if is_header or _cell_bold(cn):
                    cf.setFontWeight(QFont.Weight.Bold)
                if is_header:
                    cf.setBackground(gray)
                cc.insertText(text, cf)
            except Exception:
                pass

    editor.setTextCursor(cursor)
    return True


def _safe_indent(cursor, widget):
    """Indent the current list block by one level without creating a blank line."""
    lst = cursor.currentList()
    if not lst:
        return
    new_fmt = QTextListFormat(lst.format())
    new_fmt.setIndent(lst.format().indent() + 1)
    cursor.createList(new_fmt)


def _safe_outdent(cursor, lst, fmt, widget):
    if not lst:
        return

    cursor.beginEditBlock()
    try:
        current_indent = fmt.indent()

        if current_indent > 1:
            new_fmt = QTextListFormat(fmt)
            new_fmt.setIndent(current_indent - 1)
            cursor.createList(new_fmt)
        else:
            # Remove bullet completely — must detach from list then clear format
            current_block = cursor.block()
            lst.remove(current_block)
            bf = QTextBlockFormat()
            bf.setIndent(0)
            cursor.setBlockFormat(bf)

    finally:
        cursor.endEditBlock()


class _RichTextEditInternal(QTextEdit):
    """
    QTextEdit that handles:
      - Tab/Shift+Tab for table cell navigation and bullet indent/outdent
      - Paste interception: Word tables pasted via Ctrl+V are converted to
        proper QTextTable objects with header row styling instead of
        being dropped as plain text or broken HTML fragments
    """

    # ── format constants applied to all pasted content ───────────────────
    _PASTE_FAMILY = 'Aptos'
    _PASTE_SIZE   = 11.0
    _PASTE_COLOR  = (0x1a, 0x1a, 0x2e)   # near-black

    def _normalise_pasted(self, start_pos):
        """
        Walk every character from start_pos to the current cursor end
        and reset font to Aptos 11pt black, preserving bold/italic/bullets.
        """
        doc   = self.document()
        end   = self.textCursor().position()
        if end <= start_pos:
            return
        r, g, b = self._PASTE_COLOR
        body_fmt = QTextCharFormat()
        body_fmt.setFontFamily(self._PASTE_FAMILY)
        body_fmt.setFontPointSize(self._PASTE_SIZE)
        body_fmt.setForeground(QColor(r, g, b))
        # Walk block by block — faster than char-by-char and safe
        block = doc.findBlock(start_pos)
        while block.isValid() and block.position() <= end:
            it = block.begin()
            while not it.atEnd():
                frag = it.fragment()
                if frag.isValid():
                    frag_start = frag.position()
                    frag_end   = frag_start + frag.length()
                    # Only touch fragments within our pasted range
                    if frag_end > start_pos and frag_start < end:
                        old_fmt = frag.charFormat()
                        new_fmt = QTextCharFormat(old_fmt)
                        new_fmt.setFontFamily(self._PASTE_FAMILY)
                        new_fmt.setFontPointSize(self._PASTE_SIZE)
                        new_fmt.setForeground(QColor(r, g, b))
                        # Apply
                        c = QTextCursor(doc)
                        c.setPosition(max(start_pos, frag_start))
                        c.setPosition(min(end, frag_end),
                                      QTextCursor.MoveMode.KeepAnchor)
                        c.mergeCharFormat(new_fmt)
                it += 1
            block = block.next()

    def canInsertFromMimeData(self, source):
        if source.hasHtml() and '<table' in source.html().lower():
            return True
        return super().canInsertFromMimeData(source)

    def insertFromMimeData(self, source):
        start = self.textCursor().position()
        if source.hasHtml():
            html = source.html()
            if '<table' in html.lower():
                cursor = self.textCursor()
                if _parse_html_table_to_qt(html, cursor, self):
                    self._normalise_pasted(start)
                    return   # successfully inserted as QTextTable
        # Fall through to default paste for everything else
        super().insertFromMimeData(source)
        self._normalise_pasted(start)

    def keyPressEvent(self, event):
        key  = event.key()
        mods = event.modifiers()
        shift = bool(mods & Qt.KeyboardModifier.ShiftModifier)

        # On Windows, Shift+Tab fires as Key_Backtab rather than Key_Tab+Shift
        if key == Qt.Key.Key_Backtab:
            cursor = self.textCursor()
            tbl = cursor.currentTable()
            if tbl is not None:
                cur_cell = tbl.cellAt(cursor)
                r, c = cur_cell.row(), cur_cell.column()
                prev_cell = (tbl.cellAt(r, c - 1) if c > 0
                             else tbl.cellAt(r - 1, tbl.columns() - 1) if r > 0
                             else None)
                if prev_cell and prev_cell.isValid():
                    self.setTextCursor(prev_cell.firstCursorPosition())
                return
            lst = cursor.currentList()
            if lst is not None:
                _safe_outdent(cursor, lst, lst.format(), self)
                return
            return   # swallow Shift+Tab in plain text (no-op)

        if key == Qt.Key.Key_Tab:
            cursor = self.textCursor()
            tbl = cursor.currentTable()
            if tbl is not None:
                cur_cell = tbl.cellAt(cursor)
                r, c = cur_cell.row(), cur_cell.column()
                next_cell = (tbl.cellAt(r, c + 1) if c < tbl.columns() - 1
                             else tbl.cellAt(r + 1, 0) if r < tbl.rows() - 1
                             else None)
                if next_cell and next_cell.isValid():
                    self.setTextCursor(next_cell.firstCursorPosition())
                return

            lst = cursor.currentList()
            if lst is not None:
                _safe_indent(cursor, self)
                return

            cursor.insertText("    ")
            return

        # Header mode Enter: reset mode and clear header format from new line
        if key in (Qt.Key.Key_Return, Qt.Key.Key_Enter):
            if getattr(self, '_header_mode', False):
                self._header_mode = False
                if hasattr(self, '_rich_editor_ref'):
                    self._rich_editor_ref._update_header_btn_state()
                super().keyPressEvent(event)
                # Build body format
                body_fmt = QTextCharFormat()
                body_fmt.setFontFamily('Aptos')
                body_fmt.setFontPointSize(11.0)
                body_fmt.setFontWeight(QFont.Weight.Normal)
                body_fmt.setForeground(QColor(0x1a, 0x1a, 0x2e))
                # Apply to the new block three ways to ensure nothing leaks through:
                # 1. setCurrentCharFormat — affects next typed char
                self.setCurrentCharFormat(body_fmt)
                # 2. setBlockCharFormat — sets the block's default
                c = self.textCursor()
                c.setBlockCharFormat(body_fmt)
                # 3. mergeCharFormat on a selection covering the whole block
                c.select(QTextCursor.SelectionType.BlockUnderCursor)
                c.mergeCharFormat(body_fmt)
                c.clearSelection()
                self.setTextCursor(c)
                return

        if key in (Qt.Key.Key_Return, Qt.Key.Key_Enter):
            cursor = self.textCursor()
            lst = cursor.currentList()
            if lst is not None:
                try:
                    if not cursor.block().text().strip():
                        cursor.beginEditBlock()
                        bf = QTextBlockFormat()
                        bf.setIndent(0)
                        cursor.setBlockFormat(bf)
                        live_lst = cursor.block().textList()
                        if live_lst is not None:
                            live_lst.remove(cursor.block())
                        cursor.endEditBlock()
                        self.setTextCursor(cursor)
                    else:
                        saved_fmt = QTextListFormat(lst.format())
                        cursor.insertBlock()
                        if cursor.block().textList() is None:
                            cursor.createList(saved_fmt)
                        self.setTextCursor(cursor)
                except RuntimeError:
                    super().keyPressEvent(event)
                return

        # If header mode is on, apply header format to typed characters
        if getattr(self, '_header_mode', False):
            if key in (Qt.Key.Key_Return, Qt.Key.Key_Enter):
                pass  # handled above
            elif key not in (Qt.Key.Key_Backspace, Qt.Key.Key_Delete,
                             Qt.Key.Key_Left, Qt.Key.Key_Right,
                             Qt.Key.Key_Up, Qt.Key.Key_Down,
                             Qt.Key.Key_Home, Qt.Key.Key_End):
                # Let Qt insert the character first, then apply format
                super().keyPressEvent(event)
                c = self.textCursor()
                c.movePosition(QTextCursor.MoveOperation.Left,
                               QTextCursor.MoveMode.KeepAnchor)
                fmt = QTextCharFormat()
                from PyQt6.QtGui import QFont as _QF
                fmt.setFontFamily('Aptos')
                fmt.setFontPointSize(16.0)
                fmt.setFontWeight(_QF.Weight.Bold)
                from PyQt6.QtGui import QColor as _QC
                fmt.setForeground(_QC(0x9E, 0x1B, 0x32))
                c.setCharFormat(fmt)
                c.movePosition(QTextCursor.MoveOperation.Right)
                self.setTextCursor(c)
                self.setCurrentCharFormat(fmt)
                return
        super().keyPressEvent(event)


    def mouseDoubleClickEvent(self, event):
        cursor = self.cursorForPosition(event.pos())
        cf = cursor.charFormat()
        if cf.isImageFormat():
            img_fmt = cf.toImageFormat()
            from PyQt6.QtWidgets import (
                QDialog, QFormLayout, QSpinBox, QDialogButtonBox)
            dlg = QDialog(self)
            dlg.setWindowTitle("Resize Image")
            dlg.setFixedWidth(260)
            dlg.setStyleSheet(
                "QDialog{background:#fff;}"
                "QLabel{color:#1a0509;}"
                "QSpinBox{background:#fff;color:#1a0509;"
                "  border:1px solid #d6c0c5;border-radius:4px;padding:4px;}"
                "QPushButton{background:#f9f0f2;color:#920d2e;"
                "  border:1px solid #d6c0c5;border-radius:4px;padding:5px 14px;}")
            fl = QFormLayout(dlg)
            w_spin = QSpinBox()
            w_spin.setRange(10, 2000)
            w_spin.setValue(int(img_fmt.width()) or 400)
            h_spin = QSpinBox()
            h_spin.setRange(10, 2000)
            h_spin.setValue(int(img_fmt.height()) or 300)
            fl.addRow("Width (px):",  w_spin)
            fl.addRow("Height (px):", h_spin)
            btns = QDialogButtonBox(
                QDialogButtonBox.StandardButton.Ok |
                QDialogButtonBox.StandardButton.Cancel)
            btns.accepted.connect(dlg.accept)
            btns.rejected.connect(dlg.reject)
            fl.addRow(btns)
            if dlg.exec() == QDialog.DialogCode.Accepted:
                new_fmt = cf.toImageFormat()
                new_fmt.setWidth(w_spin.value())
                new_fmt.setHeight(h_spin.value())
                cursor.setCharFormat(new_fmt)
                self.setTextCursor(cursor)
            return
        super().mouseDoubleClickEvent(event)


class RichTextEditor(QWidget):
    """Toolbar + rich QTextEdit: Bold, Italic, Bullets, Indent/Outdent, Insert Table."""

    def __init__(self, parent=None):
        super().__init__(parent)
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(2)

        # Toolbar
        toolbar = QWidget()
        toolbar.setStyleSheet(
            "QWidget { background:#f4f6fa; border:1px solid #dde1e7; "
            "border-bottom:none; border-radius:4px 4px 0 0; }")
        tb_layout = QHBoxLayout(toolbar)
        tb_layout.setContentsMargins(6, 4, 6, 4)
        tb_layout.setSpacing(2)

        def _btn(text, tip):
            b = QPushButton(text); b.setToolTip(tip); b.setFixedSize(28, 24)
            b.setStyleSheet(
                "QPushButton{background:transparent;border:1px solid transparent;"
                "border-radius:3px;font-size:12px;color:#3a3a5c;}"
                "QPushButton:hover{background:#e0e4ef;border-color:#c0c8d8;}"
                "QPushButton:pressed{background:#d0d8ef;}")
            b.setFocusPolicy(Qt.FocusPolicy.NoFocus)
            return b

        self.btn_bold    = _btn("B",   "Bold (Ctrl+B)")
        self.btn_italic  = _btn("I",   "Italic (Ctrl+I)")
        self.btn_header  = _btn("H",   "Heading — 16pt Aptos bold dark red")
        self.btn_bullet  = _btn("•≡",  "Toggle bullet list")
        self.btn_indent  = _btn("→≡",  "Indent (Tab)")
        self.btn_outdent = _btn("←≡",  "Outdent (Shift+Tab)")
        self.btn_indent.setVisible(False)
        self.btn_outdent.setVisible(False)
        self.btn_table   = _btn("⊞",   "Insert table")
        self.btn_image   = _btn("🖼",   "Insert image")
        self.btn_image.setFixedSize(32, 24)
        self.btn_align_l = _btn("⬤≡",  "Align left")
        self.btn_align_c = _btn("≡⬤≡", "Align center")
        self.btn_align_r = _btn("≡⬤",  "Align right")

        _active = ("QPushButton:checked{background:#f5d0da;"
                   "border-color:#920d2e;color:#920d2e;}")
        self.btn_bold.setStyleSheet(
            self.btn_bold.styleSheet() + "QPushButton{font-weight:700;}" + _active)
        self.btn_italic.setStyleSheet(
            self.btn_italic.styleSheet() + "QPushButton{font-style:italic;}" + _active)
        self.btn_header.setStyleSheet(
            self.btn_header.styleSheet() +
            "QPushButton{font-weight:700;color:#9E1B32;font-size:13px;}")

        sep2 = QFrame(); sep2.setFrameShape(QFrame.Shape.VLine)
        sep2.setStyleSheet("color:#dde1e7;")
        sep3 = QFrame(); sep3.setFrameShape(QFrame.Shape.VLine)
        sep3.setStyleSheet("color:#dde1e7;")

        for w in (self.btn_bold, self.btn_italic, self.btn_header,
                  self.btn_bullet,
                  sep2, self.btn_table,
                  sep3,
                  self.btn_align_l, self.btn_align_c, self.btn_align_r):
            tb_layout.addWidget(w)
        tb_layout.addStretch()
        layout.addWidget(toolbar)

        self.editor = _RichTextEditInternal()
        self.editor._header_mode = False
        self.editor._rich_editor_ref = self   # back-reference for btn state
        # Default character format: Aptos 11pt black
        _def_fmt = QTextCharFormat()
        _def_fmt.setFontFamily('Aptos')
        _def_fmt.setFontPointSize(11.0)
        _def_fmt.setForeground(QColor(0x1a, 0x1a, 0x2e))
        self.editor.setCurrentCharFormat(_def_fmt)
        self.editor.document().setDefaultFont(
            QFont('Aptos', 11))
        # Match bullet visual indent to Word's convention (~0.25 in per level).
        self.editor.document().setIndentWidth(20)
        self.editor.setStyleSheet(
            "QTextEdit { border:1px solid #dde1e7; border-radius:0 0 4px 4px; "
            "padding:6px; font-family:'Aptos'; font-size:12px; "
            "color:#1a1a2e; background:#fafbfd; }"
            "QTextEdit:focus { border-color:#6a8fd8; background:#ffffff; }")
        self.editor.setMinimumHeight(120)
        self.editor.setAcceptRichText(True)
        layout.addWidget(self.editor)

        self.btn_bold.setCheckable(True)
        self.btn_italic.setCheckable(True)
        self.btn_bold.clicked.connect(self._toggle_bold)
        self.btn_italic.clicked.connect(self._toggle_italic)
        self.btn_header.clicked.connect(self._toggle_header)
        self.btn_bullet.clicked.connect(self._toggle_bullet)
        self.btn_indent.clicked.connect(self._indent_bullet)
        self.btn_outdent.clicked.connect(self._outdent_bullet)
        self.btn_table.clicked.connect(self._insert_table_dialog)
        self.btn_image.clicked.connect(self._insert_image)
        self.btn_align_l.clicked.connect(
            lambda: self._set_alignment(Qt.AlignmentFlag.AlignLeft))
        self.btn_align_c.clicked.connect(
            lambda: self._set_alignment(Qt.AlignmentFlag.AlignHCenter))
        self.btn_align_r.clicked.connect(
            lambda: self._set_alignment(Qt.AlignmentFlag.AlignRight))
        self.editor.currentCharFormatChanged.connect(self._sync_fmt_btns)

    def _sync_fmt_btns(self, fmt):
        self.btn_bold.setChecked(fmt.fontWeight() == QFont.Weight.Bold)
        self.btn_italic.setChecked(fmt.fontItalic())

    def _toggle_bold(self):
        fmt = QTextCharFormat()
        c = self.editor.textCursor()
        fmt.setFontWeight(
            QFont.Weight.Normal if c.charFormat().fontWeight() == QFont.Weight.Bold
            else QFont.Weight.Bold)
        c.mergeCharFormat(fmt); self.editor.setTextCursor(c)

    def _toggle_italic(self):
        fmt = QTextCharFormat()
        c = self.editor.textCursor()
        fmt.setFontItalic(not c.charFormat().fontItalic())
        c.mergeCharFormat(fmt)
        self.editor.setTextCursor(c)

    _HDR_FAMILY  = 'Aptos'
    _HDR_SIZE    = 16
    _HDR_COLOR   = (0x9E, 0x1B, 0x32)
    _BODY_COLOR  = (0x1a, 0x1a, 0x2e)
    _BODY_FAMILY = 'Aptos'
    _BODY_SIZE   = 11.0

    def _toggle_header(self):
        """Toggle sub-header style on selected text (or current line).
        Uses Qt selection-based font API for reliable cross-platform rendering.
        """
        c = self.editor.textCursor()
        # If no selection, select the whole current block
        if not c.hasSelection():
            c.select(QTextCursor.SelectionType.BlockUnderCursor)
        if not c.hasSelection():
            # Empty line: toggle header_mode for next typed chars
            self.editor._header_mode = not getattr(
                self.editor, '_header_mode', False)
            self._update_header_btn_state()
            if self.editor._header_mode:
                hdr_fmt = QTextCharFormat()
                hdr_fmt.setFontFamily(self._HDR_FAMILY)
                hdr_fmt.setFontPointSize(float(self._HDR_SIZE))
                hdr_fmt.setFontWeight(QFont.Weight.Bold)
                hdr_fmt.setForeground(QColor(*self._HDR_COLOR))
                self.editor.setCurrentCharFormat(hdr_fmt)
            return

        # Detect whether selection is already header
        probe = QTextCursor(self.editor.document())
        probe.setPosition(c.selectionStart())
        probe.setPosition(c.selectionStart() + 1,
                          QTextCursor.MoveMode.KeepAnchor)
        cf = probe.charFormat()
        r, g, b = self._HDR_COLOR
        is_hdr = (cf.foreground().color() == QColor(r, g, b)
                  and cf.fontWeight() == QFont.Weight.Bold)

        # Build target format
        fmt = QTextCharFormat()
        if is_hdr:
            fmt.setFontFamily(self._BODY_FAMILY)
            fmt.setFontPointSize(self._BODY_SIZE)
            fmt.setFontWeight(QFont.Weight.Normal)
            fmt.setForeground(QColor(*self._BODY_COLOR))
        else:
            fmt.setFontFamily(self._HDR_FAMILY)
            fmt.setFontPointSize(float(self._HDR_SIZE))
            fmt.setFontWeight(QFont.Weight.Bold)
            fmt.setForeground(QColor(r, g, b))

        # Apply using Qt's built-in selection API — most reliable method
        saved = QTextCursor(c)
        self.editor.setTextCursor(c)
        self.editor.setFontFamily(fmt.fontFamily())
        self.editor.setFontPointSize(fmt.fontPointSize())
        self.editor.setFontWeight(
            QFont.Weight.Bold if fmt.fontWeight() == QFont.Weight.Bold
            else QFont.Weight.Normal)
        # Apply colour via mergeCharFormat (colour not in setFont* API)
        colour_fmt = QTextCharFormat()
        colour_fmt.setForeground(fmt.foreground())
        c2 = self.editor.textCursor()
        c2.mergeCharFormat(colour_fmt)
        self.editor.setTextCursor(saved)
        self.editor.setCurrentCharFormat(fmt)

        self.editor._header_mode = not is_hdr
        self._update_header_btn_state()
    def _update_header_btn_state(self):
        active = getattr(self.editor, '_header_mode', False)
        style = (
            "QPushButton{font-weight:700;color:#ffffff;background:#9E1B32;"
            "border:1px solid #9E1B32;border-radius:3px;font-size:13px;font-family:Aptos;}"
            if active else
            "QPushButton{font-weight:700;color:#9E1B32;font-size:13px;}"
        )
        self.btn_header.setStyleSheet(
            "QPushButton{background:transparent;border:1px solid transparent;"
            "border-radius:3px;font-size:12px;color:#3a3a5c;}"
            "QPushButton:hover{background:#e0e4ef;border-color:#c0c8d8;}"
            "QPushButton:pressed{background:#d0d8ef;}" + style)

    def _toggle_bullet(self):
        cursor = self.editor.textCursor()
        lst = cursor.currentList()
        if lst:
            # Re-fetch list after setBlockFormat — the old reference may be stale
            bf = QTextBlockFormat()
            bf.setIndent(0)
            cursor.setBlockFormat(bf)
            live = cursor.block().textList()
            if live is not None:
                try:
                    live.remove(cursor.block())
                except RuntimeError:
                    pass
        else:
            # Clear placeholder so the bullet cursor is visible
            self.editor.setPlaceholderText("")
            fmt = QTextListFormat()
            fmt.setStyle(QTextListFormat.Style.ListDisc)
            fmt.setIndent(1)
            cursor.createList(fmt)
        self.editor.setTextCursor(cursor)

    def _indent_bullet(self):
        cursor = self.editor.textCursor()
        if cursor.currentList():
            _safe_indent(cursor, self)
        else:
            cursor.insertText("  ")

    def _outdent_bullet(self):
        cursor = self.editor.textCursor()
        lst = cursor.currentList()
        if not lst: return
        _safe_outdent(cursor, lst, lst.format(), self.editor)

    def _insert_table_dialog(self):
        dlg = InsertTableDialog(self)
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return
        rows, cols, has_total, sum_cols = dlg.get_config()
        # Add an extra row for the total if requested
        total_rows = rows + (1 if has_total else 0)
        cursor = self.editor.textCursor()
        from PyQt6.QtGui import QTextLength
        fmt = QTextTableFormat()
        fmt.setCellPadding(4)
        fmt.setCellSpacing(0)
        fmt.setBorderStyle(QTextTableFormat.BorderStyle.BorderStyle_Solid)
        fmt.setBorder(1)
        fmt.setWidth(QTextLength(QTextLength.Type.PercentageLength, 100))
        table = cursor.insertTable(total_rows, cols, fmt)
        # Style first row as header
        _style_qt_table_header_row(table, cols)
        # Style total row if requested
        if has_total:
            total_row_idx = total_rows - 1
            self._style_total_row(table, total_row_idx, cols, sum_cols, rows)
            # Store metadata so the doc converter can recalculate totals
            # We write a hint into the first total cell as a marker
            cell0 = table.cellAt(total_row_idx, 0)
            c = cell0.firstCursorPosition()
            fmt2 = QTextCharFormat()
            fmt2.setFontWeight(QFont.Weight.Bold)
            fmt2.setForeground(QColor("#1a1a2e"))
            c.mergeCharFormat(fmt2)
            c.insertText("Total")
            # Hint columns to sum in last cell property (as comment-like text)
            # Store sum_cols in a hidden way: set tool tip on the cell
            # Actually just pre-fill "=SUM" hint so user knows which cells to fill
            for col_i in sum_cols:
                cell = table.cellAt(total_row_idx, col_i)
                sc = cell.firstCursorPosition()
                sc.insertText("")

    def _style_total_row(self, table, row_idx, cols, sum_cols, data_rows):
        """Style the total row: bold dark bg, mark sum columns."""
        from PyQt6.QtGui import QTextCharFormat, QTextBlockFormat
        total_bg = QColor("#dde1e7")
        for col_i in range(cols):
            cell = table.cellAt(row_idx, col_i)
            # Gray background on total row
            fmt = cell.format()
            fmt.setBackground(total_bg)
            cell.setFormat(fmt)
            # Mark sum columns with a different shade
            if col_i in sum_cols:
                fmt2 = cell.format()
                fmt2.setBackground(QColor("#c8cedd"))
                cell.setFormat(fmt2)
            # Bold text format
            c = cell.firstCursorPosition()
            cf = QTextCharFormat()
            cf.setFontWeight(QFont.Weight.Bold)
            cf.setForeground(QColor("#1a1a2e"))
            c.mergeCharFormat(cf)

    def _insert_image(self):
        from PyQt6.QtWidgets import QFileDialog
        path, _ = QFileDialog.getOpenFileName(
            self, "Insert Image", "",
            "Images (*.png *.jpg *.jpeg *.bmp *.gif *.webp)")
        if not path:
            return
        cursor = self.editor.textCursor()
        from PyQt6.QtGui import QImage, QTextImageFormat
        img = QImage(path)
        if img.isNull():
            return
        # Scale to max 400px wide, preserve aspect
        max_w = 400
        w = min(img.width(), max_w)
        h = int(img.height() * w / img.width()) if img.width() else 80
        # Store image in document resource
        self.editor.document().addResource(
            3, __import__('PyQt6.QtCore', fromlist=['QUrl']).QUrl.fromLocalFile(path),
            img)
        fmt = QTextImageFormat()
        fmt.setName(path)
        fmt.setWidth(w)
        fmt.setHeight(h)
        cursor.insertImage(fmt)
        self.editor.setTextCursor(cursor)

    def _set_alignment(self, alignment):
        cursor = self.editor.textCursor()
        bf = cursor.blockFormat()
        bf.setAlignment(alignment)
        cursor.setBlockFormat(bf)
        self.editor.setTextCursor(cursor)

    def toPlainText(self): return self.editor.toPlainText()
    def toHtml(self):      return self.editor.toHtml()
    def setPlaceholderText(self, t): self.editor.setPlaceholderText(t)

# ── HTML → python-docx converter ──────────────────────────────────────────────
def _html_to_docx_paragraphs(doc, html: str):
    """
    Convert Qt-generated rich-text HTML into python-docx content.

    Qt's QTextEdit HTML quirks this handles:
      - Bullet lists: <ul style="..."><li style="-qt-list-indent:N"><p>text</p></li></ul>
        AND the flattened form where Qt emits <p style="-qt-list-indent:N ...">
      - Tables: <table><thead/tbody><tr><td><p>text</p></td></tr></table>
        but also nested inside a <p style="-qt-paragraph-type:empty"> wrapper
      - Bold/italic via inline style font-weight/font-style as well as <b>/<i> tags
      - Indented bullets: -qt-list-indent value maps to Word indent level
    """
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # ── pull body HTML ────────────────────────────────────────────────────────
    body_m = re.search(r'<body[^>]*>(.*?)</body>', html, re.S | re.I)
    body_html = body_m.group(1) if body_m else html

    # ── minimal DOM builder ───────────────────────────────────────────────────
    class _N:
        __slots__ = ('tag','attrs','children','text')
        def __init__(self, tag, attrs=None):
            self.tag      = (tag or '').lower()
            self.attrs    = attrs or {}
            self.children = []
            self.text     = ''
        def attr(self, k, default=''):
            return self.attrs.get(k, default)
        def style(self):
            return self.attr('style')
        def has_style(self, fragment):
            return fragment.lower() in self.style().lower()

    class _Builder(HTMLParser):
        def __init__(self):
            super().__init__()
            self.root = _N('root')
            self._stack = [self.root]
        def handle_starttag(self, tag, attrs):
            n = _N(tag, dict(attrs))
            self._stack[-1].children.append(n)
            self._stack.append(n)
        def handle_endtag(self, _):
            if len(self._stack) > 1: self._stack.pop()
        def handle_data(self, data):
            self._stack[-1].text += data
        def handle_entityref(self, name):
            mapping = {'amp':'&','lt':'<','gt':'>','quot':'"','nbsp':' ','apos':"'"}
            self._stack[-1].text += mapping.get(name, '')
        def handle_charref(self, name):
            try:
                c = chr(int(name[1:], 16) if name.startswith('x') else int(name))
                self._stack[-1].text += c
            except Exception:
                pass

    b = _Builder()
    b.feed(body_html)

    # ── run extraction ────────────────────────────────────────────────────────
    def _bold(node):
        if node.tag in ('b', 'strong'): return True
        s = node.style()
        return ('font-weight:600' in s or 'font-weight:700' in s or
                'font-weight: 600' in s or 'font-weight: 700' in s or
                'font-weight:bold' in s.lower())

    def _italic(node):
        if node.tag in ('i', 'em'): return True
        s = node.style()
        return 'font-style:italic' in s or 'font-style: italic' in s

    # 'p' and 'li' removed so their inline text IS collected as run content
    SKIP_TEXT_TAGS = frozenset(
        ('td','th','tr','table','thead','tbody','tfoot',
         'ul','ol','body','html','root','head','style','script'))

    def _font_size(node):
        """Return pt size from inline style font-size, or None."""
        m = re.search(r'font-size\s*:\s*([\d.]+)pt', node.style())
        return float(m.group(1)) if m else None

    def _font_family(node):
        """Return font-family value from inline style, or None."""
        m = re.search(r"font-family\s*:\s*'?([^;',]+)", node.style())
        return m.group(1).strip().strip("'\"") if m else None

    def _font_color(node):
        """Return hex color string (no #) from color: #RRGGBB in style, or None."""
        m = re.search(r'(?:^|;)\s*color\s*:\s*#([0-9a-fA-F]{6})', node.style())
        return m.group(1) if m else None

    def _runs(node, bold=False, italic=False, size=None, family=None, color=None):
        b2      = bold   or _bold(node)
        i2      = italic or _italic(node)
        size2   = _font_size(node)   or size
        family2 = _font_family(node) or family
        color2  = _font_color(node)  or color
        if node.text and (node.text.strip() or node.tag not in SKIP_TEXT_TAGS):
            yield (node.text, b2, i2, size2, family2, color2)
        for child in node.children:
            yield from _runs(child, b2, i2, size2, family2, color2)

    def _fill_para(p, node):
        from docx.shared import Pt, RGBColor as _RGB
        for text, b, i, size, family, color in _runs(node):
            run = p.add_run(text)
            run.bold = b; run.italic = i
            if size:   run.font.size   = Pt(size)
            if family: run.font.name   = family
            if color:
                try:
                    r = int(color[0:2], 16)
                    g = int(color[2:4], 16)
                    b_ = int(color[4:6], 16)
                    run.font.color.rgb = _RGB(r, g, b_)
                except Exception:
                    pass

    # ── list helpers ──────────────────────────────────────────────────────────
    def _list_indent(node):
        """Return Qt list indent level (1-based) from -qt-list-indent style."""
        m = re.search(r'-qt-list-indent\s*:\s*(\d+)', node.style())
        return int(m.group(1)) if m else 1

    def _is_bullet(node):
        s = node.style()
        return ('-qt-list-indent' in s or
                'list-style-type' in s.lower() or
                node.tag == 'li')

    def _list_style_type(node):
        """Detect ordered vs unordered from parent ul/ol or style."""
        s = node.style().lower()
        if 'decimal' in s or 'lower-alpha' in s or 'upper-alpha' in s:
            return 'ordered'
        return 'bullet'

    def _ensure_bullet_numid_fw(doc, ordered=False):
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _E

        key = (id(doc), ordered)
        if key in _fw_bullet_numid_cache:
            return _fw_bullet_numid_cache[key]

        try:
            numbering_part = doc.part.numbering_part
        except Exception:
            from docx.parts.numbering import NumberingPart
            numbering_part = NumberingPart.new()
            doc.part._add_relationship(
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering',
                numbering_part)

        np_el = numbering_part._element
        abs_id = '210' if not ordered else '211'
        num_id = '210' if not ordered else '211'

        for old in np_el.findall(_qn('w:abstractNum')):
            if old.get(_qn('w:abstractNumId')) == abs_id:
                np_el.remove(old)
        for old in np_el.findall(_qn('w:num')):
            if old.get(_qn('w:numId')) == num_id:
                np_el.remove(old)

        absNum = _E('w:abstractNum')
        absNum.set(_qn('w:abstractNumId'), abs_id)
        for lvl_idx in range(9):
            lvl = _E('w:lvl')
            lvl.set(_qn('w:ilvl'), str(lvl_idx))
            start = _E('w:start'); start.set(_qn('w:val'), '1')
            lvl.append(start)
            fmt_el = _E('w:numFmt')
            fmt_el.set(_qn('w:val'), 'decimal' if ordered else 'bullet')
            lvl.append(fmt_el)
            txt = _E('w:lvlText')
            txt.set(_qn('w:val'), ('%' + '1.') if ordered else '\u2022')
            lvl.append(txt)
            jc = _E('w:jc'); jc.set(_qn('w:val'), 'left')
            lvl.append(jc)
            pPr_lvl = _E('w:pPr')
            ind = _E('w:ind')
            ind.set(_qn('w:left'),    str(360 * (lvl_idx + 1)))
            ind.set(_qn('w:hanging'), '360')
            pPr_lvl.append(ind); lvl.append(pPr_lvl)
            absNum.append(lvl)
        np_el.append(absNum)

        num = _E('w:num')
        num.set(_qn('w:numId'), num_id)
        absRef = _E('w:abstractNumId'); absRef.set(_qn('w:val'), abs_id)
        num.append(absRef)
        np_el.append(num)

        _fw_bullet_numid_cache[key] = num_id
        return num_id

    def _add_bullet_para(doc, node, indent_level=1, ordered=False):
        """
        Add a bullet paragraph using self-contained numPr XML so every
        indent level reliably renders a bullet/number in the generated document.
        """
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _E

        p   = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()

        num_id = _ensure_bullet_numid_fw(doc, ordered)

        numPr = _E('w:numPr')
        ilvl  = _E('w:ilvl'); ilvl.set(_qn('w:val'), str(max(0, indent_level - 1)))
        numId = _E('w:numId'); numId.set(_qn('w:val'), num_id)
        numPr.append(ilvl); numPr.append(numId)
        pPr.append(numPr)

        spacing = _E('w:spacing')
        spacing.set(_qn('w:before'), '0')
        spacing.set(_qn('w:after'),  '0')
        pPr.append(spacing)

        left = 360 * indent_level
        ind = _E('w:ind')
        ind.set(_qn('w:left'),    str(left))
        ind.set(_qn('w:hanging'), '360')
        pPr.append(ind)

        _fill_para(p, node)
        return p

    # ── table builder ─────────────────────────────────────────────────────────
    def _build_table(doc, node):
        """Build a python-docx table from a <table> node."""
        # Collect all <tr> nodes through any tbody/thead/tfoot wrappers
        sections = [c for c in node.children
                    if c.tag in ('tbody','thead','tfoot')]
        if not sections:
            sections = [node]
        all_trs = []
        for sec in sections:
            for ch in sec.children:
                if ch.tag == 'tr':
                    all_trs.append(ch)
        if not all_trs:
            return

        max_cols = max(
            sum(1 for c in tr.children if c.tag in ('td','th'))
            for tr in all_trs
        ) if all_trs else 1

        tbl = doc.add_table(rows=len(all_trs), cols=max(max_cols, 1))
        tbl.style = 'Table Grid'

        from docx.oxml import OxmlElement as _E
        from docx.oxml.ns import qn as _qn

        for r_i, tr in enumerate(all_trs):
            cell_nodes = [c for c in tr.children if c.tag in ('td','th')]
            # First row is always the header — regardless of td vs th
            is_header_row = (r_i == 0)

            for c_i, cn in enumerate(cell_nodes[:max_cols]):
                cell = tbl.rows[r_i].cells[c_i]
                cell.text = ''
                p = cell.paragraphs[0]

                from docx.shared import Pt as _Pt, RGBColor as _TRGB
                for text, bold, italic, size, family, color in _runs(cn):
                    run = p.add_run(text)
                    run.bold   = bold or is_header_row
                    run.italic = italic
                    if size:   run.font.size = _Pt(size)
                    if family: run.font.name = family
                    if color:
                        try:
                            run.font.color.rgb = _TRGB(
                                int(color[0:2],16), int(color[2:4],16), int(color[4:6],16))
                        except Exception: pass

                # Gray fill on header row
                if is_header_row:
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = _E('w:shd')
                    shd.set(_qn('w:val'),   'clear')
                    shd.set(_qn('w:color'), 'auto')
                    shd.set(_qn('w:fill'),  'D9D9D9')
                    tcPr.append(shd)

    # ── main walker ───────────────────────────────────────────────────────────
    def _walk(node):
        tag = node.tag

        # Transparent containers — recurse
        if tag in ('html','head','body','root','style','script','meta','title'):
            for child in node.children:
                _walk(child)
            return

        # Qt wraps everything in <p> — check if it actually contains a table
        # or is a list item before treating it as a plain paragraph
        if tag in ('div', 'p'):
            # Check for nested table
            for child in node.children:
                if child.tag == 'table':
                    # Any text before the table
                    if node.text and node.text.strip():
                        p = doc.add_paragraph()
                        p.add_run(node.text)
                    _build_table(doc, child)
                    return

            # Qt bullet list paragraph: <p style="-qt-list-indent:N ...">
            if _is_bullet(node):
                indent = _list_indent(node)
                ordered = _list_style_type(node) == 'ordered'
                _add_bullet_para(doc, node, indent_level=indent, ordered=ordered)
                return

            # Plain paragraph — collect all child content
            p = doc.add_paragraph()
            _fill_para(p, node)
            return

        if tag == 'span':
            # Span outside a paragraph — wrap in one
            p = doc.add_paragraph()
            _fill_para(p, node)
            return

        if tag in ('ul', 'ol'):
            ordered = (tag == 'ol')
            for child in node.children:
                if child.tag == 'li':
                    indent = _list_indent(child)
                    _add_bullet_para(doc, child, indent_level=indent, ordered=ordered)
                else:
                    _walk(child)
            return

        if tag == 'li':
            _add_bullet_para(doc, node, indent_level=_list_indent(node))
            return

        if tag == 'table':
            _build_table(doc, node)
            return

        if tag == 'br':
            doc.add_paragraph()
            return

        # Fallback: recurse into anything else
        for child in node.children:
            _walk(child)

    _walk(b.root)

# ── Margin configuration widget ───────────────────────────────────────────────
class MarginWidget(QWidget):
    """
    Simple margin control: Margin:  [−]  [20.0%]  [+]
    Always active. Calls _refresh_cb whenever the value changes.
    """
    def __init__(self, global_margin_ref=None, parent=None):
        super().__init__(parent)
        self._value      = 20.0
        self._refresh_cb = None

        layout = QHBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        lbl = QLabel("Margin:")
        lbl.setStyleSheet("color:#3a3a5c; font-size:11px; font-weight:600;")
        layout.addWidget(lbl)
        layout.addSpacing(10)

        self._btn_minus = QPushButton("−")
        self._btn_minus.setFixedSize(30, 30)
        self._btn_minus.setStyleSheet(
            "QPushButton { background:#f4f6fa; color:#3a3a5c; border:1px solid #c8cedd;"
            "  border-radius:4px; font-size:16px; font-weight:700; }"
            "QPushButton:hover { background:#e0e4ef; }"
            "QPushButton:pressed { background:#d0d8ef; }")
        self._btn_minus.clicked.connect(self._decrement)
        layout.addWidget(self._btn_minus)

        layout.addSpacing(8)

        self._lbl = QLabel("20.0 %")
        self._lbl.setFixedWidth(74)
        self._lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self._lbl.setStyleSheet(
            "QLabel { background:#ffffff; border:1px solid #6a8fd8;"
            "  border-radius:4px; color:#1a1a2e; font-size:12px;"
            "  font-weight:600; padding:4px 6px; }")
        self._lbl.setToolTip("Margin %. Formula: Price = Cost / (1 - Margin/100)")
        layout.addWidget(self._lbl)

        layout.addSpacing(8)

        self._btn_plus = QPushButton("+")
        self._btn_plus.setFixedSize(30, 30)
        self._btn_plus.setStyleSheet(
            "QPushButton { background:#f4f6fa; color:#3a3a5c; border:1px solid #c8cedd;"
            "  border-radius:4px; font-size:16px; font-weight:700; }"
            "QPushButton:hover { background:#e0e4ef; }"
            "QPushButton:pressed { background:#d0d8ef; }")
        self._btn_plus.clicked.connect(self._increment)
        layout.addWidget(self._btn_plus)

        layout.addStretch()

        # Compatibility stubs — referenced elsewhere, must exist
        self._enable_cb     = QCheckBox()
        self._enable_cb.setChecked(True)
        self._enable_cb.setVisible(False)
        self._spin          = QDoubleSpinBox()
        self._spin.setVisible(False)
        self._spin.setRange(0, 99.9)
        self._spin.setDecimals(1)
        self._spin.setValue(self._value)
        self._use_global_cb = QCheckBox()
        self._use_global_cb.setChecked(True)   # MUST be True so global spin triggers refresh
        self._use_global_cb.setVisible(False)
        self._preview_cb    = QCheckBox()
        self._preview_cb.setChecked(True)
        self._preview_cb.setVisible(False)

    def set_refresh_callback(self, fn):
        self._refresh_cb = fn

    def _increment(self):
        self._value = min(99.0, round(self._value + 1.0, 1))
        self._lbl.setText(f"{self._value:.1f} %")
        if self._refresh_cb:
            self._refresh_cb()

    def _decrement(self):
        self._value = max(0.0, round(self._value - 1.0, 1))
        self._lbl.setText(f"{self._value:.1f} %")
        if self._refresh_cb:
            self._refresh_cb()

    def is_enabled(self):  return True
    def show_preview(self): return True
    def margin_pct(self):  return self._value
    def set_global_ref(self, fn): pass   # global ref not used; section always uses _value


# ── BOM TableSection ──────────────────────────────────────────────────────────
# Columns: Part Number | Qty | Cost | Margin % | Sale Price | [hidden Total]
# Column indices: 0=PartNum, 1=Qty, 2=CostEach, 3=MarkedUp(preview), 4=Total(hidden)

class TableSection(CollapsibleCard):
    """BOM table with optional margin preview columns."""

    _COL_EDIT      = 0   # ✏ button column
    _COL_PART      = 1
    _COL_QTY       = 2
    _COL_COST      = 3
    _COL_MARGIN    = 4   # editable per-row margin %
    _COL_MKUP      = 5   # sale price per unit
    _COL_LINE_TOT  = 6   # qty × sale price (visible)
    _COL_TOTAL     = 7   # hidden — used for doc export

    def __init__(self, on_remove=None, global_margin_ref=None):
        super().__init__("📋  BOM Table")
        self._on_remove = on_remove

        remove_btn = QPushButton("✕  Remove")
        remove_btn.setStyleSheet(REMOVE_BTN_STYLE)
        remove_btn.clicked.connect(self._delete_self)
        self._header_extras.addWidget(remove_btn)

        # Margin widget — kept as attribute for global ref, not shown in UI
        self._global_margin_ref = global_margin_ref  # callable returning global %
        self._global_spin = None
        self._margin_widget = MarginWidget(global_margin_ref=global_margin_ref)
        self._margin_widget.set_refresh_callback(self._refresh_all_totals)

        # Table (7 cols)
        self.table = AutoHeightTable(8)
        self.table.setHorizontalHeaderLabels(
            ["", "Part Number", "Qty", "Cost", "Margin %",
             "Sale Price", "Line Total", "Total"])
        self.table.setColumnHidden(self._COL_TOTAL, True)

        hh = self.table.horizontalHeader()
        hh.setSectionResizeMode(self._COL_EDIT,     QHeaderView.ResizeMode.Fixed)
        hh.setSectionResizeMode(self._COL_PART,     QHeaderView.ResizeMode.Stretch)
        hh.setSectionResizeMode(self._COL_QTY,      QHeaderView.ResizeMode.Fixed)
        hh.setSectionResizeMode(self._COL_COST,     QHeaderView.ResizeMode.Fixed)
        hh.setSectionResizeMode(self._COL_MARGIN,   QHeaderView.ResizeMode.Fixed)
        hh.setSectionResizeMode(self._COL_MKUP,     QHeaderView.ResizeMode.Fixed)
        hh.setSectionResizeMode(self._COL_LINE_TOT, QHeaderView.ResizeMode.Fixed)
        hh.setSectionResizeMode(self._COL_TOTAL,    QHeaderView.ResizeMode.Fixed)
        self.table.setColumnWidth(self._COL_EDIT,     44)
        self.table.setColumnWidth(self._COL_QTY,      75)
        self.table.setColumnWidth(self._COL_COST,     100)
        self.table.setColumnWidth(self._COL_MARGIN,   75)
        self.table.setColumnWidth(self._COL_MKUP,     100)
        self.table.setColumnWidth(self._COL_LINE_TOT, 100)
        self.table.verticalHeader().setVisible(False)

        self.table.setItemDelegateForColumn(self._COL_PART,   RichTextDelegate(self.table))
        self.table.setItemDelegateForColumn(self._COL_QTY,    NumericDelegate(self.table))
        self.table.setItemDelegateForColumn(self._COL_COST,   NumericDelegate(self.table))
        self.table.setItemDelegateForColumn(self._COL_MARGIN, NumericDelegate(self.table))

        # Note above the table — fixed height, reference-only notice
        note = QLabel(
            "ℹ  Cost and Margin % columns are for reference only "
            "and will not appear in the generated document.")
        note.setWordWrap(True)
        note.setFixedHeight(36)
        note.setStyleSheet(
            "QLabel { color:#6672a0; font-size:10px; "
            "background:#f4f6fa; border:1px solid #dde1e7; "
            "border-radius:4px; padding:4px 8px; }")
        self._body_layout.addWidget(note)

        # Table + Hours? column side-by-side
        table_row = QHBoxLayout()
        table_row.setSpacing(0)
        table_row.setContentsMargins(0, 0, 0, 0)
        table_row.addWidget(self.table)
        # Hours? panel — header label + QCheckBox per row, styled to blend with table
        self._hours_panel = QWidget()
        self._hours_panel.setFixedWidth(66)
        self._hours_panel.setStyleSheet(
            "QWidget { background:#ffffff; border:1px solid #dde1e7; "
            "border-left:none; border-radius:0 4px 4px 0; }")
        hours_vbox = QVBoxLayout(self._hours_panel)
        hours_vbox.setContentsMargins(0, 0, 0, 0)
        hours_vbox.setSpacing(0)
        # Header label matching table header exactly
        hours_hdr = QLabel("Hours?")
        hours_hdr.setAlignment(Qt.AlignmentFlag.AlignCenter)
        hours_hdr.setFixedHeight(30)
        hours_hdr.setStyleSheet(
            "QLabel { background:#f4f6fa; color:#3a3a5c; font-size:11px;"
            " font-weight:600; border-bottom:1px solid #dde1e7;"
            " border-top:none; border-left:none; border-right:none; padding:0 4px; }")
        hours_vbox.addWidget(hours_hdr)
        self._hours_hdr_label = hours_hdr
        # Checkboxes are added dynamically in add_row()
        self._hours_cb_layout = hours_vbox   # append QCheckBoxes here
        self._hours_checkboxes = []           # list of QCheckBox, one per row
        table_row.addWidget(self._hours_panel)
        self._table_row_layout = table_row

        self.table.cellChanged.connect(self._safe_update_totals)
        self._body_layout.addLayout(self._table_row_layout)

        # Grand total label
        self._grand_total_lbl = QLabel("Grand Total:  $0")
        self._grand_total_lbl.setAlignment(Qt.AlignmentFlag.AlignRight)
        self._grand_total_lbl.setStyleSheet(
            "QLabel { font-size:12px; font-weight:700; color:#1a1a2e; "
            "padding:4px 8px; }")
        self._body_layout.addWidget(self._grand_total_lbl)

        # Buttons
        btn_row = QHBoxLayout(); btn_row.setSpacing(8)
        add_btn = QPushButton("＋  Add Row"); add_btn.setStyleSheet(ACTION_BTN_STYLE)
        add_btn.clicked.connect(self.add_row)
        del_btn = QPushButton("－  Remove Selected"); del_btn.setStyleSheet(ACTION_BTN_STYLE)
        del_btn.clicked.connect(self._remove_selected_rows)
        btn_row.addWidget(add_btn); btn_row.addWidget(del_btn)
        btn_row.addStretch()
        self._body_layout.addLayout(btn_row)
        # Track hours state per row
        self._hours_rows = set()

        # Start with one blank row pre-filled
        self.add_row()

    # ── margin visibility ─────────────────────────────────────────────────────
    def showEvent(self, event):
        """Sync Hours? header height to the real table header after layout."""
        super().showEvent(event)
        hdr_h = self.table.horizontalHeader().height()
        if hdr_h > 0 and hasattr(self, '_hours_hdr_label'):
            self._hours_hdr_label.setFixedHeight(hdr_h)

    def _open_cell_editor(self, row):
        """Open CellEditorDialog for the Part Number cell of this row."""
        item = self.table.item(row, self._COL_PART)
        html  = (item.data(Qt.ItemDataRole.UserRole) or "") if item else ""
        plain = (item.text() or "") if item else ""
        dlg = CellEditorDialog(html=html, plain=plain, parent=self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            new_plain = dlg.plain().strip()
            new_html  = dlg.html()
            if item is None:
                item = QTableWidgetItem()
                item.setForeground(QColor("#1a1a2e"))
                self.table.setItem(row, self._COL_PART, item)
            self.table.blockSignals(True)
            item.setText(new_plain)
            self.table.blockSignals(False)
            item.setData(Qt.ItemDataRole.UserRole, new_html)  # outside blockSignals so dataChanged fires → delegate repaints
            # Resize row to fit rich content; sync the Hours? container too
            from PyQt6.QtWidgets import QStyleOptionViewItem
            delegate = self.table.itemDelegateForColumn(self._COL_PART)
            idx = self.table.model().index(row, self._COL_PART)
            opt = QStyleOptionViewItem()
            opt.rect = self.table.visualRect(idx)
            opt.font = self.table.font()
            sh = delegate.sizeHint(opt, idx)
            new_h = max(30, sh.height())
            self.table.setRowHeight(row, new_h)
            # Keep the Hours? checkbox container flush with the row
            if row < len(self._hours_checkboxes):
                cb_item = self._hours_cb_layout.itemAt(row + 1)
                if cb_item and cb_item.widget():
                    cb_item.widget().setFixedHeight(new_h)
            self.table.viewport().update()
            self.table.updateGeometry()
            self.updateGeometry()

    def _on_margin_toggle(self, state):
        self._refresh_all_totals()

    def _on_preview_toggle(self, state):
        pass  # Marked-Up column is always visible

    def _row_margin_pct(self, row: int) -> float:
        """Read margin % from the inline Margin % column for this row."""
        item = self.table.item(row, self._COL_MARGIN)
        if item and item.text().strip():
            try:
                return float(item.text())
            except ValueError:
                pass
        return self._margin_widget.margin_pct()

    # ── row management ────────────────────────────────────────────────────────
    def add_row(self):
        self.table.blockSignals(True)
        row = self.table.rowCount()
        self.table.insertRow(row)
        global_pct = self._margin_widget.margin_pct()
        for col in range(8):
            if col == self._COL_EDIT:
                # ✏ edit button — placed as a widget via setCellWidget
                item = QTableWidgetItem("")
                item.setFlags(Qt.ItemFlag.ItemIsEnabled)
            elif col == self._COL_MARGIN:
                item = QTableWidgetItem(
                    f"{global_pct:.1f}" if global_pct > 0 else "")
                item.setForeground(QColor("#6672a0"))
            elif col in (self._COL_MKUP, self._COL_LINE_TOT, self._COL_TOTAL):
                item = QTableWidgetItem("")
                item.setFlags(item.flags() & ~Qt.ItemFlag.ItemIsEditable)
                item.setForeground(QColor("#1a3a6e"))
                item.setBackground(QColor("#e8f0fe"))
            else:
                item = QTableWidgetItem("")
                item.setForeground(QColor("#1a1a2e"))
            self.table.setItem(row, col, item)
        self.table.blockSignals(False)
        # Edit button — pencil icon, transparent background, lights up on hover
        edit_btn = QPushButton()
        from PyQt6.QtGui import QIcon, QFont as _QF
        # Use a QLabel-based approach: set text to pencil unicode with symbol font
        edit_btn.setText("\u270f")  # ✏ pencil
        _icon_font = _QF("Segoe UI Symbol", 13)
        _icon_font.setStyleStrategy(_QF.StyleStrategy.PreferDefault)
        edit_btn.setFont(_icon_font)
        edit_btn.setFixedSize(26, 26)
        edit_btn.setToolTip("Edit item text (supports bullets & bold)")
        edit_btn.setStyleSheet(
            "QPushButton { background:transparent; border:none;"
            "  color:#b0b8cc; padding:0; }"
            "QPushButton:hover { color:#3a5bd9; }"
            "QPushButton:pressed { color:#1a3aaa; }")
        edit_btn.clicked.connect(lambda _, r=row: self._open_cell_editor(r))
        edit_container = QWidget()
        edit_container.setStyleSheet("background:transparent;")
        ecl = QHBoxLayout(edit_container)
        ecl.setContentsMargins(1, 0, 1, 0)
        ecl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        ecl.addWidget(edit_btn)
        self.table.setCellWidget(row, self._COL_EDIT, edit_container)
        # Add a QCheckBox for this row in the Hours? panel
        cb = QCheckBox()
        # Light blue-grey background so the white checkbox indicator is visible
        cb.setStyleSheet(
            "QCheckBox { background:#e8ecf5; border-radius:3px; padding:2px; }"
            "QCheckBox::indicator { width:14px; height:14px; }"
            "QCheckBox::indicator:unchecked { border:1px solid #9aa5c8; border-radius:2px; background:#ffffff; }"
            "QCheckBox::indicator:checked { border:1px solid #3a5bd9; border-radius:2px; background:#3a5bd9; }")
        row_idx = len(self._hours_checkboxes)
        cb.stateChanged.connect(
            lambda state, r=row_idx: self._on_hours_cb_changed(r, state))
        # Wrap in a centered container matching the row height
        container = QWidget()
        container.setStyleSheet("QWidget { background:#edf0f5; }")
        cl = QHBoxLayout(container)
        cl.setContentsMargins(0, 0, 0, 0)
        cl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        cl.addWidget(cb)
        row_h = self.table.rowHeight(row) if self.table.rowCount() > 0 else 30
        container.setFixedHeight(max(30, row_h))
        self._hours_cb_layout.addWidget(container)
        self._hours_checkboxes.append(cb)
        if not hasattr(self, '_last_global_pct'):
            self._last_global_pct = self._margin_widget.margin_pct()
        self.table.updateGeometry(); self.updateGeometry()

    def _remove_selected_rows(self):
        rows = sorted({i.row() for i in self.table.selectedIndexes()}, reverse=True)
        self.table.blockSignals(True)
        for row in rows:
            self.table.removeRow(row)
            self._hours_rows.discard(row)
            # Remove the checkbox container at this index
            if row < len(self._hours_checkboxes):
                self._hours_checkboxes.pop(row)
                item = self._hours_cb_layout.itemAt(row + 1)  # +1 for header
                if item and item.widget():
                    item.widget().deleteLater()
                self._hours_cb_layout.removeItem(item)
        # Rebuild _hours_rows from remaining checkboxes
        self._hours_rows = {
            i for i, cb in enumerate(self._hours_checkboxes)
            if cb.isChecked()
        }
        self.table.blockSignals(False)
        self.table.updateGeometry(); self.updateGeometry()

    def _on_hours_cb_changed(self, row, state):
        """Fired when a per-row Hours? checkbox is toggled."""
        is_hours = (state == Qt.CheckState.Checked.value)
        if is_hours:
            self._hours_rows.add(row)
        else:
            self._hours_rows.discard(row)
        self._safe_update_totals(row, self._COL_QTY)

    def _qty_for_row(self, row):
        """Return (numeric_qty, display_suffix) for the given row."""
        qty_item = self.table.item(row, self._COL_QTY)
        is_hours = row in getattr(self, '_hours_rows', set())
        try:
            raw = qty_item.text().strip() if qty_item else ''
            # Guard against a saved '(Hours)' suffix surviving in the cell text.
            raw = raw.replace('(Hours)', '').strip()
            qty = float(raw) if raw else 0.0
        except Exception:
            qty = 0.0
        return qty, (' (Hours)' if is_hours else '')

    def _safe_update_totals(self, row, col):
        if col not in (self._COL_QTY, self._COL_COST, self._COL_MARGIN):
            return
        self.table.blockSignals(True)
        qty, suffix = self._qty_for_row(row)
        # Do NOT show (Hours) in the cell — it is export-only
        cost_item = self.table.item(row, self._COL_COST)
        try:
            cost = float(cost_item.text()) if cost_item and cost_item.text() else 0.0
        except ValueError:
            cost = 0.0

        margin = self._row_margin_pct(row)
        marked_up_unit = _apply_margin(cost, margin)
        qty_num, _ = self._qty_for_row(row)
        total = qty_num * marked_up_unit

        def _set(c, text):
            item = self.table.item(row, c)
            if not item:
                item = QTableWidgetItem()
                item.setFlags(item.flags() & ~Qt.ItemFlag.ItemIsEditable)
                item.setForeground(QColor("#1a3a6e"))
                item.setBackground(QColor("#e8f0fe"))
                self.table.setItem(row, c, item)
            item.setText(text)

        line_total = qty * marked_up_unit
        _set(self._COL_MKUP,     f"${int(marked_up_unit):,}")
        _set(self._COL_LINE_TOT, f"${int(line_total):,}")
        _set(self._COL_TOTAL,    f"${int(line_total):,}")
        self.table.blockSignals(False)
        self._update_grand_total()

    def _update_grand_total(self):
        """Sum all Line Total cells and update the grand total label."""
        grand = 0.0
        for r in range(self.table.rowCount()):
            item = self.table.item(r, self._COL_LINE_TOT)
            if item and item.text():
                try: grand += float(item.text().replace('$', '').replace(',', ''))
                except ValueError: pass
        self._grand_total_lbl.setText(f"Grand Total:  ${int(grand):,}")

    def _refresh_all_totals(self):
        for row in range(self.table.rowCount()):
            self._safe_update_totals(row, self._COL_COST)
        self._on_preview_toggle(None)
        self._update_grand_total()

    def set_global_margin_ref(self, fn, spin=None):
        self._global_margin_ref = fn
        self._margin_widget.set_global_ref(fn)
        if hasattr(self, "_global_spin") and self._global_spin is not None:
            try: self._global_spin.valueChanged.disconnect(self._on_global_spin_changed)
            except Exception: pass
        self._global_spin = spin
        if spin is not None:
            spin.valueChanged.connect(self._on_global_spin_changed)

    def _on_global_spin_changed(self):
        """Called when the global margin value changes."""
        # Get value from the callable ref (works even when spin signals are blocked)
        v = self._global_margin_ref() if self._global_margin_ref else 20.0
        self._margin_widget._value = v
        if hasattr(self._margin_widget, '_lbl'):
            self._margin_widget._lbl.setText(f"{v:.1f} %")
        # Update Margin % cells — skip rows where user has manually overridden
        global_str = f"{v:.1f}"
        self.table.blockSignals(True)
        for row in range(self.table.rowCount()):
            item = self.table.item(row, self._COL_MARGIN)
            if item is None:
                continue
            cell_val = item.text().strip()
            # Only update if cell is empty OR still matches the previous global value
            # i.e. skip if the user typed something different
            prev_global = getattr(self, '_last_global_pct', None)
            if cell_val == '' or prev_global is None or cell_val == f"{prev_global:.1f}":
                item.setText(global_str if v > 0 else "")
        self._last_global_pct = v
        self.table.blockSignals(False)
        self._refresh_all_totals()

    def _delete_self(self):
        if self._on_remove: self._on_remove(self)
        self.setParent(None); self.deleteLater()

    def _update_bom_title(self, num: int):
        self.set_title(f"⊞  BOM Table — Section {num}")

    def get_data(self):
        # Export: Part Number, Qty, Sale Price (marked-up unit), Line Total
        # Cost and Margin % are UI-only and not exported
        rows = []
        for row in range(self.table.rowCount()):
            part_item = self.table.item(row, self._COL_PART)
            part_plain = part_item.text() if part_item else ""
            part_html  = (part_item.data(Qt.ItemDataRole.UserRole) or "") if part_item else ""
            rows.append([
                part_plain,
                (lambda r: (lambda q, s: (f"{int(q)}{s}" if q == int(q) else f"{q}{s}") if q else "")
                           (*self._qty_for_row(r)))(row),
                self.table.item(row, self._COL_MKUP).text()     if self.table.item(row, self._COL_MKUP)     else "",
                self.table.item(row, self._COL_LINE_TOT).text() if self.table.item(row, self._COL_LINE_TOT) else "",
                self.table.item(row, self._COL_TOTAL).text()    if self.table.item(row, self._COL_TOTAL)    else "",
                part_html,   # index 5 — rich HTML for Part Number
            ])
        return {"type": "table", "rows": rows}


# ── ParagraphSection ──────────────────────────────────────────────────────────
class ParagraphSection(CollapsibleCard):
    def __init__(self, on_remove=None):
        super().__init__("¶  Paragraph Section")
        self._on_remove = on_remove

        remove_btn = QPushButton("✕  Remove")
        remove_btn.setStyleSheet(REMOVE_BTN_STYLE)
        remove_btn.clicked.connect(self._delete_self)
        self._header_extras.addWidget(remove_btn)

        self.header = QLineEdit()
        self.header.setPlaceholderText("Section heading…")
        self.header.setStyleSheet(FIELD_STYLE + "QLineEdit { font-weight:600; }")
        self._section_num = 0
        self._dfw_ref = None
        self.header.textChanged.connect(self._on_header_text_changed)
        self._body_layout.addWidget(self.header)

        self.rich_editor = RichTextEditor()
        self.rich_editor.setPlaceholderText("Section body — supports bold, italic, bullets, and tables…")
        self.rich_editor.editor.document().contentsChanged.connect(self._notify_toc)
        self._body_layout.addWidget(self.rich_editor)

    def _on_header_text_changed(self, t):
        n = self._section_num
        self.set_title(f'{n}.  {t}' if (t and n) else
                       f'Section {n}' if n else 'Paragraph Section')
        self._notify_toc()

    def _notify_toc(self):
        # Try _dfw_ref first
        cb = getattr(getattr(self, '_dfw_ref', None), '_toc_refresh_cb', None)
        if callable(cb):
            cb(); return
        # Fallback: walk Qt parent widget chain
        p = self.parent()
        while p is not None:
            cb = getattr(p, '_toc_refresh_cb', None)
            if callable(cb):
                cb(); return
            p = p.parent()

    def update_section_label(self, num: int):
        self._section_num = num
        try:
            self.header.textChanged.disconnect()
        except Exception:
            pass
        self.header.textChanged.connect(self._on_header_text_changed)
        self._on_header_text_changed(self.header.text())
        self._notify_toc()

    def _delete_self(self):
        if self._on_remove: self._on_remove(self)
        self.setParent(None); self.deleteLater()

    def get_data(self):
        return {
            "type": "paragraph",
            "header": self.header.text(),
            "text":   self.rich_editor.toPlainText(),
            "html":   self.rich_editor.toHtml(),
        }


# ── DynamicFormWidget ─────────────────────────────────────────────────────────
class DynamicFormWidget:
    def __init__(self, scroll_layout, global_margin_ref=None):
        self.sections = []
        self.scroll_layout = scroll_layout
        self._global_margin_ref = global_margin_ref
        self._global_spin = None

    def set_global_margin_ref(self, fn, spin=None):
        self._global_margin_ref = fn
        self._global_spin = spin
        for s in self.sections:
            if isinstance(s, TableSection):
                s.set_global_margin_ref(fn, spin=spin)

    def _renumber_sections(self):
        for i, s in enumerate(self.sections):
            n = i + 1
            s._dfw_ref = self   # ensure back-reference always set
            if hasattr(s, 'update_section_label'):
                s.update_section_label(n)
            elif hasattr(s, '_update_bom_title'):
                s._update_bom_title(n)
        cb = getattr(self, '_toc_refresh_cb', None)
        if cb: cb()

    def _wire_arrows(self, section):
        section._btn_up.clicked.connect(lambda: self._move_section(section, -1))
        section._btn_down.clicked.connect(lambda: self._move_section(section, +1))

    def _move_section(self, section, direction):
        idx = self.sections.index(section)
        new_idx = idx + direction
        if new_idx < 0 or new_idx >= len(self.sections):
            return
        # Swap in list
        self.sections[idx], self.sections[new_idx] = (
            self.sections[new_idx], self.sections[idx])
        # Rebuild layout order
        for s in self.sections:
            self.scroll_layout.removeWidget(s)
        for s in self.sections:
            self.scroll_layout.addWidget(s)
        self._renumber_sections()

    def add_section(self, section_type: str):
        def on_remove(s):
            if s in self.sections: self.sections.remove(s)
            self._renumber_sections()

        if section_type == "table":
            section = TableSection(on_remove=on_remove,
                                   global_margin_ref=self._global_margin_ref)
            if hasattr(self, "_global_spin") and self._global_spin is not None:
                section.set_global_margin_ref(self._global_margin_ref, spin=self._global_spin)
        else:
            section = ParagraphSection(on_remove=on_remove)

        self._wire_arrows(section)
        section._dfw_ref = self
        self.sections.append(section)
        self.scroll_layout.addWidget(section)
        self._renumber_sections()

    def get_form_data(self):
        # Assign section numbers
        numbered = []
        sec_num = 0
        for s in self.sections:
            sec_num += 1
            d = s.get_data()
            d["section_number"] = sec_num
            numbered.append(d)
        return {"sections": numbered}

    def collect_project_state(self, top_fields: dict) -> dict:
        """Return complete project state as a plain dict (JSON-serialisable)."""
        return {
            "_version": 1,
            "fields": top_fields,
            "sections": self.get_form_data()["sections"],
        }

    def restore_project_state(self, state: dict):
        """Restore sections from a saved state dict."""
        # Remove all existing sections
        for s in list(self.sections):
            s.setParent(None)
            s.deleteLater()
        self.sections.clear()

        for sec in state.get("sections", []):
            stype = sec.get("type", "paragraph")
            self.add_section(stype)
            widget = self.sections[-1]

            if stype == "paragraph":
                widget.header.setText(sec.get("header", ""))
                html = sec.get("html", "")
                if html:
                    widget.rich_editor.editor.setHtml(html)
                else:
                    widget.rich_editor.editor.setPlainText(sec.get("text", ""))

            elif stype == "table":
                tbl = widget.table
                tbl.blockSignals(True)
                rows_data = sec.get("rows", [])
                # Fill existing row (add_section already added one blank)
                for r_i, row in enumerate(rows_data):
                    if r_i >= tbl.rowCount():
                        widget.add_row()
                    is_hours_row = False
                    for c_i, val in enumerate(row[:5]):
                        col = [widget._COL_PART, widget._COL_QTY,
                               widget._COL_MKUP, widget._COL_LINE_TOT,
                               widget._COL_TOTAL][c_i]
                        item = tbl.item(r_i, col)
                        if item:
                            if c_i == 1:  # qty — strip saved '(Hours)' suffix
                                val_str = str(val)
                                if '(Hours)' in val_str:
                                    is_hours_row = True
                                    val_str = val_str.replace('(Hours)', '').strip()
                                item.setText(val_str)
                            else:
                                item.setText(str(val))
                    if len(row) > 5:
                        item = tbl.item(r_i, widget._COL_PART)
                        if item:
                            item.setData(Qt.ItemDataRole.UserRole, row[5])
                    # Restore Hours? checkbox without triggering recalculation;
                    # _refresh_all_totals() below will use the correct state.
                    if is_hours_row and r_i < len(widget._hours_checkboxes):
                        widget._hours_checkboxes[r_i].blockSignals(True)
                        widget._hours_checkboxes[r_i].setChecked(True)
                        widget._hours_checkboxes[r_i].blockSignals(False)
                        widget._hours_rows.add(r_i)
                tbl.blockSignals(False)
                widget._refresh_all_totals()
