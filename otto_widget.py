"""
otto_widget.py — OTTO Robot BOM Configurator & Quote Generator
"""
from __future__ import annotations
import copy, json, os
from datetime import date as _date

from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QFormLayout, QLabel, QComboBox,
    QPushButton, QTableWidget, QTableWidgetItem, QHeaderView, QLineEdit,
    QSpinBox, QDoubleSpinBox, QCheckBox, QDialog, QDialogButtonBox,
    QFrame, QScrollArea, QFileDialog, QMessageBox, QSizePolicy, QAbstractItemView,
)
from PyQt6.QtCore import Qt, pyqtSignal
from PyQt6.QtGui import QColor, QBrush, QFont

from form_widgets import DragDropLabel   # reuse the existing drag-drop image widget

# Shared +/− button style — matches IBE widget exactly
_SPIN_BTN = (
    "QPushButton { background:#f4f6fa; color:#3a3a5c; border:1px solid #dde1e7;"
    "  border-radius:4px; font-size:14px; font-weight:700; padding:0; }"
    "QPushButton:hover { background:#e8ecf5; }"
    "QPushButton:pressed { background:#dde1e7; }"
    "QPushButton:disabled { color:#cccccc; border-color:#eeeeee; background:#f8f8f8; }"
)


# ─── Compact +/− quantity widget ─────────────────────────────────────────────
class _SpinWidget(QWidget):
    """[−] N [+] quantity control styled to match the IBE widget's +/− buttons."""
    valueChanged = pyqtSignal(int)

    def __init__(self, value: int = 1, min_val: int = 0, max_val: int = 9999,
                 parent=None):
        super().__init__(parent)
        self._value   = max(min_val, min(max_val, value))
        self._min     = min_val
        self._max     = max_val

        lay = QHBoxLayout(self)
        lay.setContentsMargins(2, 1, 2, 1)
        lay.setSpacing(3)

        self._btn_m = QPushButton("−")
        self._btn_m.setFixedSize(24, 24)
        self._btn_m.setStyleSheet(_SPIN_BTN)
        self._btn_m.clicked.connect(self._dec)

        self._lbl = QLabel(str(self._value))
        self._lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self._lbl.setMinimumWidth(26)
        self._lbl.setStyleSheet(
            "QLabel { font-size:11px; font-weight:600; color:#1a0509;"
            "  background:transparent; border:none; }")

        self._btn_p = QPushButton("+")
        self._btn_p.setFixedSize(24, 24)
        self._btn_p.setStyleSheet(_SPIN_BTN)
        self._btn_p.clicked.connect(self._inc)

        lay.addWidget(self._btn_m)
        lay.addWidget(self._lbl)
        lay.addWidget(self._btn_p)
        self.setStyleSheet("QWidget { background:transparent; border:none; }")

    def value(self) -> int:
        return self._value

    def setValue(self, v: int):
        self._value = max(self._min, min(self._max, int(v)))
        self._lbl.setText(str(self._value))

    def setEnabled(self, enabled: bool):
        super().setEnabled(enabled)
        self._btn_m.setEnabled(enabled)
        self._btn_p.setEnabled(enabled)
        self._lbl.setStyleSheet(
            f"QLabel {{ font-size:11px; font-weight:600;"
            f"  color:{'#1a0509' if enabled else '#aaaaaa'};"
            f"  background:transparent; border:none; }}")

    def _dec(self):
        if self._value > self._min:
            self._value -= 1
            self._lbl.setText(str(self._value))
            self.valueChanged.emit(self._value)

    def _inc(self):
        if self._value < self._max:
            self._value += 1
            self._lbl.setText(str(self._value))
            self.valueChanged.emit(self._value)


# ─── Contact sub-widget (mirrors _ContactWidget in main.py) ───────────────────
class _OTTOContactWidget(QWidget):
    _TITLE = "font-size:10px; font-weight:600; color:#3a3a5c; margin-bottom:0px; margin-top:2px;"
    _FIELD = (
        "QLineEdit { background:#ffffff; border:1px solid #d0d4de;"
        "  border-radius:4px; padding:2px 6px; font-size:11px; color:#1a1a2e; }"
        "QLineEdit:focus { border-color:#920d2e; }"
    )
    _COMPANY = "McNaughton-McKay Electric Co"

    def __init__(self, parent=None):
        super().__init__(parent)
        lay = QVBoxLayout(self)
        lay.setContentsMargins(0, 0, 0, 0); lay.setSpacing(2)

        lay.addWidget(self._hdr("Presented by:"))
        self._pb_name  = self._fld("Name");  lay.addWidget(self._pb_name)
        self._pb_email = self._fld("Email"); lay.addWidget(self._pb_email)

        lay.addWidget(self._hdr("Account Manager:"))
        self._am_name  = self._fld("Name");  lay.addWidget(self._am_name)
        self._am_email = self._fld("Email"); lay.addWidget(self._am_email)

    def _hdr(self, text):
        l = QLabel(text); l.setStyleSheet(self._TITLE); return l

    def _fld(self, ph):
        e = QLineEdit(); e.setPlaceholderText(ph); e.setStyleSheet(self._FIELD)
        e.setFixedHeight(20); return e

    def toPlainText(self) -> str:
        pb = self._pb_name.text().strip()  or "INSERTNAME"
        pe = self._pb_email.text().strip() or "INSERTEMAIL"
        am = self._am_name.text().strip()  or "INSERTNAME"
        ae = self._am_email.text().strip() or "INSERTEMAIL"
        return (f"{pb}\nSolutions Architect, {self._COMPANY}\n{pe}\n\n"
                f"{am}\nAccount Manager, {self._COMPANY}\n{ae}")

    def setPlainText(self, text: str):
        lines = (text or "").splitlines()
        def _g(i): return lines[i].strip() if i < len(lines) else ""
        def _c(v): return "" if v in ("INSERTNAME", "INSERTEMAIL") else v
        self._pb_name.setText(_c(_g(0)))
        self._pb_email.setText(_c(_g(2)))
        self._am_name.setText(_c(_g(4)))
        self._am_email.setText(_c(_g(6)))

    # Derived helpers for doc generation
    @property
    def pres_name(self):  return self._pb_name.text().strip()
    @property
    def pres_email(self): return self._pb_email.text().strip()
    @property
    def acct_name(self):  return self._am_name.text().strip()
    @property
    def acct_email(self): return self._am_email.text().strip()


# ─── BOM Data ─────────────────────────────────────────────────────────────────
_ITEMS: list[dict] = [
    dict(key="robot",      cat="hardware",  logic="Always",
         product="OTTO 100 V2.5, NA",
         part="26971",     price=42067.50,
         desc=("OTTO 100 V2.5 Series, Base Configuration. Non-CE Certified. "
               "150 kg max payload, 2 m/s max speed, integrated lift and attachment "
               "interface. No cart engagement plate. 15-month warranty from delivery.")),
    dict(key="fm",         cat="hardware",  logic="Always",
         product="OTTO 100 Fleet Manager \u2013 Perpetual License + Lifetime Upgrades",
         part="29763",     price=3000.00,
         desc=("OTTO Fleet Management Software for heterogeneous fleet up to 100 AMRs. "
               "Fleet commissioning, factory integration, job allocation, system analytics. "
               "Per-AMR license with lifetime software upgrade access.")),
    dict(key="manual",     cat="hardware",  logic="Strongly Encouraged",
         product="OTTO 100 v2.5 Manual \u2013 English",
         part="020074-01", price=61.77,
         desc="OTTO 100 v2.5 Manual Printed OMM-000081 \u2013 ENGLISH"),
    dict(key="fcharger",   cat="hardware",  logic="Strongly Encouraged",
         product="OTTO 100 V2.4+ Fast Charger \u2013 NA, 480V",
         part="18044",     price=18637.50,
         desc=("Autonomous High Speed Charging Station for OTTO 100 V2.4+ (480 VAC NA). "
               "2\u00d7 faster than standard charger. Licensed electrician required. "
               "15-month warranty.")),
    dict(key="mcharger",   cat="hardware",  logic="Always",
         product="OTTO 100 V2.4+ Manual Charger",
         part="017326",    price=1059.68,
         desc=("Manual Charging Station for OTTO 100 V2.4+. Trickle charge during "
               "maintenance or shutdown. Standard power receptacle required.")),
    dict(key="cart_base",  cat="hardware",  logic="Option 1",
         product="OTTO 100 Staging Cart \u2013 Baseplate Config",
         part="017307",    price=740.18,
         desc=("Staging Cart \u2013 Load Plate Configuration. 130 kg max payload, "
               "normally-closed ergonomic brake and load plate.")),
    dict(key="cart_uline", cat="hardware",  logic="Option 2",
         product="OTTO 100 Staging Cart \u2013 Uline Shelf Config",
         part="017309",    price=1059.68,
         desc=("Staging Cart \u2013 Shelf Configuration. 130 kg max payload, "
               "normally-closed ergonomic brake, integrated Uline shelving.")),
    dict(key="engplate",   cat="hardware",  logic="Always when carts used",
         product="OTTO 100 V2.4 and v2.5 Staging Cart Engagement Plate",
         part="017770",    price=260.93,
         desc=("Engagement Plate \u2013 mates OTTO 100 V2.4/V2.5 to staging cart. "
               "Lateral accuracy \u00b110 mm.")),
    dict(key="asem",       cat="hardware",  logic="Optional",
         product="ASEM Box PC 6300B-SW2ACAN-CCCNENNANAE-BBNNN",
         part="6300B-SW2ACAN-CCCNENNANAE-BBNNN",
         price=5701.97,
         desc=("ASEM Box PC, 24VDC. i7 Embedded, 32 GB RAM, 960 GB 2.5\u2033 SSD. "
               "Windows 11 IoT Enterprise LTSC 2024.")),
    dict(key="commission", cat="service",   logic="Optional",
         product="Robotic Commissioning",
         part="",           price=50000.00,
         desc="Install Fleet Manager, Site Survey, Initial Integration, Robot Dealer Prep"),
]

_LEVEL_DEFAULTS: dict[str, dict[str, tuple[bool, int]]] = {
    "Bare Min": {
        "robot":      (True,  1), "fm":         (True,  1), "manual":     (True,  1),
        "fcharger":   (False, 1), "mcharger":   (True,  1),
        "cart_base":  (True,  1), "cart_uline": (False, 1), "engplate":   (True,  1),
        "asem":       (False, 1), "commission": (False, 1),
    },
    "Basic": {
        "robot":      (True,  1), "fm":         (True,  1), "manual":     (True,  1),
        "fcharger":   (True,  1), "mcharger":   (True,  1),
        "cart_base":  (True,  1), "cart_uline": (False, 1), "engplate":   (True,  1),
        "asem":       (True,  1), "commission": (True,  1),
    },
    "MCMC Special": {
        "robot":      (True,  1), "fm":         (True,  1), "manual":     (True,  1),
        "fcharger":   (True,  1), "mcharger":   (True,  1),
        "cart_base":  (True,  1), "cart_uline": (False, 1), "engplate":   (True,  1),
        "asem":       (True,  1), "commission": (True,  1),
    },
}

_COMMISSION_PRICE: dict[str, float] = {
    "Bare Min": 0.0, "Basic": 1000.0, "MCMC Special": 50000.0,
}

# Row background / foreground per logic type
_LOGIC_STYLE: dict[str, tuple[str, str]] = {
    "Always":                 ("#f4f4f4", "#999999"),   # neutral grey — locked/required
    "Strongly Encouraged":    ("#fffae8", "#7a5a00"),   # amber
    "Option 1":               ("#e8f0ff", "#0a3a7a"),   # blue  — mutually exclusive pair
    "Option 2":               ("#fff4e0", "#7a4200"),   # amber — mutually exclusive pair
    "Always when carts used": ("#f0f5f0", "#2a7a4a"),   # soft green
    "Optional":               ("#ffffff", "#555577"),
}

# Checkbox styles — match IBE widget's red brand theme
_CB_ENABLED = (
    "QCheckBox { spacing:4px; }"
    "QCheckBox::indicator { width:16px; height:16px; border-radius:3px; }"
    "QCheckBox::indicator:unchecked { border:2px solid #d6c0c5; background:#ffffff; }"
    "QCheckBox::indicator:unchecked:hover { border-color:#920d2e; background:#fdf4f6; }"
    "QCheckBox::indicator:checked { border:2px solid #920d2e; background:#920d2e; }"
    "QCheckBox::indicator:checked:hover { border-color:#b01036; background:#b01036; }"
)
# Locked (Always) rows — checked + disabled, neutral grey fill
_CB_LOCKED = (
    "QCheckBox { spacing:4px; }"
    "QCheckBox::indicator { width:16px; height:16px; border-radius:3px; }"
    "QCheckBox::indicator:checked:disabled { border:2px solid #aaaaaa; background:#aaaaaa; }"
    "QCheckBox::indicator:unchecked:disabled { border:2px solid #cccccc; background:#eeeeee; }"
)

_CC, _CPROD, _CCAT, _CLOGIC, _CQTY, _CPRICE, _CTOT = range(7)

# ─── Shared style constants ───────────────────────────────────────────────────
_BTN = (
    "QPushButton { background:#920d2e; color:#ffffff; border:none;"
    "  border-radius:5px; padding:5px 16px; font-size:11px; font-weight:600; }"
    "QPushButton:hover { background:#b01036; }"
    "QPushButton:disabled { background:#cca0a8; }"
)
_GHOST = (
    "QPushButton { background:#f4f6fa; color:#3a3a5c; border:1px solid #dde1e7;"
    "  border-radius:5px; padding:4px 14px; font-size:11px; }"
    "QPushButton:hover { background:#eaecf4; }"
)
_COMBO = (
    "QComboBox { background:#ffffff; border:1px solid #d0d4de; border-radius:4px;"
    "  padding:3px 8px; font-size:11px; color:#1a1a2e; }"
    "QComboBox::drop-down { border:none; width:18px; }"
    "QComboBox:focus { border-color:#920d2e; }"
    # Dropdown list — force white bg and dark text (fixes black-on-black on Windows)
    "QComboBox QAbstractItemView { background:#ffffff; color:#1a1a2e;"
    "  border:1px solid #d0d4de; selection-background-color:#f5d0da;"
    "  selection-color:#920d2e; outline:none; }"
)
_FIELD = (
    "QLineEdit { background:#ffffff; border:1px solid #d0d4de; border-radius:5px;"
    "  padding:6px 12px; font-size:12px; font-family:'Aptos'; color:#1a0509; }"
    "QLineEdit:focus { border-color:#920d2e; }"
)
_GROUP = (
    "QGroupBox { font-size:11px; font-weight:600; color:#920d2e;"
    "  border:1px solid #d6c0c5; border-radius:6px; margin-top:6px; padding-top:4px; }"
    "QGroupBox::title { subcontrol-origin:margin; left:8px; top:0px; }"
)
_FORM_LBL = "font-weight:700; color:#6a4050; font-size:12px; font-family:'Aptos';"


# ─── Main Widget ──────────────────────────────────────────────────────────────
class OTTOWidget(QWidget):
    """OTTO Robot BOM configurator and quote generator."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self._custom_items: list[dict] = []
        self._block_update = False
        self._checkboxes:     dict[int, QCheckBox]    = {}
        self._qty_spins:      dict[int, _SpinWidget]  = {}
        self._row_keys:       dict[int, str | None]   = {}
        self._row_custom_idx: dict[int, int]          = {}
        self._product_widgets:dict[int, QWidget]      = {}
        self._active_tab = "hardware"
        self._setup_ui()
        self._rebuild_table()

    # ─── UI construction ──────────────────────────────────────────────────────

    def _setup_ui(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        _scroll = QScrollArea()
        _scroll.setWidgetResizable(True)
        _scroll.setFrameShape(QFrame.Shape.NoFrame)
        _scroll.setStyleSheet(
            "QScrollArea { background:transparent; border:none; }"
            "QScrollBar:vertical { background:#f0f2f7; width:8px; margin:0; }"
            "QScrollBar::handle:vertical { background:#c8cedd; border-radius:4px;"
            "  min-height:20px; }"
            "QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical { height:0; }")
        _content = QWidget()
        _cl = QVBoxLayout(_content)
        _cl.setContentsMargins(12, 10, 12, 10)
        _cl.setSpacing(8)

        # ── Collapsible project info panel ───────────────────────────────
        _oti_hdr = QWidget()
        _oti_hdr.setCursor(Qt.CursorShape.PointingHandCursor)
        _oti_hdr.setFixedHeight(32)
        _oti_hdr.setStyleSheet(
            "QWidget { background:#f9f0f2; border:1px solid #d6c0c5;"
            "  border-radius:6px 6px 0 0; }")
        _oti_row = QHBoxLayout(_oti_hdr)
        _oti_row.setContentsMargins(10, 0, 10, 0); _oti_row.setSpacing(6)
        _oti_lbl = QLabel("Project Information")
        _oti_lbl.setStyleSheet(
            "font-size:12px; font-weight:700; color:#920d2e;"
            "  background:transparent; border:none;")
        self._oti_toggle = QPushButton("▼")
        self._oti_toggle.setFixedSize(22, 22)
        self._oti_toggle.setStyleSheet(
            "QPushButton { background:transparent; border:none;"
            "  font-size:12px; color:#920d2e; }"
            "QPushButton:hover { color:#7a0b27; }")
        _oti_row.addWidget(self._oti_toggle); _oti_row.addWidget(_oti_lbl)
        _oti_row.addStretch()

        info_panel = QWidget()
        info_panel.setStyleSheet(
            "QWidget#otto_info { background:#ffffff; border:1px solid #d6c0c5;"
            "  border-top:none; border-radius:0 0 6px 6px; }"
        )
        info_panel.setObjectName("otto_info")
        ip_layout = QHBoxLayout(info_panel)
        ip_layout.setContentsMargins(10, 5, 10, 5)
        ip_layout.setSpacing(12)

        # Left: form fields
        form = QFormLayout()
        form.setHorizontalSpacing(10)
        form.setVerticalSpacing(3)
        form.setFieldGrowthPolicy(QFormLayout.FieldGrowthPolicy.ExpandingFieldsGrow)

        def _fe(ph: str) -> QLineEdit:
            e = QLineEdit(); e.setPlaceholderText(ph)
            e.setStyleSheet(_FIELD); e.setMinimumWidth(240); e.setFixedHeight(22); return e

        self.edit_project  = _fe("Enter project name\u2026")
        self.edit_proposal = _fe("MCMX-CUSTNAME-YYYYMMDD V1.0")
        self.edit_customer = _fe("Enter customer name\u2026")
        self.edit_location = _fe("City, State / Country\u2026")
        self._contact_w    = _OTTOContactWidget()
        self._contact_w.setMinimumWidth(280)

        for lbl_txt, w in [
            ("Project Name",    self.edit_project),
            ("Proposal Number", self.edit_proposal),
            ("Customer Name",   self.edit_customer),
            ("Customer Location", self.edit_location),
            ("Contact Info",    self._contact_w),
        ]:
            lbl = QLabel(lbl_txt); lbl.setStyleSheet(_FORM_LBL)
            form.addRow(lbl, w)
        ip_layout.addLayout(form, stretch=1)
        ip_layout.addStretch()

        # Right: customer logo (same pink container as Quote Generator)
        logo_box = QWidget()
        logo_box.setMaximumWidth(102)
        logo_box.setStyleSheet(
            "QWidget { background:#f9f0f2; border:1px solid #d6c0c5;"
            "  border-radius:6px; padding:3px; }"
            "QPushButton { border-radius:4px; }"
        )
        lb_layout = QVBoxLayout(logo_box)
        lb_layout.setContentsMargins(2, 2, 2, 2)
        lb_layout.setSpacing(3)

        self._picture_label = DragDropLabel(parent=logo_box)
        self._picture_label.setFixedSize(90, 90)
        lb_layout.addWidget(self._picture_label)

        btn_upload = QPushButton("\u2b06  Upload")
        btn_upload.setObjectName("button_upload_picture")   # picks up APP_STYLE rules
        btn_upload.setFixedSize(90, 22)
        btn_upload.clicked.connect(self._upload_picture)
        lb_layout.addWidget(btn_upload)

        btn_clear = QPushButton("\u2715  Clear")
        btn_clear.setObjectName("button_clear_picture")
        btn_clear.setFixedSize(90, 22)
        btn_clear.clicked.connect(self._picture_label.clear_image)
        lb_layout.addWidget(btn_clear)

        ip_layout.addWidget(logo_box)

        _oti_wrap = QWidget()
        _oti_wrap.setStyleSheet("QWidget{background:transparent;border:none;}")
        _oti_vl = QVBoxLayout(_oti_wrap)
        _oti_vl.setContentsMargins(0, 0, 0, 0); _oti_vl.setSpacing(0)
        _oti_vl.addWidget(_oti_hdr)
        _oti_vl.addWidget(info_panel)
        _cl.addWidget(_oti_wrap)

        def _oti_collapse():
            info_panel.setVisible(False)
            self._oti_toggle.setText("►")
            _oti_hdr.setStyleSheet(
                "QWidget { background:#f9f0f2; border:1px solid #d6c0c5;"
                "  border-radius:6px; }")

        def _oti_do_toggle():
            vis = not info_panel.isVisible()
            if vis:
                _sel_collapse()
            info_panel.setVisible(vis)
            self._oti_toggle.setText("▼" if vis else "►")
            _oti_hdr.setStyleSheet(
                "QWidget { background:#f9f0f2; border:1px solid #d6c0c5;"
                f"  border-radius:{'6px 6px 0 0' if vis else '6px'}; }}")
        self._oti_toggle.clicked.connect(_oti_do_toggle)
        _oti_hdr.mousePressEvent = lambda _e: _oti_do_toggle()

        # ── Version row ──────────────────────────────────────────────────
        _VER_BTN = (
            "QPushButton { background:#f4f6fa; color:#3a3a5c; border:1px solid #c8cedd;"
            "  border-radius:4px; font-size:16px; font-weight:700; }"
            "QPushButton:hover { background:#e0e4ef; }"
            "QPushButton:pressed { background:#d0d8ef; }")
        _VER_VAL = (
            "QLabel { background:#ffffff; border:1px solid #d6c0c5;"
            "  border-radius:4px; color:#1a0509; font-size:12px;"
            "  font-weight:700; padding:4px 8px; }")

        class _VerSpin:
            """Lightweight +/− spinner matching the Quote Generator style."""
            def __init__(self_, lo, hi, initial, on_change):
                self_._val = initial; self_._lo = lo; self_._hi = hi; self_._cb = on_change
                self_.widget = QWidget()
                _r = QHBoxLayout(self_.widget)
                _r.setContentsMargins(0, 0, 0, 0); _r.setSpacing(3)
                self_._minus = QPushButton("−"); self_._minus.setFixedSize(28, 28)
                self_._minus.setStyleSheet(_VER_BTN); self_._minus.clicked.connect(self_._dec)
                _r.addWidget(self_._minus)
                self_._lbl = QLabel(str(initial)); self_._lbl.setFixedWidth(36)
                self_._lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
                self_._lbl.setStyleSheet(_VER_VAL)
                _r.addWidget(self_._lbl)
                self_._plus = QPushButton("+"); self_._plus.setFixedSize(28, 28)
                self_._plus.setStyleSheet(_VER_BTN); self_._plus.clicked.connect(self_._inc)
                _r.addWidget(self_._plus)
            def value(self_): return self_._val
            def setValue(self_, v):
                self_._val = max(self_._lo, min(self_._hi, int(v)))
                self_._lbl.setText(str(self_._val))
            def _inc(self_): self_.setValue(self_._val + 1); self_._cb()
            def _dec(self_): self_.setValue(self_._val - 1); self_._cb()

        _ver_row_w = QWidget()
        _ver_row_w.setFixedHeight(44)
        _ver_row_w.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Fixed)
        _vrl = QHBoxLayout(_ver_row_w)
        _vrl.setContentsMargins(0, 6, 0, 6); _vrl.setSpacing(0)

        _vlbl = QLabel("Version:")
        _vlbl.setStyleSheet(
            "font-size:12px; font-weight:700; color:#920d2e;"
            " background:transparent; border:none;")
        _vrl.addWidget(_vlbl); _vrl.addSpacing(10)

        self._version_major_spin = _VerSpin(1, 99, 1, self._update_version_label)
        _vrl.addWidget(self._version_major_spin.widget)

        _vdot = QLabel(".")
        _vdot.setStyleSheet(
            "font-size:16px; font-weight:700; color:#1a0509;"
            " background:transparent; border:none;")
        _vdot.setFixedWidth(10); _vdot.setAlignment(Qt.AlignmentFlag.AlignCenter)
        _vrl.addWidget(_vdot)

        self._version_minor_spin = _VerSpin(0, 99, 0, self._update_version_label)
        _vrl.addWidget(self._version_minor_spin.widget)

        _vrl.addSpacing(12)
        self._version_badge = QLabel("V1.0")
        self._version_badge.setStyleSheet(
            "QLabel { background:#9E1B32; color:#ffffff; font-size:11px;"
            "  font-weight:700; border-radius:4px; padding:3px 10px; }")
        _vrl.addWidget(self._version_badge)
        _vrl.addStretch()
        _cl.addWidget(_ver_row_w)

        # ── Change History panel ─────────────────────────────────────────
        from PyQt6.QtWidgets import QListWidget as _QLW
        self._history_collapsed = True
        self._hist_outer = QWidget()
        _hist_vbox = QVBoxLayout(self._hist_outer)
        _hist_vbox.setContentsMargins(0, 0, 0, 4); _hist_vbox.setSpacing(4)

        self._btn_collapse_hist = QPushButton("▶  Change History")
        self._btn_collapse_hist.setStyleSheet(
            "QPushButton { background:#ffffff; border:1px solid #d6c0c5;"
            "  border-radius:5px; padding:6px 14px; font-size:12px;"
            "  font-weight:600; color:#920d2e; text-align:left; }"
            "QPushButton:hover { background:#fdf0f3; }")
        self._btn_collapse_hist.clicked.connect(self._toggle_history_panel)
        _hist_vbox.addWidget(self._btn_collapse_hist)

        self._history_panel = QWidget()
        self._history_panel.setVisible(False)
        _hp_vbox = QVBoxLayout(self._history_panel)
        _hp_vbox.setContentsMargins(0, 2, 0, 2); _hp_vbox.setSpacing(4)

        self._version_history = _QLW()
        self._version_history.setMinimumHeight(80)
        self._version_history.setMaximumHeight(130)
        self._version_history.setStyleSheet(
            "QListWidget { background:#ffffff; border:1px solid #d6c0c5;"
            "  border-radius:4px; font-size:11px; color:#1a0509; }"
            "QListWidget::item { padding:4px 8px; }"
            "QListWidget::item:selected { background:#f5d0da; color:#920d2e; }")
        _hp_vbox.addWidget(self._version_history)

        _hbr = QWidget(); _hbrl = QHBoxLayout(_hbr)
        _hbrl.setContentsMargins(0, 0, 0, 0); _hbrl.setSpacing(8)
        _hadd = QPushButton("+ Add Entry")
        _hadd.setStyleSheet(
            "QPushButton { background:#f4f6fa; color:#3a3a5c;"
            "  border:1px solid #dde1e7; border-radius:4px;"
            "  padding:4px 12px; font-size:11px; }"
            "QPushButton:hover { background:#e8ecf5; border-color:#b0b8d0; }")
        _hadd.clicked.connect(self._add_version_entry)
        _hbrl.addWidget(_hadd)
        _hdel = QPushButton("Remove")
        _hdel.setStyleSheet(
            "QPushButton { background:transparent; color:#e05252;"
            "  border:1px solid #e05252; border-radius:4px;"
            "  padding:4px 12px; font-size:11px; }"
            "QPushButton:hover { background:#fdf2f2; }")
        _hdel.clicked.connect(self._remove_version_entry)
        _hbrl.addWidget(_hdel); _hbrl.addStretch()
        _hp_vbox.addWidget(_hbr)

        _hist_vbox.addWidget(self._history_panel)
        _cl.addWidget(self._hist_outer)

        # ── "Selections" collapsible (Robot + BOM) ───────────────────────
        _sel_hdr = QWidget()
        _sel_hdr.setCursor(Qt.CursorShape.PointingHandCursor)
        _sel_hdr.setFixedHeight(32)
        _sel_hdr.setStyleSheet(
            "QWidget { background:#f9f0f2; border:1px solid #d6c0c5;"
            "  border-radius:6px; }")          # collapsed by default → all corners rounded
        _sel_hdr_row = QHBoxLayout(_sel_hdr)
        _sel_hdr_row.setContentsMargins(10, 0, 10, 0); _sel_hdr_row.setSpacing(6)
        _sel_lbl = QLabel("Selections")
        _sel_lbl.setStyleSheet(
            "font-size:12px; font-weight:700; color:#920d2e;"
            "  background:transparent; border:none;")
        self._sel_toggle = QPushButton("►")
        self._sel_toggle.setFixedSize(22, 22)
        self._sel_toggle.setStyleSheet(
            "QPushButton { background:transparent; border:none;"
            "  font-size:12px; color:#920d2e; }"
            "QPushButton:hover { color:#7a0b27; }")
        _sel_hdr_row.addWidget(self._sel_toggle); _sel_hdr_row.addWidget(_sel_lbl)
        _sel_hdr_row.addStretch()

        sel_body = QWidget()
        sel_body.setVisible(False)             # collapsed by default
        sel_body.setStyleSheet("QWidget{background:transparent;border:none;}")
        sel_body_layout = QVBoxLayout(sel_body)
        sel_body_layout.setContentsMargins(0, 4, 0, 0)
        sel_body_layout.setSpacing(4)

        # ── Robot + level selectors ───────────────────────────────────────
        sel_bar = QWidget()
        sel_bar.setStyleSheet(
            "QWidget { background:#f9f0f2; border:1px solid #d6c0c5;"
            "  border-radius:6px; }"
            "QLabel { background:transparent; border:none; color:#920d2e;"
            "  font-weight:700; font-size:11px; font-family:'Aptos'; }")
        sel = QHBoxLayout(sel_bar)
        sel.setContentsMargins(10, 6, 10, 6); sel.setSpacing(10)

        lbl_r = QLabel("Robot:")
        self._robot_cb = QComboBox()
        self._robot_cb.setStyleSheet(_COMBO); self._robot_cb.setFixedWidth(130)
        self._robot_cb.addItem("OTTO 100")

        lbl_l = QLabel("Service Level:")
        self._level_cb = QComboBox()
        self._level_cb.setStyleSheet(_COMBO); self._level_cb.setFixedWidth(150)
        for lvl in ("Bare Min", "Basic", "MCMC Special"):
            self._level_cb.addItem(lvl)
        self._level_cb.setCurrentText("Basic")
        self._level_cb.currentTextChanged.connect(self._on_level_changed)

        sel.addWidget(lbl_r); sel.addWidget(self._robot_cb)
        sel.addSpacing(12)
        sel.addWidget(lbl_l); sel.addWidget(self._level_cb)
        sel.addStretch()
        sel_body_layout.addWidget(sel_bar)

        # ── Hardware / Services tab bar ────────────────────────────────────
        _TAB_ON  = ("QPushButton { background:#920d2e; color:#ffffff; border:none;"
                    "  border-radius:5px 5px 0 0; padding:6px 22px;"
                    "  font-size:12px; font-weight:700; }")
        _TAB_OFF = ("QPushButton { background:#f9f0f2; color:#920d2e;"
                    "  border:1px solid #d6c0c5; border-bottom:none;"
                    "  border-radius:5px 5px 0 0; padding:6px 22px;"
                    "  font-size:12px; font-weight:600; }"
                    "QPushButton:hover { background:#fde8ed; }")
        _tab_bar = QWidget()
        _tab_bar.setStyleSheet("QWidget { background:transparent; border:none; }")
        _tb_lay  = QHBoxLayout(_tab_bar)
        _tb_lay.setContentsMargins(0, 4, 0, 0); _tb_lay.setSpacing(2)
        self._tab_hw_btn  = QPushButton("Hardware")
        self._tab_svc_btn = QPushButton("Services")
        self._tab_hw_btn.setStyleSheet(_TAB_ON)
        self._tab_svc_btn.setStyleSheet(_TAB_OFF)
        self._tab_hw_btn.clicked.connect(lambda: self._set_bom_tab("hardware"))
        self._tab_svc_btn.clicked.connect(lambda: self._set_bom_tab("service"))
        _tb_lay.addWidget(self._tab_hw_btn); _tb_lay.addWidget(self._tab_svc_btn)
        _tb_lay.addStretch()
        sel_body_layout.addWidget(_tab_bar)

        # ── BOM table ─────────────────────────────────────────────────────
        self._table = QTableWidget()
        self._table.setMinimumHeight(220)
        self._table.setColumnCount(7)
        self._table.setHorizontalHeaderLabels(
            ["\u2713", "Product / Description", "Category", "Option", "Qty", "Unit Price", "Line Total"])
        hdr = self._table.horizontalHeader()
        hdr.setDefaultAlignment(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter)
        hdr.setSectionResizeMode(QHeaderView.ResizeMode.Fixed)
        hdr.setSectionResizeMode(_CPROD, QHeaderView.ResizeMode.Stretch)
        self._table.setColumnWidth(_CC,     20)
        self._table.setColumnWidth(_CCAT,   90)
        self._table.setColumnWidth(_CLOGIC, 80)
        self._table.setColumnWidth(_CQTY,   86)
        self._table.setColumnWidth(_CPRICE, 96)
        self._table.setColumnWidth(_CTOT,   96)
        self._table.setColumnHidden(_CCAT, True)   # Category shown via tab instead
        self._table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        self._table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self._table.verticalHeader().setVisible(False)
        self._table.setShowGrid(False)
        self._table.setStyleSheet(
            "QTableWidget { border:1px solid #d6c0c5; border-radius:6px;"
            "  font-size:10px; color:#1a0509; background:#ffffff; }"
            "QHeaderView::section { background:#f9f0f2; color:#920d2e; font-weight:700;"
            "  border-bottom:2px solid #d6c0c5; border-right:none;"
            "  padding:5px 6px; font-size:10px; font-family:'Aptos'; }"
            "QTableWidget::item { padding:3px 6px; border-bottom:1px solid #f0e8ea; }"
            "QTableWidget::item:selected { background:#fde8ed; color:#920d2e; }"
            "QTableWidget::item:hover { background:#fdf4f6; }"
        )
        sel_body_layout.addWidget(self._table, stretch=3)

        # ── Add / remove custom items ─────────────────────────────────────
        add_row = QHBoxLayout(); add_row.setSpacing(8)
        btn_add_hw  = QPushButton("\u002b Add Hardware Item")
        btn_add_svc = QPushButton("\u002b Add Service Item")
        btn_remove  = QPushButton("\u2715 Remove Selected")
        for b in (btn_add_hw, btn_add_svc, btn_remove):
            b.setStyleSheet(_GHOST)
        btn_add_hw.clicked.connect(lambda: self._add_custom("hardware"))
        btn_add_svc.clicked.connect(lambda: self._add_custom("service"))
        btn_remove.clicked.connect(self._remove_custom)
        add_row.addWidget(btn_add_hw); add_row.addWidget(btn_add_svc)
        add_row.addWidget(btn_remove); add_row.addStretch()
        sel_body_layout.addLayout(add_row)

        _sel_wrap = QWidget()
        _sel_wrap.setStyleSheet("QWidget{background:transparent;border:none;}")
        _sel_vl = QVBoxLayout(_sel_wrap)
        _sel_vl.setContentsMargins(0, 0, 0, 0); _sel_vl.setSpacing(0)
        _sel_vl.addWidget(_sel_hdr)
        _sel_vl.addWidget(sel_body, stretch=1)
        _cl.addWidget(_sel_wrap)

        def _sel_collapse():
            sel_body.setVisible(False)
            self._sel_toggle.setText("►")
            _sel_hdr.setStyleSheet(
                "QWidget { background:#f9f0f2; border:1px solid #d6c0c5;"
                "  border-radius:6px; }")

        def _sel_do_toggle():
            vis = not sel_body.isVisible()
            if vis:
                _oti_collapse()
            sel_body.setVisible(vis)
            self._sel_toggle.setText("▼" if vis else "►")
            _sel_hdr.setStyleSheet(
                "QWidget { background:#f9f0f2; border:1px solid #d6c0c5;"
                f"  border-radius:{'6px 6px 0 0' if vis else '6px'}; }}")
        self._sel_toggle.clicked.connect(_sel_do_toggle)
        _sel_hdr.mousePressEvent = lambda _e: _sel_do_toggle()

        # ── Summary / invoice panel ───────────────────────────────────────
        from PyQt6.QtWidgets import QGroupBox
        sum_box = QGroupBox("Cost Summary \u2013 Invoice Schedule")
        sum_box.setStyleSheet(_GROUP)
        sl = QHBoxLayout(sum_box)
        sl.setContentsMargins(16, 10, 16, 10); sl.setSpacing(0)

        def _sum_row(caption, lbl_widget, bold_caption=False):
            r = QHBoxLayout(); r.setSpacing(8)
            t = QLabel(caption)
            t.setStyleSheet(
                f"font-size:{'12' if bold_caption else '11'}px;"
                f"font-weight:{'700' if bold_caption else '400'};"
                " color:#555577;")
            r.addWidget(t); r.addWidget(lbl_widget); r.addStretch()
            return r

        self._inv: dict[str, QLabel] = {}

        # ── Hardware column ───────────────────────────────────────────
        hw_col = QVBoxLayout(); hw_col.setSpacing(4)
        hw_hdr = QLabel("Hardware")
        hw_hdr.setStyleSheet("font-size:13px; font-weight:700; color:#920d2e;")
        hw_col.addWidget(hw_hdr)
        self._lbl_hw = self._mk_sum_lbl(bold=True, large=True)
        hw_col.addLayout(_sum_row("Total:", self._lbl_hw, bold_caption=True))
        sep1 = QFrame(); sep1.setFrameShape(QFrame.Shape.HLine)
        sep1.setStyleSheet("color:#e0e4ec; margin:2px 0;"); hw_col.addWidget(sep1)
        hw_inv_lbl = QLabel("Invoice Schedule:")
        hw_inv_lbl.setStyleSheet("font-size:11px; font-weight:700; color:#555577;")
        hw_col.addWidget(hw_inv_lbl)
        for k in ("HDW-1ST (50%)", "HDW-2ND (40%)", "HDW-3RD (10%)"):
            lbl = self._mk_sum_lbl(); self._inv[k] = lbl
            hw_col.addLayout(_sum_row(k + ":", lbl))
        sl.addLayout(hw_col, stretch=1)
        sl.addSpacing(24)

        # ── Services column ───────────────────────────────────────────
        svc_col = QVBoxLayout(); svc_col.setSpacing(4)
        svc_hdr = QLabel("Services")
        svc_hdr.setStyleSheet("font-size:13px; font-weight:700; color:#920d2e;")
        svc_col.addWidget(svc_hdr)
        self._lbl_svc = self._mk_sum_lbl(bold=True, large=True)
        svc_col.addLayout(_sum_row("Total:", self._lbl_svc, bold_caption=True))
        sep2 = QFrame(); sep2.setFrameShape(QFrame.Shape.HLine)
        sep2.setStyleSheet("color:#e0e4ec; margin:2px 0;"); svc_col.addWidget(sep2)
        svc_inv_lbl = QLabel("Invoice Schedule:")
        svc_inv_lbl.setStyleSheet("font-size:11px; font-weight:700; color:#555577;")
        svc_col.addWidget(svc_inv_lbl)
        for k in ("SRV-1ST (50%)", "SRV-2ND (50%)"):
            lbl = self._mk_sum_lbl(); self._inv[k] = lbl
            svc_col.addLayout(_sum_row(k + ":", lbl))
        sl.addLayout(svc_col, stretch=1)
        sl.addSpacing(24)

        # ── Grand total column ────────────────────────────────────────
        gt_col = QVBoxLayout(); gt_col.setSpacing(4)
        gt_col.addStretch()
        gt_hdr = QLabel("Grand Total")
        gt_hdr.setStyleSheet("font-size:13px; font-weight:700; color:#920d2e;")
        gt_col.addWidget(gt_hdr)
        self._lbl_tot = self._mk_sum_lbl(bold=True, large=True)
        self._lbl_tot.setStyleSheet(
            "font-size:18px; font-weight:700; color:#920d2e;")
        gt_col.addWidget(self._lbl_tot)
        gt_col.addStretch()
        sl.addLayout(gt_col, stretch=1)

        _cl.addWidget(sum_box)
        _cl.addStretch()

        _scroll.setWidget(_content)
        root.addWidget(_scroll, 1)

        # ── Bottom action bar ─────────────────────────────────────────────
        _bar = QWidget()
        _bar.setStyleSheet("QWidget { background:#f9f0f2; border-top:1px solid #d6c0c5; }")
        _bar.setFixedHeight(52)
        br = QHBoxLayout(_bar)
        br.setContentsMargins(16, 8, 16, 8)
        br.setSpacing(10)

        br.addStretch()

        btn_load = QPushButton("⬆  Import .mcmxo")
        btn_load.setStyleSheet(
            "QPushButton { background:#ffffff; color:#920d2e;"
            "  border:1px solid #d6c0c5; border-radius:5px;"
            "  padding:7px 18px; font-size:12px; font-weight:600; }"
            "QPushButton:hover { background:#fdf0f3; border-color:#920d2e; }"
        )
        btn_load.clicked.connect(self.load_project)
        br.addWidget(btn_load)

        btn_save = QPushButton("💾  Save .mcmxo")
        btn_save.setStyleSheet(
            "QPushButton { background:#ffffff; color:#3a3a5c;"
            "  border:1px solid #d6c0c5; border-radius:5px;"
            "  padding:7px 18px; font-size:12px; font-weight:600; }"
            "QPushButton:hover { background:#f4f6fa; }"
        )
        btn_save.clicked.connect(self.save_project)
        br.addWidget(btn_save)

        btn_gen = QPushButton("Generate Document")
        btn_gen.setStyleSheet(
            "QPushButton { background:#920d2e; color:#ffffff;"
            "  border:none; border-radius:5px;"
            "  padding:7px 22px; font-size:13px; font-weight:700; }"
            "QPushButton:hover { background:#7a0b27; }"
            "QPushButton:pressed { background:#600820; }"
        )
        btn_gen.clicked.connect(self.generate_document)
        br.addWidget(btn_gen)

        root.addWidget(_bar)

    @staticmethod
    def _mk_sum_lbl(bold=False, large=False) -> QLabel:
        lbl = QLabel("\u2014")
        lbl.setStyleSheet(
            f"font-size:{'13' if large else '11'}px;"
            f"font-weight:{'700' if bold else '400'};"
            f"color:{'#920d2e' if bold else '#1a1a2e'};")
        return lbl

    def _upload_picture(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Upload Customer Logo", "",
            "Images (*.png *.jpg *.jpeg *.bmp *.gif *.svg *.webp)")
        if path:
            self._picture_label.set_image(path)

    # ─── Table building ───────────────────────────────────────────────────────

    def _snapshot_preset_states(self) -> dict[str, tuple[bool, int]]:
        snap = {}
        for row, key in self._row_keys.items():
            if key is None:
                continue
            cb = self._checkboxes.get(row)
            sp = self._qty_spins.get(row)
            snap[key] = (cb.isChecked() if cb else True, sp.value() if sp else 1)
        return snap

    def _rebuild_table(self, keep_preset_states: bool = False):
        saved = self._snapshot_preset_states() if keep_preset_states else {}

        self._block_update = True
        self._table.setRowCount(0)
        self._checkboxes.clear(); self._qty_spins.clear()
        self._row_keys.clear();   self._row_custom_idx.clear()
        self._product_widgets.clear()

        level    = self._level_cb.currentText()
        defaults = _LEVEL_DEFAULTS.get(level, _LEVEL_DEFAULTS["Basic"])

        for item in _ITEMS:
            key = item["key"]
            if keep_preset_states and key in saved:
                checked, qty = saved[key]
            else:
                checked, qty = defaults.get(key, (False, 1))
            price = _COMMISSION_PRICE.get(level, item["price"]) if key == "commission" else item["price"]
            self._insert_preset_row(item, checked, qty, price)

        for ci, citem in enumerate(self._custom_items):
            self._insert_custom_row(ci, citem)

        self._block_update = False
        self._recalc_totals()
        self._apply_tab_filter()

    def _insert_preset_row(self, item: dict, checked: bool, qty: int, price: float):
        logic = item["logic"]
        row   = self._table.rowCount()
        self._table.insertRow(row)
        self._row_keys[row] = item["key"]
        bg, fg = _LOGIC_STYLE.get(logic, ("#ffffff", "#333333"))
        bgc    = QColor(bg)

        is_locked = logic in ("Always", "Always when carts used")
        cb = QCheckBox()
        cb.setChecked(checked)
        cb.setStyleSheet(_CB_LOCKED if is_locked else _CB_ENABLED)
        cb.setEnabled(not is_locked)
        cb.setToolTip({
            "Always":                 "Always included",
            "Strongly Encouraged":    "Strongly recommended \u2014 uncheck to exclude",
            "Option 1":               "Mutually exclusive with the Uline Shelf cart (Option 2)",
            "Option 2":               "Mutually exclusive with the Baseplate cart (Option 1)",
            "Always when carts used": "Auto-included when any staging cart is selected",
            "Optional":               "Optional \u2014 check to include",
        }.get(logic, ""))
        wrap = QWidget(); wl = QHBoxLayout(wrap)
        wl.setContentsMargins(4, 0, 0, 0); wl.addWidget(cb); wl.addStretch()
        # Tint the cell widget background to match the row color
        wrap.setStyleSheet(f"QWidget {{ background:{bg}; border:none; }}")
        self._table.setCellWidget(row, _CC, wrap)
        self._checkboxes[row] = cb
        cb.stateChanged.connect(self._on_cb_changed)

        # ── Product + Description two-line cell widget ───────────────────
        prod_w = QWidget()
        prod_w.setStyleSheet(f"QWidget {{ background:{bg}; border:none; }}")
        pw_lay = QVBoxLayout(prod_w)
        pw_lay.setContentsMargins(6, 4, 6, 4); pw_lay.setSpacing(1)

        part_str = str(item.get("part", "")).strip()
        name_txt = f"{item['product']} - {part_str}" if part_str else item["product"]
        name_lbl = QLabel(name_txt)
        name_lbl.setWordWrap(True)
        name_lbl.setFont(QFont("Aptos", 12, QFont.Weight.DemiBold if is_locked else QFont.Weight.Medium))
        name_lbl.setObjectName("prod_name")
        pw_lay.addWidget(name_lbl)

        desc_text = item.get("desc", "")
        if desc_text:
            trunc = (desc_text[:130] + "…") if len(desc_text) > 130 else desc_text
            desc_lbl = QLabel(trunc)
            desc_lbl.setWordWrap(True)
            desc_lbl.setFont(QFont("Aptos", 11))
            desc_lbl.setStyleSheet("color:#6a5060;")
            desc_lbl.setObjectName("prod_desc")
            pw_lay.addWidget(desc_lbl)

        self._table.setCellWidget(row, _CPROD, prod_w)
        self._product_widgets[row] = prod_w
        self._apply_prod_colors(prod_w, bg, checked, is_locked)

        cat_txt = "Hardware" if item["cat"] == "hardware" else "Service"
        cat_fg  = "#1a6a3a" if item["cat"] == "service" else "#1a3a6e"
        ci_item = QTableWidgetItem(cat_txt)
        ci_item.setBackground(bgc); ci_item.setForeground(QColor(cat_fg))
        ci_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
        ci_item.setFont(QFont("Aptos", 9))
        self._table.setItem(row, _CCAT, ci_item)

        li = QTableWidgetItem(logic)
        li.setBackground(bgc); li.setForeground(QColor(fg))
        li.setFont(QFont("Aptos", 9))
        self._table.setItem(row, _CLOGIC, li)

        # ── +/− quantity widget ───────────────────────────────────────────
        sp = _SpinWidget(value=qty)
        sp.setEnabled(checked)
        sp.setStyleSheet(f"QWidget {{ background:{bg}; border:none; }}")
        sp.valueChanged.connect(self._recalc_totals)
        self._table.setCellWidget(row, _CQTY, sp)
        self._qty_spins[row] = sp

        pri = QTableWidgetItem(f"${price:,.2f}")
        pri.setBackground(bgc); pri.setFont(QFont("Aptos", 10))
        pri.setTextAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
        self._table.setItem(row, _CPRICE, pri)

        tot = QTableWidgetItem("")
        tot.setBackground(QColor("#f4f6fa")); tot.setForeground(QColor("#1a3a6e"))
        tot.setFont(QFont("Aptos", 10))
        tot.setTextAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
        self._table.setItem(row, _CTOT, tot)

        self._table.setRowHeight(row, 72 if desc_text else 42)
        self._dim_row(row, checked)

    def _insert_custom_row(self, ci: int, citem: dict):
        row = self._table.rowCount()
        self._table.insertRow(row)
        self._row_keys[row] = None
        self._row_custom_idx[row] = ci
        bgc = QColor("#fffef8")

        cb = QCheckBox(); cb.setChecked(citem.get("checked", True))
        cb.setStyleSheet(_CB_ENABLED)
        wrap = QWidget(); wl = QHBoxLayout(wrap)
        wl.setContentsMargins(4, 0, 0, 0); wl.addWidget(cb); wl.addStretch()
        wrap.setStyleSheet("QWidget { background:#fffef8; border:none; }")
        self._table.setCellWidget(row, _CC, wrap)
        self._checkboxes[row] = cb
        cb.stateChanged.connect(self._on_cb_changed)

        _prod = citem.get("product", ""); _part = citem.get("part", "").strip()
        _disp = f"{_prod} - {_part}" if _part else _prod
        for col, text in [(_CPROD, _disp)]:
            it = QTableWidgetItem(text); it.setBackground(bgc)
            self._table.setItem(row, col, it)

        li = QTableWidgetItem("Custom"); li.setBackground(bgc)
        li.setForeground(QColor("#555577")); li.setFont(QFont("Aptos", 9))
        self._table.setItem(row, _CLOGIC, li)   # Option column

        cat_cb = QComboBox(); cat_cb.addItems(["Hardware", "Service"])
        cat_cb.setCurrentText("Hardware" if citem.get("cat", "hardware") == "hardware" else "Service")
        cat_cb.setStyleSheet(
            "QComboBox { font-size:10px; border:1px solid #d0d4de; border-radius:3px; }"
            "QComboBox QAbstractItemView { background:#ffffff; color:#1a1a2e;"
            "  selection-background-color:#f5d0da; selection-color:#920d2e; }")
        cat_cb.currentTextChanged.connect(lambda _: self._recalc_totals())
        self._table.setCellWidget(row, _CCAT, cat_cb)

        sp = _SpinWidget(value=int(citem.get("qty", 1)))
        sp.setStyleSheet("QWidget { background:#fffef8; border:none; }")
        sp.valueChanged.connect(self._recalc_totals)
        self._table.setCellWidget(row, _CQTY, sp); self._qty_spins[row] = sp

        pri = QTableWidgetItem(f"${float(citem.get('price', 0)):,.2f}")
        pri.setBackground(bgc)
        pri.setTextAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
        self._table.setItem(row, _CPRICE, pri)

        tot = QTableWidgetItem("")
        tot.setBackground(QColor("#f4f6fa")); tot.setForeground(QColor("#1a3a6e"))
        tot.setTextAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
        self._table.setItem(row, _CTOT, tot)
        self._table.setRowHeight(row, 28)

    @staticmethod
    def _apply_prod_colors(prod_w: QWidget, bg: str, active: bool, is_locked: bool):
        """Update the two-line product cell widget's label colours."""
        name_fg = "#888888" if is_locked else ("#1a0509" if active else "#aaaaaa")
        desc_fg = "#aaaaaa" if is_locked else ("#778899" if active else "#cccccc")
        for child in prod_w.findChildren(QLabel):
            if child.objectName() == "prod_name":
                child.setStyleSheet(
                    f"QLabel {{ font-size:13px; color:{name_fg};"
                    f"  background:{bg}; border:none; }}")
            elif child.objectName() == "prod_desc":
                child.setStyleSheet(
                    f"QLabel {{ font-size:11px; color:{desc_fg};"
                    f"  background:{bg}; border:none; }}")

    # ─── Tab filter ───────────────────────────────────────────────────────────

    def _get_row_cat(self, row: int) -> str:
        key = self._row_keys.get(row)
        if key is not None:
            return next((x["cat"] for x in _ITEMS if x["key"] == key), "hardware")
        cat_w = self._table.cellWidget(row, _CCAT)
        return "service" if (cat_w and cat_w.currentText() == "Service") else "hardware"

    def _apply_tab_filter(self):
        for row in range(self._table.rowCount()):
            self._table.setRowHidden(row, self._get_row_cat(row) != self._active_tab)

    def _set_bom_tab(self, cat: str):
        self._active_tab = cat
        _on  = ("QPushButton { background:#920d2e; color:#ffffff; border:none;"
                "  border-radius:5px 5px 0 0; padding:6px 22px;"
                "  font-size:12px; font-weight:700; }")
        _off = ("QPushButton { background:#f9f0f2; color:#920d2e;"
                "  border:1px solid #d6c0c5; border-bottom:none;"
                "  border-radius:5px 5px 0 0; padding:6px 22px;"
                "  font-size:12px; font-weight:600; }"
                "QPushButton:hover { background:#fde8ed; }")
        self._tab_hw_btn.setStyleSheet(_on  if cat == "hardware" else _off)
        self._tab_svc_btn.setStyleSheet(_on  if cat == "service"  else _off)
        self._apply_tab_filter()

    def _dim_row(self, row: int, active: bool):
        key       = self._row_keys.get(row)
        item      = next((x for x in _ITEMS if x["key"] == key), None)
        logic     = item["logic"] if item else "Optional"
        is_locked = logic in ("Always", "Always when carts used")

        # Only visually dim rows whose checkbox is locked (can't click).
        # Rows that are simply unchecked but still clickable stay at full colour.
        visual_active = active if is_locked else True

        prod_w = self._product_widgets.get(row)
        if prod_w:
            bg, _ = _LOGIC_STYLE.get(logic, ("#ffffff", "#333333"))
            self._apply_prod_colors(prod_w, bg, visual_active, is_locked)

        alpha = 255 if visual_active else 100
        for col in (_CCAT, _CLOGIC, _CPRICE):
            it = self._table.item(row, col)
            if it:
                c = it.foreground().color(); c.setAlpha(alpha)
                it.setForeground(QBrush(c))

    # ─── Event handlers ───────────────────────────────────────────────────────

    def _on_level_changed(self, _lvl: str):
        self._rebuild_table(keep_preset_states=False)

    def _on_cb_changed(self, _state: int):
        if self._block_update:
            return
        sender = self.sender()
        row = next((r for r, cb in self._checkboxes.items() if cb is sender), None)
        if row is None:
            return
        key     = self._row_keys.get(row)
        checked = sender.isChecked()

        if key in ("cart_base", "cart_uline") and checked:
            other = "cart_uline" if key == "cart_base" else "cart_base"
            for r, k in self._row_keys.items():
                if k == other:
                    self._checkboxes[r].blockSignals(True)
                    self._checkboxes[r].setChecked(False)
                    self._checkboxes[r].blockSignals(False)
                    sp = self._qty_spins.get(r)
                    if sp: sp.setEnabled(False)
                    self._dim_row(r, False)
                    break
            # Ensure cart has a qty of at least 1 so engplate sync works
            sp_cart = self._qty_spins.get(row)
            if sp_cart and sp_cart.value() == 0:
                sp_cart.setValue(1)

        self._sync_engplate()
        sp = self._qty_spins.get(row)
        if sp: sp.setEnabled(checked)
        self._dim_row(row, checked)
        self._recalc_totals()

    def _sync_engplate(self):
        cart_qty = 0
        for r, k in self._row_keys.items():
            if k in ("cart_base", "cart_uline"):
                cb = self._checkboxes.get(r)
                if cb and cb.isChecked():
                    sp = self._qty_spins.get(r)
                    cart_qty += sp.value() if sp else 0
        cart_active = cart_qty > 0
        for r, k in self._row_keys.items():
            if k == "engplate":
                cb = self._checkboxes.get(r)
                sp = self._qty_spins.get(r)
                if cb:
                    cb.blockSignals(True); cb.setChecked(cart_active); cb.blockSignals(False)
                if sp:
                    sp.blockSignals(True); sp.setValue(cart_qty if cart_active else 0)
                    sp.setEnabled(cart_active); sp.blockSignals(False)
                self._dim_row(r, cart_active)
                break

    # ─── Totals ───────────────────────────────────────────────────────────────

    def _recalc_totals(self):
        if self._block_update:
            return
        hw = svc = 0.0
        for row in range(self._table.rowCount()):
            cb = self._checkboxes.get(row)
            if not (cb and cb.isChecked()):
                it = self._table.item(row, _CTOT)
                if it: it.setText("\u2014")
                continue
            sp  = self._qty_spins.get(row)
            qty = sp.value() if sp else 0
            pri = self._table.item(row, _CPRICE)
            try:
                price = float(pri.text().replace("$", "").replace(",", "")) if pri else 0.0
            except Exception:
                price = 0.0
            line = qty * price
            it = self._table.item(row, _CTOT)
            if it: it.setText(f"${line:,.2f}")

            key = self._row_keys.get(row)
            if key is not None:
                cat = next((x["cat"] for x in _ITEMS if x["key"] == key), "hardware")
            else:
                cat_w = self._table.cellWidget(row, _CCAT)
                cat   = "service" if (cat_w and cat_w.currentText() == "Service") else "hardware"
            (svc if cat == "service" else hw).__class__   # no-op just for readability
            if cat == "service":  svc += line
            else:                 hw  += line

        grand = hw + svc
        self._lbl_hw.setText(f"${hw:,.2f}")
        self._lbl_svc.setText(f"${svc:,.2f}")
        self._lbl_tot.setText(f"${grand:,.2f}")
        self._inv["HDW-1ST (50%)"].setText(f"${hw  * 0.50:,.2f}")
        self._inv["HDW-2ND (40%)"].setText(f"${hw  * 0.40:,.2f}")
        self._inv["HDW-3RD (10%)"].setText(f"${hw  * 0.10:,.2f}")
        self._inv["SRV-1ST (50%)"].setText(f"${svc * 0.50:,.2f}")
        self._inv["SRV-2ND (50%)"].setText(f"${svc * 0.50:,.2f}")

    # ─── Custom items ─────────────────────────────────────────────────────────

    def _add_custom(self, default_cat: str):
        dlg = _CustomItemDialog(default_cat, self)
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return
        self._custom_items.append(dlg.result())
        self._rebuild_table(keep_preset_states=True)

    def _remove_custom(self):
        rows = sorted({i.row() for i in self._table.selectedIndexes()}, reverse=True)
        to_del = sorted({self._row_custom_idx[r] for r in rows
                         if self._row_keys.get(r) is None and r in self._row_custom_idx},
                        reverse=True)
        if not to_del:
            QMessageBox.information(self, "Remove Item",
                "Select a custom item row to remove.\n(Preset items cannot be removed.)")
            return
        for ci in to_del:
            if 0 <= ci < len(self._custom_items):
                del self._custom_items[ci]
        self._rebuild_table(keep_preset_states=True)

    # ─── Serialization ────────────────────────────────────────────────────────

    def get_data(self) -> dict:
        rows = []
        for row in range(self._table.rowCount()):
            cb = self._checkboxes.get(row)
            sp = self._qty_spins.get(row)
            rows.append({
                "key":     self._row_keys.get(row),
                "checked": cb.isChecked() if cb else False,
                "qty":     sp.value()     if sp else 1,
            })
        return {
            "project":         self.edit_project.text(),
            "customer":        self.edit_customer.text(),
            "location":        self.edit_location.text(),
            "proposal":        self.edit_proposal.text(),
            "contact_info":    self._contact_w.toPlainText(),
            "customer_picture":getattr(self._picture_label, "image_path", None) or "",
            "robot":           self._robot_cb.currentText(),
            "level":           self._level_cb.currentText(),
            "rows":            rows,
            "custom":          copy.deepcopy(self._custom_items),
            "version_major":   self._version_major_spin.value(),
            "version_minor":   self._version_minor_spin.value(),
            "version_history": [self._version_history.item(i).text()
                                for i in range(self._version_history.count())],
        }

    def restore_data(self, d: dict):
        self.edit_project.setText(d.get("project",  ""))
        self.edit_customer.setText(d.get("customer", ""))
        self.edit_location.setText(d.get("location", ""))
        self.edit_proposal.setText(d.get("proposal", ""))
        self._contact_w.setPlainText(d.get("contact_info", ""))

        pic = d.get("customer_picture", "")
        if pic and os.path.exists(pic):
            self._picture_label.set_image(pic)

        idx = self._robot_cb.findText(d.get("robot", "OTTO 100"))
        if idx >= 0: self._robot_cb.setCurrentIndex(idx)

        self._custom_items = copy.deepcopy(d.get("custom", []))
        self._level_cb.blockSignals(True)
        idx = self._level_cb.findText(d.get("level", "Basic"))
        if idx >= 0: self._level_cb.setCurrentIndex(idx)
        self._level_cb.blockSignals(False)

        self._rebuild_table(keep_preset_states=False)

        for row, rdata in enumerate(d.get("rows", [])):
            if row >= self._table.rowCount():
                break
            cb = self._checkboxes.get(row)
            sp = self._qty_spins.get(row)
            if cb and cb.isEnabled():
                cb.blockSignals(True); cb.setChecked(rdata.get("checked", True))
                cb.blockSignals(False); self._dim_row(row, rdata.get("checked", True))
            if sp:
                sp.blockSignals(True); sp.setValue(rdata.get("qty", 1))
                sp.setEnabled(cb.isChecked() if cb else True); sp.blockSignals(False)

        self._recalc_totals()
        self._version_major_spin.setValue(int(d.get("version_major", 1)))
        self._version_minor_spin.setValue(int(d.get("version_minor", 0)))
        self._update_version_label()
        self._version_history.clear()
        for _entry in d.get("version_history", []):
            self._version_history.addItem(_entry)

    # ── Version control ───────────────────────────────────────────────────────

    def _version_str(self) -> str:
        return f"V{self._version_major_spin.value()}.{self._version_minor_spin.value()}"

    def _update_version_label(self):
        self._version_badge.setText(self._version_str())

    def _toggle_history_panel(self):
        self._history_collapsed = not self._history_collapsed
        self._history_panel.setVisible(not self._history_collapsed)
        self._btn_collapse_hist.setText(
            "▶  Change History" if self._history_collapsed
            else "▼  Change History")

    def _add_version_entry(self):
        from PyQt6.QtWidgets import QVBoxLayout as _QVL2
        v = self._version_str()
        dlg = QDialog(self)
        dlg.setWindowTitle(f"Add Change Entry — {v}")
        dlg.setFixedWidth(420)
        dlg.setStyleSheet(
            "QDialog { background:#ffffff; }"
            "QLabel  { color:#1a0509; font-size:12px; }"
            "QLineEdit { border:1px solid #d6c0c5; border-radius:4px;"
            "  padding:6px 10px; font-size:12px; color:#1a0509; background:#fafbfd; }"
            "QLineEdit:focus { border-color:#920d2e; }"
            "QPushButton { background:#f4f6fa; color:#1a0509;"
            "  border:1px solid #d6c0c5; border-radius:4px;"
            "  padding:6px 18px; min-width:72px; font-size:12px; }"
            "QPushButton:hover { background:#e0e6f0; }")
        vbox = _QVL2(dlg)
        vbox.setSpacing(10); vbox.setContentsMargins(16, 16, 16, 16)
        vbox.addWidget(QLabel(f"<b>Version:</b>  {v}"))
        desc = QLineEdit(); desc.setPlaceholderText("Describe this change…")
        vbox.addWidget(desc)
        btns = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok |
            QDialogButtonBox.StandardButton.Cancel)
        btns.accepted.connect(dlg.accept); btns.rejected.connect(dlg.reject)
        vbox.addWidget(btns)
        if dlg.exec() == QDialog.DialogCode.Accepted and desc.text().strip():
            today = _date.today().strftime("%m/%d/%Y")
            self._version_history.addItem(
                f"{v}  —  {desc.text().strip()}  ({today})")

    def _remove_version_entry(self):
        row = self._version_history.currentRow()
        if row >= 0:
            self._version_history.takeItem(row)

    # ─── File I/O ─────────────────────────────────────────────────────────────

    def save_project(self):
        desktop  = os.path.join(os.path.expanduser("~"), "Desktop")
        proposal = self.edit_proposal.text().strip()
        _ver     = f" {self._version_str()}"
        path, _  = QFileDialog.getSaveFileName(
            self, "Save OTTO Project",
            os.path.join(desktop, f"{proposal or 'otto'}{_ver}.mcmxo"),
            "OTTO Project (*.mcmxo)")
        if not path: return
        try:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(self.get_data(), f, indent=2, ensure_ascii=False)
            QMessageBox.information(self, "Saved", f"\u2714  Project saved:\n{path}")
        except Exception as e:
            QMessageBox.critical(self, "Save Failed", str(e))

    def load_project(self):
        desktop = os.path.join(os.path.expanduser("~"), "Desktop")
        path, _ = QFileDialog.getOpenFileName(
            self, "Load OTTO Project", desktop, "OTTO Project (*.mcmxo)")
        if not path: return
        try:
            with open(path, "r", encoding="utf-8") as f:
                self.restore_data(json.load(f))
        except Exception as e:
            QMessageBox.critical(self, "Load Failed", str(e))

    # ─── Document generation ──────────────────────────────────────────────────

    def generate_document(self):
        import re as _re
        desktop   = os.path.join(os.path.expanduser("~"), "Desktop")
        proposal  = self.edit_proposal.text().strip()
        safe      = _re.sub(r'[\\/:*?"<>|]', "-", proposal) if proposal else "OTTO"
        _ver      = f" {self._version_str()}"
        _date_str = _date.today().strftime("%Y%m%d")
        out, _    = QFileDialog.getSaveFileName(
            self, "Save OTTO Document",
            os.path.join(desktop, f"{safe}-{_date_str}{_ver}.docx"),
            "Word Documents (*.docx)")
        if not out: return
        try:
            generate_otto_doc(self.get_data(), out)
            mcmxo = os.path.splitext(out)[0] + ".mcmxo"
            try:
                with open(mcmxo, "w", encoding="utf-8") as f:
                    json.dump(self.get_data(), f, indent=2, ensure_ascii=False)
            except Exception:
                pass
            QMessageBox.information(
                self, "Generated",
                f"\u2714  Document:\n{out}\n\nProject file:\n{mcmxo}")
        except Exception as e:
            QMessageBox.critical(self, "Generation Failed", str(e))


# ─── Custom Item Dialog ───────────────────────────────────────────────────────
class _CustomItemDialog(QDialog):
    def __init__(self, default_cat: str = "hardware", parent=None):
        super().__init__(parent)
        self.setWindowTitle("Add Custom BOM Item")
        self.setMinimumWidth(380)
        self.setStyleSheet("QDialog { background:#ffffff; }")
        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 12, 16, 10); layout.setSpacing(8)
        form = QFormLayout(); form.setSpacing(6)
        _F = ("QLineEdit { background:#fff; border:1px solid #d0d4de;"
              "  border-radius:4px; padding:3px 8px; font-size:11px; }")
        _L = "font-size:10px; color:#555577;"
        self._prod  = QLineEdit(); self._prod.setStyleSheet(_F)
        self._part  = QLineEdit(); self._part.setStyleSheet(_F)
        self._desc  = QLineEdit(); self._desc.setStyleSheet(_F)
        self._cat   = QComboBox(); self._cat.addItems(["Hardware", "Service"])
        self._cat.setCurrentText("Hardware" if default_cat == "hardware" else "Service")
        self._cat.setStyleSheet(_COMBO)
        self._qty   = QSpinBox(); self._qty.setRange(0, 9999); self._qty.setValue(1)
        self._price = QDoubleSpinBox()
        self._price.setRange(0, 9_999_999); self._price.setDecimals(2)
        self._price.setPrefix("$"); self._price.setValue(0.0)
        for lbl, w in [("Product Name:", self._prod), ("Part #:", self._part),
                        ("Description:", self._desc), ("Category:", self._cat),
                        ("Qty:", self._qty), ("Unit Price:", self._price)]:
            l = QLabel(lbl); l.setStyleSheet(_L); form.addRow(l, w)
        layout.addLayout(form)
        btns = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok |
                                QDialogButtonBox.StandardButton.Cancel)
        btns.accepted.connect(self.accept); btns.rejected.connect(self.reject)
        layout.addWidget(btns)

    def result(self) -> dict:
        return {
            "product": self._prod.text().strip(), "part": self._part.text().strip(),
            "desc": self._desc.text().strip(),
            "cat": "hardware" if self._cat.currentText() == "Hardware" else "service",
            "qty": self._qty.value(), "price": self._price.value(), "checked": True,
        }


# ─── Ordinal date formatter ───────────────────────────────────────────────────
def _ordinal_date(d: _date) -> str:
    day = d.day
    suffix = {1:"st",2:"nd",3:"rd"}.get(day if day < 20 else day % 10, "th")
    return d.strftime(f"%B {day}{suffix}, %Y")


# ─── Document Generator ───────────────────────────────────────────────────────
def generate_otto_doc(data: dict, output_path: str):
    """Generate OTTO Word document via docxtpl then python-docx BOM tables."""
    import copy
    from docxtpl import DocxTemplate, InlineImage
    from docx.shared import Mm
    from docx import Document
    from docx.oxml.ns import qn
    from lxml import etree

    import sys as _sys
    # Use sys._MEIPASS when running from a PyInstaller bundle so the
    # template is found inside the extracted _internal folder.
    _here = getattr(_sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))
    candidates = [
        os.path.join(_here, "otto_template.docx"),
    ]
    template = next((c for c in candidates if os.path.exists(c)), None)
    if not template:
        raise FileNotFoundError(
            "OTTO template not found. Expected 'otto_template.docx' in the scripts folder.")

    proposal  = data.get("proposal", "")
    customer  = data.get("customer", "")
    location  = data.get("location", "")
    today_str = _ordinal_date(_date.today())

    contact_lines = (data.get("contact_info", "") or "").splitlines()
    def _cl(i): return contact_lines[i].strip() if i < len(contact_lines) else ""
    pres_name = _cl(0); pres_email = _cl(2)
    acct_name = _cl(4); acct_email = _cl(6)

    pic_path = data.get("customer_picture", "")
    tpl = DocxTemplate(template)
    customer_image = InlineImage(tpl, pic_path, width=Mm(50)) \
        if pic_path and os.path.exists(pic_path) else ""

    tpl.render({
        "customer_name":     customer,
        "customer_location": location,
        "proposal_number":   proposal,
        "today_date":        today_str,
        "pres_name":         pres_name,
        "pres_phone":        "",
        "pres_email":        pres_email,
        "acct_name":         acct_name,
        "acct_phone":        "",
        "acct_email":        acct_email,
        "customer_picture":  customer_image,
    })
    tpl.save(output_path)

    # ── Re-open with python-docx to populate BOM & invoice tables ────────
    doc = Document(output_path)

    # ── Build active BOM lists ────────────────────────────────────────────
    level  = data.get("level", "Basic")
    rows   = data.get("rows",  [])
    custom = data.get("custom", [])

    hw_rows:  list[tuple[str, str, int, float]] = []
    svc_rows: list[tuple[str, str, int, float]] = []

    for rdata in rows:
        if not rdata.get("checked"): continue
        key  = rdata.get("key")
        if not key: continue
        item = next((x for x in _ITEMS if x["key"] == key), None)
        if not item: continue
        qty  = rdata.get("qty", 1)
        if qty <= 0: continue
        price = _COMMISSION_PRICE.get(level, item["price"]) if key == "commission" else item["price"]
        entry = (item["product"], str(item["part"]) if item["part"] else "", qty, price)
        (svc_rows if item["cat"] == "service" else hw_rows).append(entry)

    preset_count = len(_ITEMS)
    for ci, citem in enumerate(custom):
        ri = preset_count + ci
        if ri < len(rows) and not rows[ri].get("checked", True): continue
        qty   = int(citem.get("qty", 1)); price = float(citem.get("price", 0))
        if qty <= 0: continue
        entry = (citem.get("product",""), citem.get("part",""), qty, price)
        (svc_rows if citem.get("cat") == "service" else hw_rows).append(entry)

    hw_total  = sum(q * p for _, _, q, p in hw_rows)
    svc_total = sum(q * p for _, _, q, p in svc_rows)
    grand     = hw_total + svc_total

    def _cell_set(cell, text: str):
        tc    = cell._tc
        paras = tc.findall(qn("w:p"))
        for p in paras[1:]: tc.remove(p)
        p = paras[0] if paras else etree.SubElement(tc, qn("w:p"))
        for r in p.findall(qn("w:r"))[1:]: p.remove(r)
        runs = p.findall(qn("w:r"))
        if runs:
            t_node = runs[0].find(qn("w:t"))
            if t_node is None:
                t_node = etree.SubElement(runs[0], qn("w:t"))
            t_node.text = text
            t_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        else:
            rn = etree.SubElement(p, qn("w:r"))
            tn = etree.SubElement(rn, qn("w:t"))
            tn.text = text
            tn.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    def _rebuild_bom_table(tbl_idx: int, bom_rows: list):
        tbl  = doc.tables[tbl_idx]
        tmpl = copy.deepcopy(tbl.rows[1]._element) if len(tbl.rows) > 1 \
               else copy.deepcopy(tbl.rows[0]._element)
        while len(tbl.rows) > 1:
            tbl._element.remove(tbl.rows[-1]._element)
        entries = bom_rows or [("(none)", "", 0, 0.0)]
        for product, part, qty, _p in entries:
            new_tr = copy.deepcopy(tmpl)
            tbl._element.append(new_tr)
            new_row = tbl.rows[-1]
            part_str = f"{part} \u2013 {product}" if part else product
            _cell_set(new_row.cells[0], part_str)
            _cell_set(new_row.cells[1], str(qty))

    _rebuild_bom_table(0, hw_rows)
    _rebuild_bom_table(1, svc_rows)


    if len(doc.tables) > 2:
        inv_tbl  = doc.tables[2]
        inv_data = [
            (f"{proposal}-HDW-1ST", hw_total  * 0.50),
            (f"{proposal}-HDW-2ND", hw_total  * 0.40),
            (f"{proposal}-HDW-3RD", hw_total  * 0.10),
            (f"{proposal}-SRV-1ST", svc_total * 0.50),
            (f"{proposal}-SRV-2ND", svc_total * 0.50),
        ]
        for i, (pn, amt) in enumerate(inv_data):
            ri = i + 1
            if ri < len(inv_tbl.rows):
                _cell_set(inv_tbl.rows[ri].cells[0], pn)
                _cell_set(inv_tbl.rows[ri].cells[1], f"${amt:,.2f}")
        if len(inv_tbl.rows) >= 7:
            total_str = f"Total Price: ${grand:,.2f}"
            _cell_set(inv_tbl.rows[6].cells[0], total_str)
            _cell_set(inv_tbl.rows[6].cells[1], total_str)

    # ── Refresh all fields (TOC page numbers, cross-refs, etc.) on open ──
    # Mark every field-begin character dirty so Word recalculates each
    # field the first time the document is opened.
    for _fc in doc.element.iter(qn("w:fldChar")):
        if _fc.get(qn("w:fldCharType")) == "begin":
            _fc.set(qn("w:dirty"), "true")

    # Set updateFields in document settings so Word triggers the refresh
    # automatically on open (no manual right-click required).
    _stg = doc.settings.element
    for _uf in _stg.findall(qn("w:updateFields")):
        _stg.remove(_uf)
    _uf = etree.SubElement(_stg, qn("w:updateFields"))
    _uf.set(qn("w:val"), "true")

    doc.save(output_path)
