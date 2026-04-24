"""
thermal_widget.py — Thermal Imaging Scheduler & Cost Estimator
"""
from __future__ import annotations
import math, json, os
from datetime import date as _date

from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit, QPushButton,
    QTableWidget, QTableWidgetItem, QHeaderView, QScrollArea, QFrame,
    QSizePolicy, QComboBox, QCheckBox, QAbstractItemView, QFileDialog,
    QMessageBox, QDoubleSpinBox, QFormLayout, QTextEdit,
)
from PyQt6.QtCore import Qt
from PyQt6.QtGui import QColor, QFont, QDoubleValidator, QIntValidator

# Reuse Card and style tokens from cost_estimator
from cost_estimator import _Card, _FIELD_STYLE, _CALC_STYLE, _BTN_STYLE, _TABLE_STYLE
from form_widgets import DragDropLabel

_RED    = "#920d2e"
_TEXT   = "#1a0509"
_SUBTEXT= "#3a3a5c"
_BORDER = "#d6c0c5"
_BG_HDR = "#f9f0f2"

_FORM_LBL = "QLabel { font-size:12px; font-weight:700; color:#1a0509; border:none; }"

_DAY_NAMES = ["Mon","Tue","Wed","Thu","Fri","Sat","Sun"]


def _n(text, default=0.0):
    try: return float((text or "").replace(",","").replace("$","").strip() or default)
    except: return default


def _ordinal_date(d) -> str:
    suffix = {1:"st",2:"nd",3:"rd"}.get(d.day % 10 if d.day not in (11,12,13) else 0, "th")
    return d.strftime(f"%B {d.day}{suffix}, %Y")


_ACTN_STYLE = (
    "QPushButton { background:#920d2e; color:#ffffff; border:none;"
    "  border-radius:5px; padding:6px 18px; font-size:12px; font-weight:700; }"
    "QPushButton:hover { background:#b01036; }"
    "QPushButton:disabled { background:#c8c8c8; color:#888888; }"
)
_SPIN_S = (
    "QPushButton { background:#f4f6fa; color:#3a3a5c; border:1px solid #dde1e7;"
    "  border-radius:4px; font-size:14px; font-weight:700; padding:0; }"
    "QPushButton:hover { background:#e8ecf5; }"
)
_CB_STYLE = (
    "QCheckBox { spacing:6px; font-size:11px; color:#1a0509; }"
    "QCheckBox::indicator { width:16px; height:16px; border-radius:3px; }"
    "QCheckBox::indicator:unchecked { border:2px solid #d6c0c5; background:#ffffff; }"
    "QCheckBox::indicator:unchecked:hover { border-color:#920d2e; }"
    "QCheckBox::indicator:checked { border:2px solid #920d2e; background:#920d2e; }"
)
_INFO_LBL = f"font-size:10px; color:{_SUBTEXT}; background:transparent; border:none;"

_ASE_TIERS = [
    ("Up to 50 images",    1_600.00),
    ("51–100 images",      2_666.67),
    ("101–150 images",     4_000.00),
    ("151–200 images",     5_333.33),
    ("201+ images (custom)", None),
]


class _ThermalContactWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        lay = QVBoxLayout(self); lay.setContentsMargins(0,0,0,0); lay.setSpacing(4)
        self._pres_name  = QLineEdit(); self._pres_name.setPlaceholderText("Presented By name…"); self._pres_name.setStyleSheet(_FIELD_STYLE)
        self._pres_email = QLineEdit(); self._pres_email.setPlaceholderText("Presented By email…"); self._pres_email.setStyleSheet(_FIELD_STYLE)
        self._acct_name  = QLineEdit(); self._acct_name.setPlaceholderText("Account Manager name…"); self._acct_name.setStyleSheet(_FIELD_STYLE)
        self._acct_email = QLineEdit(); self._acct_email.setPlaceholderText("Account Manager email…"); self._acct_email.setStyleSheet(_FIELD_STYLE)
        for w in (self._pres_name, self._pres_email, self._acct_name, self._acct_email):
            lay.addWidget(w)

    def toPlainText(self) -> str:
        return "\n".join([self._pres_name.text(), "", self._pres_email.text(), "",
                          self._acct_name.text(), "", self._acct_email.text()])

    def setPlainText(self, text: str):
        lines = text.splitlines()
        def _cl(i): return lines[i].strip() if i < len(lines) else ""
        self._pres_name.setText(_cl(0)); self._pres_email.setText(_cl(2))
        self._acct_name.setText(_cl(4)); self._acct_email.setText(_cl(6))


class ThermalImagingWidget(QWidget):
    """Thermal Imaging scheduler and cost estimator tab."""

    LABOR_RATE   = 125.0
    TRAVEL_RATE  = 100.0
    HOTEL_RATE   = 150.0
    MEAL_RATE    =  25.0
    MILE_RATE    =   0.72

    def __init__(self, parent=None, compact=False):
        super().__init__(parent)
        self._compact              = compact
        self._num_techs            = 1
        self._tech_rows         = []
        self._schedule          = []
        self._hotel_checks      = []
        self._travel_in_checks  = {}
        self._travel_out_checks = {}
        self._confirmed         = {}

        # Project info attributes
        self._ir_project  = None
        self._ir_proposal = None
        self._ir_customer = None
        self._ir_location = None
        self._ir_scope    = None  # removed from UI; kept for backward-compat save/restore
        self._ir_contact  = None
        self._ir_picture  = None

        # Version control
        self._version_major_spin = None
        self._version_minor_spin = None
        self._version_badge      = None
        self._version_history    = None
        self._history_collapsed  = True
        self._history_panel      = None
        self._btn_collapse_hist  = None

        self._setup_ui()

    # ── UI construction ───────────────────────────────────────────────────────

    def _setup_ui(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        if self._compact:
            # Compact mode: no scroll wrapper, no bottom bar
            cl = root
            cl.setSpacing(8)
        else:
            scroll = QScrollArea()
            scroll.setWidgetResizable(True)
            scroll.setStyleSheet("QScrollArea { border:none; background:transparent; }")

            content = QWidget()
            cl = QVBoxLayout(content)
            cl.setContentsMargins(20, 16, 20, 16)
            cl.setSpacing(12)

        # ── Collapsible Project Info panel ────────────────────────────────────
        if not self._compact:
            _tir_hdr = QWidget()
            _tir_hdr.setCursor(Qt.CursorShape.PointingHandCursor)
            _tir_hdr.setFixedHeight(32)
            _tir_hdr.setStyleSheet(
                "QWidget { background:#f9f0f2; border:1px solid #d6c0c5;"
                "  border-radius:6px 6px 0 0; }")
            _tir_row = QHBoxLayout(_tir_hdr)
            _tir_row.setContentsMargins(10, 0, 10, 0); _tir_row.setSpacing(6)
            _tir_lbl = QLabel("Project Information")
            _tir_lbl.setStyleSheet(
                f"font-size:12px; font-weight:700; color:{_RED};"
                "  background:transparent; border:none;")
            self._tir_toggle = QPushButton("▼")
            self._tir_toggle.setFixedSize(22, 22)
            self._tir_toggle.setStyleSheet(
                f"QPushButton {{ background:transparent; border:none;"
                f"  font-size:12px; color:{_RED}; }}"
                f"QPushButton:hover {{ color:#7a0b27; }}")
            _tir_row.addWidget(self._tir_toggle); _tir_row.addWidget(_tir_lbl)
            _tir_row.addStretch()

        info_panel = QWidget()
        info_panel.setObjectName("thermal_info")
        if self._compact:
            info_panel.setStyleSheet(
                "QWidget#thermal_info { background:#ffffff; border:1px solid #e8dde0; border-radius:8px; }")
        else:
            info_panel.setStyleSheet(
                "QWidget#thermal_info { background:#ffffff; border:1px solid #d6c0c5;"
                "  border-top:none; border-radius:0 0 6px 6px; }")
        ip_lay = QHBoxLayout(info_panel)
        ip_lay.setContentsMargins(16, 14, 16, 14); ip_lay.setSpacing(20)
        self._build_project_info(ip_lay)

        if self._compact:
            cl.addWidget(info_panel)
        else:
            _tir_wrap = QWidget()
            _tir_wrap.setStyleSheet("QWidget{background:transparent;border:none;}")
            _tir_vl = QVBoxLayout(_tir_wrap)
            _tir_vl.setContentsMargins(0, 0, 0, 0); _tir_vl.setSpacing(0)
            _tir_vl.addWidget(_tir_hdr)
            _tir_vl.addWidget(info_panel)
            cl.addWidget(_tir_wrap)

            def _tir_do_toggle():
                vis = not info_panel.isVisible()
                info_panel.setVisible(vis)
                self._tir_toggle.setText("▼" if vis else "►")
                _tir_hdr.setStyleSheet(
                    "QWidget { background:#f9f0f2; border:1px solid #d6c0c5;"
                    f"  border-radius:{'6px 6px 0 0' if vis else '6px'}; }}")
            self._tir_toggle.clicked.connect(_tir_do_toggle)
            _tir_hdr.mousePressEvent = lambda _e: _tir_do_toggle()

            # ── Version row ───────────────────────────────────────────────
            _VER_BTN = (
                "QPushButton { background:#f4f6fa; color:#3a3a5c; border:1px solid #c8cedd;"
                "  border-radius:4px; font-size:16px; font-weight:700; }"
                "QPushButton:hover { background:#e0e4ef; }"
                "QPushButton:pressed { background:#d0d8ef; }")
            _VER_VAL = (
                "QLabel { background:#ffffff; border:1px solid #d6c0c5;"
                "  border-radius:4px; color:#1a0509; font-size:12px;"
                "  font-weight:700; padding:4px 8px; }")

            class _VerSpin:
                """Lightweight +/− spinner matching the Quote Generator style."""
                def __init__(self_, lo, hi, initial, on_change):
                    self_._val = initial; self_._lo = lo; self_._hi = hi; self_._cb = on_change
                    self_.widget = QWidget()
                    _r = QHBoxLayout(self_.widget)
                    _r.setContentsMargins(0, 0, 0, 0); _r.setSpacing(3)
                    self_._minus = QPushButton("−"); self_._minus.setFixedSize(28, 28)
                    self_._minus.setStyleSheet(_VER_BTN); self_._minus.clicked.connect(self_._dec)
                    _r.addWidget(self_._minus)
                    self_._lbl = QLabel(str(initial)); self_._lbl.setFixedWidth(36)
                    self_._lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
                    self_._lbl.setStyleSheet(_VER_VAL)
                    _r.addWidget(self_._lbl)
                    self_._plus = QPushButton("+"); self_._plus.setFixedSize(28, 28)
                    self_._plus.setStyleSheet(_VER_BTN); self_._plus.clicked.connect(self_._inc)
                    _r.addWidget(self_._plus)
                def value(self_): return self_._val
                def setValue(self_, v):
                    self_._val = max(self_._lo, min(self_._hi, int(v)))
                    self_._lbl.setText(str(self_._val))
                def _inc(self_): self_.setValue(self_._val + 1); self_._cb()
                def _dec(self_): self_.setValue(self_._val - 1); self_._cb()

            _ver_row_w = QWidget()
            _ver_row_w.setFixedHeight(44)
            _ver_row_w.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Fixed)
            _vrl = QHBoxLayout(_ver_row_w)
            _vrl.setContentsMargins(0, 6, 0, 6); _vrl.setSpacing(0)

            _vlbl = QLabel("Version:")
            _vlbl.setStyleSheet(
                "font-size:12px; font-weight:700; color:#920d2e;"
                " background:transparent; border:none;")
            _vrl.addWidget(_vlbl); _vrl.addSpacing(10)

            self._version_major_spin = _VerSpin(1, 99, 1, self._update_version_label)
            _vrl.addWidget(self._version_major_spin.widget)

            _vdot = QLabel(".")
            _vdot.setStyleSheet(
                "font-size:16px; font-weight:700; color:#1a0509;"
                " background:transparent; border:none;")
            _vdot.setFixedWidth(10); _vdot.setAlignment(Qt.AlignmentFlag.AlignCenter)
            _vrl.addWidget(_vdot)

            self._version_minor_spin = _VerSpin(0, 99, 0, self._update_version_label)
            _vrl.addWidget(self._version_minor_spin.widget)

            _vrl.addSpacing(12)
            self._version_badge = QLabel("V1.0")
            self._version_badge.setStyleSheet(
                "QLabel { background:#9E1B32; color:#ffffff; font-size:11px;"
                "  font-weight:700; border-radius:4px; padding:3px 10px; }")
            _vrl.addWidget(self._version_badge)
            _vrl.addStretch()
            cl.addWidget(_ver_row_w)

            # ── Change History panel ──────────────────────────────────────
            from PyQt6.QtWidgets import QListWidget as _QLW
            self._hist_outer = QWidget()
            _hist_vbox = QVBoxLayout(self._hist_outer)
            _hist_vbox.setContentsMargins(0, 0, 0, 4); _hist_vbox.setSpacing(4)

            self._btn_collapse_hist = QPushButton("▶  Change History")
            self._btn_collapse_hist.setStyleSheet(
                "QPushButton { background:#ffffff; border:1px solid #d6c0c5;"
                "  border-radius:5px; padding:6px 14px; font-size:12px;"
                "  font-weight:600; color:#920d2e; text-align:left; }"
                "QPushButton:hover { background:#fdf0f3; }")
            self._btn_collapse_hist.clicked.connect(self._toggle_history_panel)
            _hist_vbox.addWidget(self._btn_collapse_hist)

            self._history_panel = QWidget()
            self._history_panel.setVisible(False)
            _hp_vbox = QVBoxLayout(self._history_panel)
            _hp_vbox.setContentsMargins(0, 2, 0, 2); _hp_vbox.setSpacing(4)

            self._version_history = _QLW()
            self._version_history.setMinimumHeight(80)
            self._version_history.setMaximumHeight(130)
            self._version_history.setStyleSheet(
                "QListWidget { background:#ffffff; border:1px solid #d6c0c5;"
                "  border-radius:4px; font-size:11px; color:#1a0509; }"
                "QListWidget::item { padding:4px 8px; }"
                "QListWidget::item:selected { background:#f5d0da; color:#920d2e; }")
            _hp_vbox.addWidget(self._version_history)

            _hbr = QWidget(); _hbrl = QHBoxLayout(_hbr)
            _hbrl.setContentsMargins(0, 0, 0, 0); _hbrl.setSpacing(8)
            _hadd = QPushButton("+ Add Entry")
            _hadd.setStyleSheet(
                "QPushButton { background:#f4f6fa; color:#3a3a5c;"
                "  border:1px solid #dde1e7; border-radius:4px;"
                "  padding:4px 12px; font-size:11px; }"
                "QPushButton:hover { background:#e8ecf5; border-color:#b0b8d0; }")
            _hadd.clicked.connect(self._add_version_entry)
            _hbrl.addWidget(_hadd)
            _hdel = QPushButton("Remove")
            _hdel.setStyleSheet(
                "QPushButton { background:transparent; color:#e05252;"
                "  border:1px solid #e05252; border-radius:4px;"
                "  padding:4px 12px; font-size:11px; }"
                "QPushButton:hover { background:#fdf2f2; }")
            _hdel.clicked.connect(self._remove_version_entry)
            _hbrl.addWidget(_hdel); _hbrl.addStretch()
            _hp_vbox.addWidget(_hbr)

            _hist_vbox.addWidget(self._history_panel)
            cl.addWidget(self._hist_outer)

        # ── Config card ───────────────────────────────────────────────────
        self._cfg_card = _Card("Technician Configuration")
        self._build_config(self._cfg_card._body_layout)
        cl.addWidget(self._cfg_card)

        # ── Outlook card (hidden until generated) ─────────────────────────
        self._outlook_card = _Card("Work Week Overview")
        self._build_outlook_card(self._outlook_card._body_layout)
        self._outlook_card.setVisible(False)
        cl.addWidget(self._outlook_card)

        # ── Summary card (hidden until confirmed) ─────────────────────────
        self._summary_card = _Card("Schedule Summary & Costs")
        self._build_summary_card(self._summary_card._body_layout)
        self._summary_card.setVisible(False)
        cl.addWidget(self._summary_card)

        # ── ASE add-on card ───────────────────────────────────────────────
        self._ase_card = _Card("ASE Reporting Add-On  (Optional)")
        self._build_ase_card(self._ase_card._body_layout)
        cl.addWidget(self._ase_card)

        if self._compact:
            cl.addStretch()
        else:
            cl.addStretch()
            scroll.setWidget(content)
            root.addWidget(scroll)

            # ── Bottom bar ────────────────────────────────────────────────
            bar = QWidget()
            bar.setStyleSheet("QWidget { background:#f9f0f2; border-top:1px solid #d6c0c5; }")
            bar.setFixedHeight(52)
            bar_row = QHBoxLayout(bar)
            bar_row.setContentsMargins(16, 8, 16, 8)
            bar_row.setSpacing(10)

            bar_row.addStretch()

            import_btn = QPushButton("⬆  Import .mcmxt…")
            import_btn.setStyleSheet(
                "QPushButton { background:#ffffff; color:#920d2e;"
                "  border:1px solid #d6c0c5; border-radius:5px;"
                "  padding:7px 18px; font-size:12px; font-weight:600; }"
                "QPushButton:hover { background:#fdf0f3; border-color:#920d2e; }"
            )
            import_btn.clicked.connect(self.load_data)
            bar_row.addWidget(import_btn)

            save_btn = QPushButton("💾  Save .mcmxt")
            save_btn.setStyleSheet(
                "QPushButton { background:#ffffff; color:#3a3a5c;"
                "  border:1px solid #d6c0c5; border-radius:5px;"
                "  padding:7px 18px; font-size:12px; font-weight:600; }"
                "QPushButton:hover { background:#f4f6fa; }"
            )
            save_btn.clicked.connect(self.save_data)
            bar_row.addWidget(save_btn)

            self._export_btn = QPushButton("Generate Document")
            self._export_btn.setStyleSheet(
                "QPushButton { background:#920d2e; color:#ffffff;"
                "  border:none; border-radius:5px;"
                "  padding:7px 18px; font-size:13px; font-weight:700; }"
                "QPushButton:hover { background:#7a0b27; }"
                "QPushButton:pressed { background:#600820; }"
                "QPushButton:disabled { background:#c8c8c8; color:#888888; }"
            )
            self._export_btn.setEnabled(False)
            self._export_btn.setToolTip("Confirm the schedule first.")
            self._export_btn.clicked.connect(self._export_report)
            bar_row.addWidget(self._export_btn)

            root.addWidget(bar)

    def _export_report(self):
        data = self.get_data()

        # ── Fall back to main-window fields when thermal-specific ones are empty ──
        # The thermal tab has its own contact/picture widgets; if the user filled
        # in the main Quote form instead, we pull from there automatically.
        top = self.window()
        if not data.get("ir_contact", "").strip():
            # Try the thermal-contact widget's individual fields first, then main form
            contact_lines = []
            for attr in ("_pres_name", "_pres_email", "_acct_name", "_acct_email"):
                val = ""
                if self._ir_contact and hasattr(self._ir_contact, attr):
                    val = getattr(self._ir_contact, attr).text().strip()
                contact_lines.append(val)
            if any(contact_lines):
                data["ir_contact"] = "\n".join([
                    contact_lines[0], "", contact_lines[1], "",
                    contact_lines[2], "", contact_lines[3]
                ])
            elif hasattr(top, "textEdit_contact"):
                # Raw main-form contact block — use first non-blank line as pres_name
                raw = top.textEdit_contact.toPlainText().strip()
                lines = [l.strip() for l in raw.splitlines() if l.strip()]
                pn = lines[0] if len(lines) > 0 else ""
                pe = lines[1] if len(lines) > 1 else ""
                an = lines[2] if len(lines) > 2 else ""
                ae = lines[3] if len(lines) > 3 else ""
                data["ir_contact"] = f"{pn}\n\n{pe}\n\n{an}\n\n{ae}"

        if not data.get("ir_picture", ""):
            if hasattr(top, "_picture_label") and hasattr(top._picture_label, "image_path"):
                data["ir_picture"] = top._picture_label.image_path or ""

        if not data.get("ir_proposal", ""):
            if hasattr(top, "lineEdit_proposal"):
                data["ir_proposal"] = top.lineEdit_proposal.text().strip()

        if not data.get("ir_customer", ""):
            if hasattr(top, "lineEdit_customer"):
                data["ir_customer"] = top.lineEdit_customer.text().strip()

        if not data.get("ir_location", ""):
            if hasattr(top, "lineEdit_location"):
                data["ir_location"] = top.lineEdit_location.text().strip()

        desktop  = os.path.join(os.path.expanduser("~"), "Desktop")
        proposal = data.get("ir_proposal", "THERMAL-REPORT").strip() or "THERMAL-REPORT"
        _ver     = f" {self._version_str()}"
        path, _ = QFileDialog.getSaveFileName(
            self, "Export Thermal Imaging Report",
            os.path.join(desktop, f"{proposal}{_ver}.docx"),
            "Word Document (*.docx)")
        if not path:
            return
        try:
            generate_thermal_doc(data, path)
        except Exception as e:
            import traceback
            QMessageBox.critical(self, "Export Failed", f"{e}\n\n{traceback.format_exc()}")
            return

        # Auto-save .mcmxt alongside the .docx
        mcmxt_path = os.path.splitext(path)[0] + ".mcmxt"
        mcmxt_saved = False
        try:
            with open(mcmxt_path, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            mcmxt_saved = True
        except Exception:
            pass

        if mcmxt_saved:
            QMessageBox.information(self, "Exported",
                f"Report saved to:\n{path}\n\nProject file:\n{mcmxt_path}")
        else:
            QMessageBox.information(self, "Exported", f"Report saved to:\n{path}")

    # ── Project Info panel ────────────────────────────────────────────────────

    def _build_project_info(self, layout):
        # Left: form fields
        form_w = QWidget()
        form_lay = QFormLayout(form_w)
        form_lay.setContentsMargins(0, 0, 0, 0)
        form_lay.setSpacing(8)
        form_lay.setLabelAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)

        def _lbl(text):
            l = QLabel(text)
            l.setStyleSheet(_FORM_LBL)
            return l

        self._ir_project = QLineEdit()
        self._ir_project.setStyleSheet(_FIELD_STYLE)
        form_lay.addRow(_lbl("Project Name:"), self._ir_project)

        self._ir_proposal = QLineEdit()
        self._ir_proposal.setPlaceholderText("MCMX-CUSTNAME-IRSTUDY-YYYYMMDD")
        self._ir_proposal.setStyleSheet(_FIELD_STYLE)
        form_lay.addRow(_lbl("Proposal Number:"), self._ir_proposal)

        self._ir_customer = QLineEdit()
        self._ir_customer.setStyleSheet(_FIELD_STYLE)
        form_lay.addRow(_lbl("Customer Name:"), self._ir_customer)

        self._ir_location = QLineEdit()
        self._ir_location.setPlaceholderText("City, State")
        self._ir_location.setStyleSheet(_FIELD_STYLE)
        form_lay.addRow(_lbl("Customer Location:"), self._ir_location)

        self._ir_contact = _ThermalContactWidget()
        form_lay.addRow(_lbl("Contact Info:"), self._ir_contact)

        layout.addWidget(form_w, stretch=1)

        # Right: logo drag-drop area
        logo_box = QWidget()
        logo_box.setFixedWidth(200)
        logo_box.setStyleSheet(
            "QWidget { background:#fff0f3; border:1px solid #e8dde0; border-radius:8px; }"
        )
        logo_vl = QVBoxLayout(logo_box)
        logo_vl.setContentsMargins(10, 10, 10, 10)
        logo_vl.setSpacing(6)
        logo_vl.setAlignment(Qt.AlignmentFlag.AlignCenter)

        logo_lbl = QLabel("Customer Logo")
        logo_lbl.setStyleSheet(f"font-size:11px; font-weight:700; color:{_RED}; border:none; background:transparent;")
        logo_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        logo_vl.addWidget(logo_lbl)

        self._ir_picture = DragDropLabel(parent=logo_box)
        self._ir_picture.setFixedSize(150, 150)
        logo_vl.addWidget(self._ir_picture, alignment=Qt.AlignmentFlag.AlignCenter)

        btn_row = QHBoxLayout(); btn_row.setSpacing(6)
        upload_btn = QPushButton("Upload")
        upload_btn.setStyleSheet(
            "QPushButton { background:#920d2e; color:#fff; border:none;"
            "  border-radius:4px; padding:4px 10px; font-size:11px; font-weight:700; }"
            "QPushButton:hover { background:#b01036; }"
        )
        upload_btn.clicked.connect(self._upload_ir_picture)
        btn_row.addWidget(upload_btn)

        clear_btn = QPushButton("Clear")
        clear_btn.setStyleSheet(
            "QPushButton { background:#f4f6fa; color:#3a3a5c; border:1px solid #d6c0c5;"
            "  border-radius:4px; padding:4px 10px; font-size:11px; }"
            "QPushButton:hover { background:#e8ecf5; }"
        )
        clear_btn.clicked.connect(lambda: self._ir_picture.set_image(None) if hasattr(self._ir_picture, 'set_image') else None)
        btn_row.addWidget(clear_btn)
        logo_vl.addLayout(btn_row)

        layout.addWidget(logo_box)

    def _upload_ir_picture(self):
        path, _ = QFileDialog.getOpenFileName(self, "Upload Customer Logo", "", "Images (*.png *.jpg *.jpeg *.bmp *.svg)")
        if path: self._ir_picture.set_image(path)

    # ── Config card body ──────────────────────────────────────────────────────

    def _build_config(self, layout):
        # Hours per day + Working days
        row1 = QHBoxLayout(); row1.setSpacing(16)

        hl = QLabel("Hours / Day:")
        hl.setStyleSheet(f"font-size:12px; font-weight:700; color:{_RED};")
        row1.addWidget(hl)
        self._hours_edit = QLineEdit("8")
        self._hours_edit.setValidator(QDoubleValidator(0.5, 24.0, 1))
        self._hours_edit.setFixedWidth(70)
        self._hours_edit.setStyleSheet(_FIELD_STYLE)
        self._hours_edit.textChanged.connect(self._mark_outlook_stale)
        row1.addWidget(self._hours_edit)

        row1.addSpacing(16)
        dl = QLabel("Working Days:")
        dl.setStyleSheet(f"font-size:12px; font-weight:700; color:{_RED};")
        row1.addWidget(dl)
        self._days_edit = QLineEdit()
        self._days_edit.setPlaceholderText("e.g. 5")
        self._days_edit.setValidator(QIntValidator(1, 9999))
        self._days_edit.setFixedWidth(80)
        self._days_edit.setStyleSheet(_FIELD_STYLE)
        self._days_edit.textChanged.connect(self._mark_outlook_stale)
        self._days_edit.textChanged.connect(self._update_gen_btn_state)
        row1.addWidget(self._days_edit)
        row1.addStretch()
        layout.addLayout(row1)

        # Technician count
        tc_row = QHBoxLayout(); tc_row.setSpacing(8)
        tl = QLabel("Technicians:")
        tl.setStyleSheet(f"font-size:12px; font-weight:700; color:{_RED};")
        tl.setFixedWidth(90)
        tc_row.addWidget(tl)

        self._tech_minus = QPushButton("−")
        self._tech_minus.setFixedSize(28, 28)
        self._tech_minus.setStyleSheet(_SPIN_S)
        self._tech_minus.clicked.connect(self._dec_techs)
        tc_row.addWidget(self._tech_minus)

        self._tech_count_lbl = QLabel("1")
        self._tech_count_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self._tech_count_lbl.setStyleSheet(f"font-size:13px; font-weight:700; color:{_TEXT}; min-width:24px;")
        tc_row.addWidget(self._tech_count_lbl)

        self._tech_plus = QPushButton("+")
        self._tech_plus.setFixedSize(28, 28)
        self._tech_plus.setStyleSheet(_SPIN_S)
        self._tech_plus.clicked.connect(self._inc_techs)
        tc_row.addWidget(self._tech_plus)
        tc_row.addStretch()
        layout.addLayout(tc_row)

        # Tech rows container
        self._tech_container = QWidget()
        self._tech_container.setStyleSheet("background:transparent;")
        self._tech_vbox = QVBoxLayout(self._tech_container)
        self._tech_vbox.setContentsMargins(0, 4, 0, 0)
        self._tech_vbox.setSpacing(6)
        layout.addWidget(self._tech_container)
        self._rebuild_tech_rows()

        # Day-of-week checkboxes
        dw_row = QHBoxLayout(); dw_row.setSpacing(8)
        dw_lbl = QLabel("Work Days:")
        dw_lbl.setStyleSheet(f"font-size:12px; font-weight:700; color:{_RED};")
        dw_lbl.setFixedWidth(78)
        dw_row.addWidget(dw_lbl)
        self._day_checks = []
        for name, default in zip(_DAY_NAMES, [True,True,True,True,True,False,False]):
            cb = QCheckBox(name)
            cb.setChecked(default)
            cb.setStyleSheet(f"QCheckBox {{ font-size:12px; color:{_TEXT}; }}")
            cb.toggled.connect(self._mark_outlook_stale)
            dw_row.addWidget(cb)
            self._day_checks.append(cb)
        dw_row.addStretch()
        layout.addLayout(dw_row)

        # ── Rate overrides ────────────────────────────────────────────────────
        sep0 = QFrame(); sep0.setFrameShape(QFrame.Shape.HLine)
        sep0.setStyleSheet(f"color:{_BORDER};"); layout.addWidget(sep0)

        rate_lbl = QLabel("Billing Rates  (override defaults)")
        rate_lbl.setStyleSheet(f"font-size:11px; font-weight:700; color:{_SUBTEXT};")
        layout.addWidget(rate_lbl)

        _rate_field_s = (
            "QLineEdit { color:#1a0509; background:#ffffff; border:1px solid #d6c0c5;"
            "  border-radius:4px; padding:3px 6px; font-size:11px; }"
            "QLineEdit:focus { border-color:#920d2e; }"
        )
        rates_row = QHBoxLayout(); rates_row.setSpacing(16)
        for attr, caption, default, suffix in [
            ("_rate_labor",   "Labor/hr (Work)",   "125",  "/hr"),
            ("_rate_travel",  "Labor/hr (Travel)",  "100",  "/hr"),
            ("_rate_hotel",   "Hotel/night",         "150",  "/night"),
            ("_rate_meal",    "Meal/day",            "25",   "/meal"),
            ("_rate_mile",    "Mileage/mi",          "0.72", "/mi"),
        ]:
            col = QVBoxLayout(); col.setSpacing(2)
            lbl = QLabel(caption); lbl.setStyleSheet(f"font-size:10px; color:{_SUBTEXT};")
            col.addWidget(lbl)
            row_inner = QHBoxLayout(); row_inner.setSpacing(2)
            dollar = QLabel("$"); dollar.setStyleSheet(f"font-size:11px; color:{_TEXT};")
            row_inner.addWidget(dollar)
            fld = QLineEdit(default)
            fld.setFixedWidth(68)
            fld.setValidator(QDoubleValidator(0.0, 99999.0, 2))
            fld.setStyleSheet(_rate_field_s)
            fld.textChanged.connect(self._mark_outlook_stale)
            row_inner.addWidget(fld)
            suf = QLabel(suffix); suf.setStyleSheet(f"font-size:10px; color:{_SUBTEXT};")
            row_inner.addWidget(suf)
            col.addLayout(row_inner)
            rates_row.addLayout(col)
            setattr(self, attr, fld)
        rates_row.addStretch()
        layout.addLayout(rates_row)

        # Separator
        sep = QFrame(); sep.setFrameShape(QFrame.Shape.HLine)
        sep.setStyleSheet(f"color:{_BORDER};"); layout.addWidget(sep)

        # Generate button — disabled until working days are entered
        gen_row = QHBoxLayout(); gen_row.addStretch()
        self._gen_btn = QPushButton("Generate Work Week Overview  →")
        self._gen_btn.setStyleSheet(_ACTN_STYLE)
        self._gen_btn.setEnabled(False)
        self._gen_btn.setToolTip("Enter the number of working days first.")
        self._gen_btn.clicked.connect(self._generate_outlook)
        gen_row.addWidget(self._gen_btn)
        layout.addLayout(gen_row)

    # ── Outlook card body ─────────────────────────────────────────────────────

    def _build_outlook_card(self, layout):
        self._outlook_table = QTableWidget()
        self._outlook_table.setStyleSheet(_TABLE_STYLE)
        self._outlook_table.setSelectionMode(QAbstractItemView.SelectionMode.NoSelection)
        self._outlook_table.verticalHeader().setVisible(False)
        self._outlook_table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        layout.addWidget(self._outlook_table)

        confirm_row = QHBoxLayout(); confirm_row.addStretch()
        self._confirm_btn = QPushButton("Confirm Schedule  →")
        self._confirm_btn.setStyleSheet(_ACTN_STYLE)
        self._confirm_btn.setEnabled(False)
        self._confirm_btn.setToolTip("Generate the Work Week Overview first.")
        self._confirm_btn.clicked.connect(self._confirm_schedule)
        confirm_row.addWidget(self._confirm_btn)
        layout.addLayout(confirm_row)

    # ── Summary card body ─────────────────────────────────────────────────────

    def _build_summary_card(self, layout):
        # Stats row
        stats_row = QHBoxLayout(); stats_row.setSpacing(20)
        for attr, caption in [
            ("_sum_days",    "Working Days"),
            ("_sum_hotel",   "Hotel Nights"),
            ("_sum_meals",   "Total Meals"),
            ("_sum_flights", "Air Trips"),
            ("_sum_miles",   "Drive Miles"),
        ]:
            col = QVBoxLayout(); col.setSpacing(2)
            lbl = QLabel(caption); lbl.setStyleSheet(f"font-size:10px; color:{_SUBTEXT};")
            col.addWidget(lbl)
            val = QLineEdit("—"); val.setReadOnly(True)
            val.setStyleSheet(_CALC_STYLE); val.setFixedWidth(100)
            col.addWidget(val)
            setattr(self, attr, val)
            stats_row.addLayout(col)
        stats_row.addStretch()
        layout.addLayout(stats_row)

        # Meals count is manually adjustable — remove read-only and wire recalc
        self._sum_meals.setReadOnly(False)
        self._sum_meals.setStyleSheet(
            "QLineEdit { color:#1a3a6e; background:#ffffff;"
            "  border:2px solid #7b8cde; border-radius:4px;"
            "  padding:4px 8px; font-size:14px; font-weight:700; }"
            "QLineEdit:focus { border-color:#920d2e; }"
        )
        self._sum_meals.setToolTip("Total Meals — edit to manually override")
        self._sum_meals.editingFinished.connect(self._on_meals_edited)

        sep = QFrame(); sep.setFrameShape(QFrame.Shape.HLine)
        sep.setStyleSheet(f"color:{_BORDER}; margin-top:4px;"); layout.addWidget(sep)

        # Cost breakdown
        cost_lbl = QLabel("Cost Breakdown")
        cost_lbl.setStyleSheet(f"font-size:12px; font-weight:700; color:{_RED}; margin-top:4px;")
        layout.addWidget(cost_lbl)

        cost_grid = QHBoxLayout(); cost_grid.setSpacing(20)
        for attr, caption in [
            ("_cost_labor",   "Labor (Work)"),
            ("_cost_travel",  "Labor (Travel)"),
            ("_cost_hotel",   "Hotel"),
            ("_cost_meals",   "Meals"),
            ("_cost_miles",   "Mileage"),
            ("_cost_flights", "Flights"),
        ]:
            col = QVBoxLayout(); col.setSpacing(2)
            lbl = QLabel(caption); lbl.setStyleSheet(f"font-size:10px; color:{_SUBTEXT};")
            col.addWidget(lbl)
            val = QLineEdit("—"); val.setReadOnly(True)
            val.setStyleSheet(_CALC_STYLE); val.setFixedWidth(100)
            col.addWidget(val)
            setattr(self, attr, val)
            cost_grid.addLayout(col)
        cost_grid.addStretch()
        layout.addLayout(cost_grid)

        # Margin control + totals
        margin_row = QHBoxLayout(); margin_row.setSpacing(12)

        ml = QLabel("Margin %:")
        ml.setStyleSheet(f"font-size:12px; font-weight:700; color:{_RED};")
        margin_row.addWidget(ml)

        self._margin_minus = QPushButton("−")
        self._margin_minus.setFixedSize(28, 28)
        self._margin_minus.setStyleSheet(_SPIN_S)
        self._margin_minus.clicked.connect(lambda: self._adj_margin(-1))
        margin_row.addWidget(self._margin_minus)

        self._margin_lbl = QLineEdit("0.0")
        self._margin_lbl.setFixedWidth(60)
        self._margin_lbl.setStyleSheet(_FIELD_STYLE)
        self._margin_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self._margin_lbl.setValidator(QDoubleValidator(0.0, 99.9, 1))
        self._margin_lbl.textChanged.connect(self._recalc_cost_totals)
        margin_row.addWidget(self._margin_lbl)

        pct_lbl = QLabel("%")
        pct_lbl.setStyleSheet(f"font-size:12px; color:{_TEXT};")
        margin_row.addWidget(pct_lbl)

        self._margin_plus = QPushButton("+")
        self._margin_plus.setFixedSize(28, 28)
        self._margin_plus.setStyleSheet(_SPIN_S)
        self._margin_plus.clicked.connect(lambda: self._adj_margin(+1))
        margin_row.addWidget(self._margin_plus)

        margin_row.addSpacing(24)

        for attr, caption, is_grand in [
            ("_cost_subtotal", "Subtotal (before margin)", False),
            ("_cost_grand",    "Grand Total",              True),
        ]:
            lbl = QLabel(caption + ":")
            lbl.setStyleSheet(f"font-size:12px; color:{_SUBTEXT};")
            margin_row.addWidget(lbl)
            val = QLineEdit("—"); val.setReadOnly(True)
            val.setStyleSheet(
                _CALC_STYLE if not is_grand else
                "QLineEdit { color:#920d2e; background:#f9f0f2;"
                "  border:1px solid #d6c0c5; border-radius:4px;"
                "  padding:4px 8px; font-size:13px; font-weight:700; }"
            )
            val.setFixedWidth(120 if is_grand else 110)
            setattr(self, attr, val)
            margin_row.addWidget(val)
        margin_row.addStretch()
        layout.addLayout(margin_row)

    # ── ASE add-on card body ──────────────────────────────────────────────────

    def _build_ase_card(self, layout):
        desc = QLabel(
            "Optional ASE report development service. Fixed pricing per image tier "
            "for IR scanning report development. Margin applied separately below."
        )
        desc.setStyleSheet(f"font-size:11px; color:{_SUBTEXT};")
        desc.setWordWrap(True)
        layout.addWidget(desc)

        # Enable toggle
        enable_row = QHBoxLayout(); enable_row.setSpacing(12)
        self._ase_enabled = QCheckBox("Include ASE Reporting")
        self._ase_enabled.setStyleSheet(_CB_STYLE)
        self._ase_enabled.toggled.connect(self._on_ase_toggled)
        enable_row.addWidget(self._ase_enabled)
        enable_row.addStretch()
        layout.addLayout(enable_row)

        # Tier + custom price (wrapped in a container so we can hide easily)
        self._ase_body = QWidget()
        self._ase_body.setVisible(False)
        ase_bl = QVBoxLayout(self._ase_body)
        ase_bl.setContentsMargins(0, 6, 0, 0); ase_bl.setSpacing(8)

        # Pricing table display
        tbl_lbl = QLabel("Pricing Tiers:")
        tbl_lbl.setStyleSheet(f"font-size:11px; font-weight:700; color:{_RED};")
        ase_bl.addWidget(tbl_lbl)

        tier_tbl = QTableWidget(len(_ASE_TIERS), 2)
        tier_tbl.setHorizontalHeaderLabels(["# of Images Included in Scope", "ASE to Mc-Mc Price"])
        tier_tbl.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        tier_tbl.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.Fixed)
        tier_tbl.setColumnWidth(1, 150)
        tier_tbl.verticalHeader().setVisible(False)
        tier_tbl.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        tier_tbl.setSelectionMode(QAbstractItemView.SelectionMode.NoSelection)
        tier_tbl.setFixedHeight(len(_ASE_TIERS) * 28 + 34)
        tier_tbl.setStyleSheet(_TABLE_STYLE)
        for i, (label, price) in enumerate(_ASE_TIERS):
            it_lbl = QTableWidgetItem(label)
            it_lbl.setFlags(Qt.ItemFlag.ItemIsEnabled)
            tier_tbl.setItem(i, 0, it_lbl)
            it_p = QTableWidgetItem("Custom" if price is None else f"$  {price:,.2f}")
            it_p.setFlags(Qt.ItemFlag.ItemIsEnabled)
            it_p.setTextAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
            if price is None:
                it_p.setForeground(QColor(_SUBTEXT))
            tier_tbl.setItem(i, 1, it_p)
            tier_tbl.setRowHeight(i, 28)
        ase_bl.addWidget(tier_tbl)

        # Tier selector row
        sel_row = QHBoxLayout(); sel_row.setSpacing(12)
        sel_lbl = QLabel("Selected Tier:")
        sel_lbl.setStyleSheet(f"font-size:12px; font-weight:700; color:{_RED};")
        sel_row.addWidget(sel_lbl)

        self._ase_tier_combo = QComboBox()
        self._ase_tier_combo.setStyleSheet(
            "QComboBox { font-size:11px; color:#1a0509; background:#ffffff;"
            "  border:1px solid #d6c0c5; border-radius:4px; padding:4px 10px; min-width:200px; }"
            "QComboBox::drop-down { border:none; width:18px; background:#f4f6fa;"
            "  border-left:1px solid #d6c0c5; border-top-right-radius:4px;"
            "  border-bottom-right-radius:4px; }"
            "QComboBox::down-arrow { width:0; height:0;"
            "  border-left:4px solid transparent; border-right:4px solid transparent;"
            "  border-top:5px solid #920d2e; }"
            "QComboBox QAbstractItemView { font-size:11px; color:#1a0509; background:#ffffff;"
            "  border:1px solid #d6c0c5; selection-background-color:#f5d0da;"
            "  selection-color:#920d2e; outline:none; }"
        )
        for label, _ in _ASE_TIERS:
            self._ase_tier_combo.addItem(label)
        self._ase_tier_combo.currentIndexChanged.connect(self._on_ase_tier_changed)
        sel_row.addWidget(self._ase_tier_combo)
        sel_row.addStretch()
        ase_bl.addLayout(sel_row)

        # Custom price input (visible only for 201+)
        custom_row = QHBoxLayout(); custom_row.setSpacing(8)
        self._ase_custom_lbl = QLabel("Custom Price ($):")
        self._ase_custom_lbl.setStyleSheet(f"font-size:12px; color:{_SUBTEXT};")
        custom_row.addWidget(self._ase_custom_lbl)
        self._ase_custom_edit = QLineEdit()
        self._ase_custom_edit.setPlaceholderText("Enter price…")
        self._ase_custom_edit.setValidator(QDoubleValidator(0.0, 9_999_999.0, 2))
        self._ase_custom_edit.setFixedWidth(130)
        self._ase_custom_edit.setStyleSheet(_FIELD_STYLE)
        self._ase_custom_edit.textChanged.connect(self._recalc_cost_totals)
        custom_row.addWidget(self._ase_custom_edit)
        custom_row.addStretch()
        ase_bl.addLayout(custom_row)
        self._ase_custom_lbl.setVisible(False)
        self._ase_custom_edit.setVisible(False)

        # ASE margin
        ase_margin_row = QHBoxLayout(); ase_margin_row.setSpacing(12)
        aml = QLabel("ASE Margin %:")
        aml.setStyleSheet(f"font-size:12px; font-weight:700; color:{_RED};")
        ase_margin_row.addWidget(aml)

        self._ase_margin_minus = QPushButton("−")
        self._ase_margin_minus.setFixedSize(28, 28)
        self._ase_margin_minus.setStyleSheet(_SPIN_S)
        self._ase_margin_minus.clicked.connect(lambda: self._adj_ase_margin(-1))
        ase_margin_row.addWidget(self._ase_margin_minus)

        self._ase_margin_lbl = QLineEdit("0.0")
        self._ase_margin_lbl.setFixedWidth(60)
        self._ase_margin_lbl.setStyleSheet(_FIELD_STYLE)
        self._ase_margin_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self._ase_margin_lbl.setValidator(QDoubleValidator(0.0, 99.9, 1))
        self._ase_margin_lbl.textChanged.connect(self._recalc_cost_totals)
        ase_margin_row.addWidget(self._ase_margin_lbl)

        ase_pct = QLabel("%"); ase_pct.setStyleSheet(f"font-size:12px; color:{_TEXT};")
        ase_margin_row.addWidget(ase_pct)

        self._ase_margin_plus = QPushButton("+")
        self._ase_margin_plus.setFixedSize(28, 28)
        self._ase_margin_plus.setStyleSheet(_SPIN_S)
        self._ase_margin_plus.clicked.connect(lambda: self._adj_ase_margin(+1))
        ase_margin_row.addWidget(self._ase_margin_plus)

        ase_margin_row.addSpacing(24)
        ase_cost_lbl = QLabel("ASE Grand Total:")
        ase_cost_lbl.setStyleSheet(f"font-size:12px; color:{_SUBTEXT};")
        ase_margin_row.addWidget(ase_cost_lbl)
        self._ase_total_edit = QLineEdit("—")
        self._ase_total_edit.setReadOnly(True)
        self._ase_total_edit.setStyleSheet(
            "QLineEdit { color:#920d2e; background:#f9f0f2;"
            "  border:1px solid #d6c0c5; border-radius:4px;"
            "  padding:4px 8px; font-size:13px; font-weight:700; }"
        )
        self._ase_total_edit.setFixedWidth(120)
        ase_margin_row.addWidget(self._ase_total_edit)
        ase_margin_row.addStretch()
        ase_bl.addLayout(ase_margin_row)

        layout.addWidget(self._ase_body)

    # ── Technician row management ─────────────────────────────────────────────

    def _rebuild_tech_rows(self):
        while self._tech_vbox.count():
            item = self._tech_vbox.takeAt(0)
            if item.widget(): item.widget().deleteLater()
        self._tech_rows = []

        _mode_style = (
            "QComboBox { font-size:11px; color:#1a0509; background:#ffffff;"
            "  border:1px solid #d6c0c5; border-radius:4px; padding:3px 8px; }"
            "QComboBox:focus { border-color:#920d2e; }"
            "QComboBox::drop-down { border:none; width:18px; background:#f4f6fa;"
            "  border-left:1px solid #d6c0c5; border-top-right-radius:4px;"
            "  border-bottom-right-radius:4px; }"
            "QComboBox::down-arrow { width:0; height:0;"
            "  border-left:4px solid transparent; border-right:4px solid transparent;"
            "  border-top:5px solid #920d2e; }"
            "QComboBox QAbstractItemView { font-size:11px; color:#1a0509; background:#ffffff;"
            "  border:1px solid #d6c0c5; selection-background-color:#f5d0da;"
            "  selection-color:#920d2e; }"
        )

        for i in range(self._num_techs):
            row_w = QWidget()
            row_w.setStyleSheet("background:transparent;")
            hl = QHBoxLayout(row_w)
            hl.setContentsMargins(0, 0, 0, 0); hl.setSpacing(8)

            num_lbl = QLabel(f"Tech {i+1}:")
            num_lbl.setStyleSheet(f"font-size:12px; color:{_SUBTEXT}; min-width:52px;")
            hl.addWidget(num_lbl)

            name_ed = QLineEdit()
            name_ed.setPlaceholderText(f"Technician {i+1}")
            name_ed.setStyleSheet(_FIELD_STYLE)
            name_ed.setFixedWidth(150)
            hl.addWidget(name_ed)

            tl = QLabel("Travel Time:")
            tl.setStyleSheet(f"font-size:12px; color:{_SUBTEXT};")
            hl.addWidget(tl)

            travel_ed = QLineEdit()
            travel_ed.setPlaceholderText("hrs")
            travel_ed.setValidator(QDoubleValidator(0.0, 48.0, 1))
            travel_ed.setFixedWidth(60)
            travel_ed.setStyleSheet(_FIELD_STYLE)
            travel_ed.textChanged.connect(self._mark_outlook_stale)
            hl.addWidget(travel_ed)

            mode_combo = QComboBox()
            mode_combo.addItem("Driving", "driving")
            mode_combo.addItem("Flying",  "flying")
            mode_combo.setFixedWidth(90)
            mode_combo.setStyleSheet(_mode_style)
            mode_combo.currentTextChanged.connect(self._mark_outlook_stale)
            hl.addWidget(mode_combo)

            cost_lbl = QLabel("Flight Cost:")
            cost_lbl.setStyleSheet(f"font-size:12px; color:{_SUBTEXT};")
            hl.addWidget(cost_lbl)
            cost_ed = QLineEdit()
            cost_ed.setPlaceholderText("$ per trip")
            cost_ed.setValidator(QDoubleValidator(0.0, 99999.0, 2))
            cost_ed.setFixedWidth(90)
            cost_ed.setStyleSheet(_FIELD_STYLE)
            hl.addWidget(cost_ed)

            mile_lbl = QLabel("Mileage:")
            mile_lbl.setStyleSheet(f"font-size:12px; color:{_SUBTEXT};")
            hl.addWidget(mile_lbl)
            mile_ed = QLineEdit()
            mile_ed.setPlaceholderText("mi one-way")
            mile_ed.setValidator(QDoubleValidator(0.0, 99999.0, 1))
            mile_ed.setFixedWidth(90)
            mile_ed.setStyleSheet(_FIELD_STYLE)
            hl.addWidget(mile_ed)

            hl.addStretch()

            def _toggle(text, cl=cost_lbl, ce=cost_ed, ml=mile_lbl, me=mile_ed):
                flying = text == "Flying"
                cl.setVisible(flying); ce.setVisible(flying)
                ml.setVisible(not flying); me.setVisible(not flying)

            mode_combo.currentTextChanged.connect(_toggle)
            _toggle(mode_combo.currentText())

            self._tech_vbox.addWidget(row_w)
            self._tech_rows.append({
                "name": name_ed, "travel": travel_ed, "mode": mode_combo,
                "flight_cost": cost_ed, "mileage": mile_ed, "_toggle": _toggle,
            })

    def _snapshot_tech_data(self):
        return [{"name": r["name"].text(), "travel": r["travel"].text(),
                 "mode": r["mode"].currentText(), "flight_cost": r["flight_cost"].text(),
                 "mileage": r["mileage"].text()} for r in self._tech_rows]

    def _restore_tech_data(self, saved):
        for i, s in enumerate(saved):
            if i >= len(self._tech_rows): break
            r = self._tech_rows[i]
            r["name"].setText(s["name"]); r["travel"].setText(s["travel"])
            idx = r["mode"].findText(s["mode"])
            if idx >= 0: r["mode"].setCurrentIndex(idx)
            r["flight_cost"].setText(s["flight_cost"]); r["mileage"].setText(s["mileage"])

    def _mark_outlook_stale(self, *_):
        """Called whenever config changes — disables Confirm and Export until regenerated."""
        self._confirmed = {}
        if hasattr(self, '_confirm_btn'):
            self._confirm_btn.setEnabled(False)
            self._confirm_btn.setToolTip("Regenerate the Work Week Overview after changing config.")
        if hasattr(self, '_export_btn'):
            self._export_btn.setEnabled(False)
            self._export_btn.setToolTip("Confirm the schedule first.")

    def _update_gen_btn_state(self, *_):
        """Enable Generate button only when a day count has been entered."""
        if hasattr(self, '_gen_btn'):
            has_days = bool(self._days_edit.text().strip())
            self._gen_btn.setEnabled(has_days)
            self._gen_btn.setToolTip("" if has_days else "Enter the number of working days first.")

    def _inc_techs(self):
        if self._num_techs < 10:
            saved = self._snapshot_tech_data()
            self._num_techs += 1
            self._tech_count_lbl.setText(str(self._num_techs))
            self._rebuild_tech_rows(); self._restore_tech_data(saved)
            self._mark_outlook_stale()

    def _dec_techs(self):
        if self._num_techs > 1:
            saved = self._snapshot_tech_data()
            self._num_techs -= 1
            self._tech_count_lbl.setText(str(self._num_techs))
            self._rebuild_tech_rows(); self._restore_tech_data(saved)
            self._mark_outlook_stale()

    # ── Schedule generation ───────────────────────────────────────────────────

    def _get_work_days(self):
        return [i for i, cb in enumerate(self._day_checks) if cb.isChecked()]

    def _generate_outlook(self):
        work_days = int(_n(self._days_edit.text(), 0))
        wdays = self._get_work_days()
        if work_days <= 0 or not wdays:
            QMessageBox.warning(self, "Missing Input",
                "Enter the number of working days and select at least one work day.")
            return

        if len(wdays) > 1:
            gap_days = (wdays[0] + 7 - wdays[-1] - 1) % 7
        else:
            gap_days = 6
        self._weekend_days = gap_days

        sched = []

        # Travel In/Out only when at least one tech's journey warrants it
        # (flying, or drive time >= 2 hrs one-way).  Local techs drive RT
        # each day — no overnight stays, so no dedicated travel days.
        _n_techs = self._num_techs
        _any_hotel_needed = any(
            _n(self._tech_rows[i]["travel"].text(), 0.0) >= 2.0 or
            self._tech_rows[i]["mode"].currentText() == "Flying"
            for i in range(_n_techs)
            if i < len(self._tech_rows)
        )

        travel_in_day = _DAY_NAMES[(wdays[0] - 1) % 7]
        if _any_hotel_needed:
            sched.append((travel_in_day, "Travel In", False, True))

        done = 0; last_wd = wdays[0]; first_week = True
        while done < work_days:
            if not first_week and gap_days > 0:
                sched.append(("Weekend", "— Travel Home / Return —", False, False))
            first_week = False
            for wd in wdays:
                if done >= work_days: break
                done += 1; last_wd = wd; is_last = done == work_days
                sched.append((_DAY_NAMES[wd], f"Work  Day {done}", True, True))

        travel_out_day = _DAY_NAMES[(last_wd + 1) % 7]
        if _any_hotel_needed:
            sched.append((travel_out_day, "Travel Out", False, False))

        self._schedule = sched
        self._confirmed = {}
        self._refresh_outlook_table(sched, self._num_techs)
        self._outlook_card.setVisible(True)
        self._summary_card.setVisible(False)
        self._confirm_btn.setEnabled(True)
        self._confirm_btn.setToolTip("")
        if hasattr(self, '_export_btn'):
            self._export_btn.setEnabled(False)
            self._export_btn.setToolTip("Confirm the schedule first.")

    def _refresh_outlook_table(self, sched, num_techs):
        tech_names = [
            self._tech_rows[i]["name"].text().strip() or f"Tech {i+1}"
            for i in range(num_techs)
        ]
        tech_info = []
        for i in range(num_techs):
            if i < len(self._tech_rows):
                row = self._tech_rows[i]
                tech_info.append({
                    "mode":        row["mode"].currentText(),
                    "travel_hrs":  _n(row["travel"].text()),
                    "mileage_ow":  _n(row["mileage"].text()),
                    "flight_cost": _n(row["flight_cost"].text()),
                })
            else:
                tech_info.append({"mode":"Driving","travel_hrs":0.0,"mileage_ow":0.0,"flight_cost":0.0})

        headers = ["Day", "Activity"] + tech_names
        self._outlook_table.setColumnCount(len(headers))
        self._outlook_table.setHorizontalHeaderLabels(headers)
        self._outlook_table.setRowCount(len(sched))

        hh = self._outlook_table.horizontalHeader()
        hh.setMinimumHeight(46)
        hh.setDefaultAlignment(Qt.AlignmentFlag.AlignCenter | Qt.TextFlag.TextWordWrap)
        hh.setSectionResizeMode(0, QHeaderView.ResizeMode.Fixed)
        self._outlook_table.setColumnWidth(0, 95)
        hh.setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        for c in range(2, len(headers)):
            hh.setSectionResizeMode(c, QHeaderView.ResizeMode.Fixed)
            self._outlook_table.setColumnWidth(c, 155)

        self._hotel_checks      = [[None] * len(sched) for _ in range(num_techs)]
        self._travel_in_checks  = {}
        self._travel_out_checks = {}

        for r, (day_lbl, activity, is_work, def_hotel) in enumerate(sched):
            is_weekend    = day_lbl == "Weekend"
            is_travel_in  = activity == "Travel In"
            is_travel_out = activity == "Travel Out"
            is_travel     = is_travel_in or is_travel_out

            day_item = QTableWidgetItem(day_lbl)
            day_item.setFlags(Qt.ItemFlag.ItemIsEnabled)
            day_item.setForeground(QColor(
                _SUBTEXT if is_weekend else (_RED if is_work else _TEXT)))
            if is_weekend:
                day_item.setBackground(QColor("#f4f6fa"))
            self._outlook_table.setItem(r, 0, day_item)

            act_item = QTableWidgetItem(activity)
            act_item.setFlags(Qt.ItemFlag.ItemIsEnabled)
            act_item.setForeground(QColor(_SUBTEXT if is_weekend else _TEXT))
            if is_weekend:
                act_item.setBackground(QColor("#f4f6fa"))
                act_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self._outlook_table.setItem(r, 1, act_item)

            for t in range(num_techs):
                info   = tech_info[t]
                flying = info["mode"] == "Flying"
                t_hrs  = info["travel_hrs"]
                mi_ow  = info["mileage_ow"]
                effective_def_hotel = def_hotel and (flying or t_hrs >= 2.0)

                cell_w  = QWidget()
                cell_vl = QVBoxLayout(cell_w)
                cell_vl.setContentsMargins(8, 4, 8, 4)
                cell_vl.setSpacing(2)
                cell_vl.setAlignment(Qt.AlignmentFlag.AlignVCenter)

                if is_weekend:
                    cell_w.setStyleSheet("background:#f4f6fa;")
                    if flying:
                        home_txt = "Flying home  (RT)"
                    else:
                        rt = mi_ow * 2
                        home_txt = f"Drive home  ({rt:,.0f} mi RT)" if rt > 0 else "Drive home  (RT)"
                    lbl_h = QLabel(home_txt)
                    lbl_h.setStyleSheet(_INFO_LBL)
                    lbl_h.setAlignment(Qt.AlignmentFlag.AlignCenter)
                    cell_vl.addWidget(lbl_h)
                    if t_hrs > 0:
                        lbl_t = QLabel(f"{t_hrs*2:.1f} hrs travel")
                        lbl_t.setStyleSheet(_INFO_LBL)
                        lbl_t.setAlignment(Qt.AlignmentFlag.AlignCenter)
                        cell_vl.addWidget(lbl_t)

                elif is_travel:
                    cell_w.setStyleSheet("background:#fff0f3;" if effective_def_hotel else "background:#ffffff;")

                    # Hotel checkbox always shown on travel days
                    this_hotel_cb = QCheckBox("Book hotel")
                    this_hotel_cb.setChecked(effective_def_hotel)
                    this_hotel_cb.setStyleSheet(_CB_STYLE)
                    this_hotel_cb.toggled.connect(
                        lambda checked, w=cell_w: w.setStyleSheet(
                            "background:#fff0f3;" if checked else "background:#ffffff;"
                        )
                    )
                    cell_vl.addWidget(this_hotel_cb)
                    self._hotel_checks[t][r] = this_hotel_cb

                    if t_hrs > 0:
                        _tl = QLabel(f"Travel: {t_hrs:.1f} hrs OW")
                        _tl.setStyleSheet(_INFO_LBL)
                        cell_vl.addWidget(_tl)

                    if flying:
                        f_cost = info["flight_cost"]
                        _ml = QLabel(f"Flight: {'${:,.0f}'.format(f_cost) if f_cost > 0 else '—'}")
                    else:
                        if is_travel_out:
                            _ml = QLabel(f"Drive home: {mi_ow:,.0f} mi OW" if mi_ow > 0 else "Drive home: OW")
                        else:
                            _ml = QLabel(f"Drive: {mi_ow:,.0f} mi OW" if mi_ow > 0 else "Drive: — mi")
                    _ml.setStyleSheet(_INFO_LBL)
                    cell_vl.addWidget(_ml)

                else:
                    # Working day
                    _has_hotel = effective_def_hotel
                    cell_w.setStyleSheet("background:#fff0f3;" if _has_hotel else "background:#ffffff;")

                    cb = QCheckBox("Book hotel")
                    cb.setChecked(_has_hotel)
                    cb.setStyleSheet(_CB_STYLE)
                    cell_vl.addWidget(cb)
                    self._hotel_checks[t][r] = cb

                    _r_text = self._rate_labor.text().strip() or str(int(self.LABOR_RATE))
                    work_rate_lbl = QLabel(f"${_r_text}/hr  ×  {_n(self._hours_edit.text(), 8.0):.0f}h")
                    work_rate_lbl.setStyleSheet(_INFO_LBL)
                    cell_vl.addWidget(work_rate_lbl)

                    meals_lbl = QLabel("2 meals")
                    meals_lbl.setStyleSheet(_INFO_LBL)
                    cell_vl.addWidget(meals_lbl)

                    _ow_mi = mi_ow
                    _ow_hr = t_hrs
                    _rt_mi = mi_ow * 2
                    _rt_hr = t_hrs * 2

                    # Non-hotel day: OW home if previous row had hotel, RT otherwise
                    # Look back past Weekend rows (they have no hotel checkbox)
                    _prev_cb = None
                    for _prev_r in range(r - 1, -1, -1):
                        if self._hotel_checks[t][_prev_r] is not None:
                            _prev_cb = self._hotel_checks[t][_prev_r]
                            break
                    _from_hotel = _prev_cb is not None and _prev_cb.isChecked()
                    if _from_hotel:
                        _drive_lbl = QLabel(f"Drive home: {_ow_mi:,.0f} mi OW")
                        _time_lbl  = QLabel(f"Travel home: {_ow_hr:.1f} hrs OW")
                    else:
                        _drive_lbl = QLabel(f"Drive: {_rt_mi:,.0f} mi RT")
                        _time_lbl  = QLabel(f"Travel: {_rt_hr:.1f} hrs RT")
                    _drive_lbl.setStyleSheet(_INFO_LBL)
                    _drive_lbl.setVisible(not _has_hotel and not flying and (_rt_mi > 0 or _ow_mi > 0))
                    cell_vl.addWidget(_drive_lbl)

                    _time_lbl.setStyleSheet(_INFO_LBL)
                    _time_lbl.setVisible(not _has_hotel and (_rt_hr > 0 or _ow_hr > 0))
                    cell_vl.addWidget(_time_lbl)

                    def _on_hotel_toggled(checked, _w=cell_w, _dl=_drive_lbl, _tl=_time_lbl,
                                          _fly=flying, _ow_m=_ow_mi, _rt_m=_rt_mi,
                                          _ow_h=_ow_hr, _rt_h=_rt_hr, _pc=_prev_cb):
                        _w.setStyleSheet("background:#fff0f3;" if checked else "background:#ffffff;")
                        _from_h = _pc is not None and _pc.isChecked()
                        if checked:
                            _dl.setVisible(False)
                            _tl.setVisible(False)
                        else:
                            if _from_h:
                                _dl.setText(f"Drive home: {_ow_m:,.0f} mi OW")
                                _tl.setText(f"Travel home: {_ow_h:.1f} hrs OW")
                            else:
                                _dl.setText(f"Drive: {_rt_m:,.0f} mi RT")
                                _tl.setText(f"Travel: {_rt_h:.1f} hrs RT")
                            _dl.setVisible(not _fly and (_rt_m > 0 or _ow_m > 0))
                            _tl.setVisible(_rt_h > 0 or _ow_h > 0)
                    cb.toggled.connect(_on_hotel_toggled)

                    # When the PREVIOUS row's hotel changes, update this row's labels
                    if _prev_cb is not None:
                        def _on_prev_toggled(prev_checked, _dl=_drive_lbl, _tl=_time_lbl,
                                             _fly=flying, _ow_m=_ow_mi, _rt_m=_rt_mi,
                                             _ow_h=_ow_hr, _rt_h=_rt_hr, _cur_cb=cb):
                            if not _cur_cb.isChecked():
                                if prev_checked:
                                    _dl.setText(f"Drive home: {_ow_m:,.0f} mi OW")
                                    _tl.setText(f"Travel home: {_ow_h:.1f} hrs OW")
                                    _dl.setVisible(not _fly and _ow_m > 0)
                                    _tl.setVisible(_ow_h > 0)
                                else:
                                    _dl.setText(f"Drive: {_rt_m:,.0f} mi RT")
                                    _tl.setText(f"Travel: {_rt_h:.1f} hrs RT")
                                    _dl.setVisible(not _fly and _rt_m > 0)
                                    _tl.setVisible(_rt_h > 0)
                        _prev_cb.toggled.connect(_on_prev_toggled)

                cell_vl.addStretch()
                self._outlook_table.setCellWidget(r, 2 + t, cell_w)

        ROW_H = 110
        for r in range(len(sched)):
            self._outlook_table.setRowHeight(r, ROW_H)
        self._outlook_table.setMinimumHeight(ROW_H * len(sched) + 50)
        self._outlook_table.setMaximumHeight(16_777_215)

    # ── Confirm schedule ──────────────────────────────────────────────────────

    def _confirm_schedule(self):
        if not self._schedule:
            return
        # Guard: tech count changed without regenerating — hotel_checks would be stale
        if len(self._hotel_checks) < self._num_techs:
            QMessageBox.warning(self, "Stale Schedule",
                "The technician count changed since the overview was generated.\n"
                "Please regenerate the Work Week Overview first.")
            return

        n          = self._num_techs
        work_days  = sum(1 for _, _, is_w, _ in self._schedule if is_w)
        hours_pd   = _n(self._hours_edit.text(), 8.0)
        weekend_rows = [r for r, (dl, _, _, _) in enumerate(self._schedule) if dl == "Weekend"]

        # Hotel
        total_hotel = 0
        for t in range(n):
            for cb in self._hotel_checks[t]:
                if cb is not None and cb.isChecked():
                    total_hotel += 1

        # Meals
        on_site_rows = sum(1 for dl, _, _, _ in self._schedule if dl != "Weekend")
        total_meals  = 0
        for mi in range(n):
            ti = self._travel_in_checks.get(mi)
            to = self._travel_out_checks.get(mi)
            uses_ti = ti is None or ti.isChecked()
            uses_to = to is None or to.isChecked()
            tech_on_site = on_site_rows - (0 if uses_ti else 1) - (0 if uses_to else 1)
            total_meals += tech_on_site * 2

        # Special row indices for hotel-based travel logic
        _ti_idx    = next((r for r, (_, act, _, _) in enumerate(self._schedule) if act == "Travel In"),  None)
        _insp_idxs = [r for r, (_, _, is_w, _) in enumerate(self._schedule) if is_w]
        _last_insp = _insp_idxs[-1] if _insp_idxs else None

        # Miles / flights — driven by hotel checkboxes
        flights = 0; total_drive_miles = 0.0
        for i in range(min(n, len(self._tech_rows))):
            row  = self._tech_rows[i]
            mode = row["mode"].currentText()
            hc   = self._hotel_checks[i] if i < len(self._hotel_checks) else []

            def _hv(r_idx, _hc=hc):
                cb = _hc[r_idx] if 0 <= r_idx < len(_hc) else None
                return cb is not None and cb.isChecked()

            ti_hotel   = _hv(_ti_idx)    if _ti_idx   is not None else False
            last_hotel = _hv(_last_insp) if _last_insp is not None else False

            if mode == "Flying":
                if ti_hotel:   flights += 1
                if last_hotel: flights += 1
            else:
                mi_ow = _n(row["mileage"].text())
                dm    = 0.0
                if ti_hotel:
                    dm += mi_ow
                for r_idx, (dl, _, is_w, _) in enumerate(self._schedule):
                    if is_w:
                        prev = _hv(r_idx - 1) if r_idx > 0 else False
                        here = _hv(r_idx)
                        if not prev and not here:  dm += mi_ow * 2
                        elif prev and not here:    dm += mi_ow
                        elif not prev and here:    dm += mi_ow
                    elif dl == "Weekend":
                        dm += mi_ow * 2
                if last_hotel:
                    dm += mi_ow
                total_drive_miles += dm

        # Hours
        onsite_hrs = work_days * hours_pd * n
        travel_hrs = 0.0
        for i in range(min(n, len(self._tech_rows))):
            t_hrs = _n(self._tech_rows[i]["travel"].text())
            hc    = self._hotel_checks[i] if i < len(self._hotel_checks) else []

            def _hv(r_idx, _hc=hc):
                cb = _hc[r_idx] if 0 <= r_idx < len(_hc) else None
                return cb is not None and cb.isChecked()

            ti_hotel   = _hv(_ti_idx)    if _ti_idx   is not None else False
            last_hotel = _hv(_last_insp) if _last_insp is not None else False

            th = 0.0
            if ti_hotel:
                th += t_hrs
            for r_idx, (dl, _, is_w, _) in enumerate(self._schedule):
                if is_w:
                    prev = _hv(r_idx - 1) if r_idx > 0 else False
                    here = _hv(r_idx)
                    if not prev and not here:  th += t_hrs * 2
                    elif prev and not here:    th += t_hrs
                    elif not prev and here:    th += t_hrs
                elif dl == "Weekend":
                    th += t_hrs * 2
            if last_hotel:
                th += t_hrs
            travel_hrs += th

        # Costs (before margin)
        labor_cost   = onsite_hrs        * _n(self._rate_labor.text(),  self.LABOR_RATE)
        travel_cost  = travel_hrs        * _n(self._rate_travel.text(), self.TRAVEL_RATE)
        hotel_cost   = total_hotel       * _n(self._rate_hotel.text(),  self.HOTEL_RATE)
        meal_cost    = total_meals       * _n(self._rate_meal.text(),   self.MEAL_RATE)
        mile_cost    = total_drive_miles * _n(self._rate_mile.text(),   self.MILE_RATE)
        flight_total = 0.0
        for i in range(min(n, len(self._tech_rows))):
            if self._tech_rows[i]["mode"].currentText() == "Flying":
                hc = self._hotel_checks[i] if i < len(self._hotel_checks) else []
                ti_hotel   = (hc[_ti_idx].isChecked()
                              if _ti_idx is not None and _ti_idx < len(hc) and hc[_ti_idx]
                              else False)
                last_hotel = (hc[_last_insp].isChecked()
                              if _last_insp is not None and _last_insp < len(hc) and hc[_last_insp]
                              else False)
                fc = _n(self._tech_rows[i]["flight_cost"].text())
                if ti_hotel:   flight_total += fc
                if last_hotel: flight_total += fc

        self._confirmed = {
            "work_days": work_days, "hotel_nights": total_hotel,
            "meals": total_meals,   "flights": flights,
            "drive_miles": total_drive_miles, "onsite_hrs": onsite_hrs,
            "travel_hrs": travel_hrs,
            "labor_cost": labor_cost, "travel_cost": travel_cost,
            "hotel_cost": hotel_cost, "meal_cost": meal_cost,
            "mile_cost": mile_cost,   "flight_total": flight_total,
        }

        self._sum_days.setText(str(work_days))
        self._sum_hotel.setText(str(total_hotel))
        self._sum_meals.setText(str(total_meals))
        self._sum_flights.setText(str(flights))
        self._sum_miles.setText(f"{total_drive_miles:,.0f}")

        self._cost_labor.setText(f"${labor_cost:,.2f}")
        self._cost_travel.setText(f"${travel_cost:,.2f}")
        self._cost_hotel.setText(f"${hotel_cost:,.2f}")
        self._cost_meals.setText(f"${meal_cost:,.2f}")
        self._cost_miles.setText(f"${mile_cost:,.2f}")
        self._cost_flights.setText(f"${flight_total:,.2f}")

        self._summary_card.setVisible(True)
        self._recalc_cost_totals()
        if hasattr(self, '_export_btn'):
            self._export_btn.setEnabled(True)
            self._export_btn.setToolTip("")

    # ── Cost totals ───────────────────────────────────────────────────────────

    def _recalc_cost_totals(self):
        if not self._confirmed:
            return
        c  = self._confirmed
        try:
            margin_pct = float(self._margin_lbl.text() or "0") / 100.0
        except:
            margin_pct = 0.0

        subtotal = (c["labor_cost"] + c["travel_cost"] + c["hotel_cost"] +
                    c["meal_cost"]  + c["mile_cost"]   + c["flight_total"])
        grand    = subtotal * (1.0 + margin_pct)

        self._cost_subtotal.setText(f"${subtotal:,.2f}")
        self._cost_grand.setText(f"${grand:,.2f}")
        self._recalc_ase_total()

    def _on_meals_edited(self):
        """Recalculate meal cost when the user manually edits the meals count."""
        if not self._confirmed:
            return
        try:
            new_meals = int(float(self._sum_meals.text().replace(",", "").strip() or "0"))
        except (ValueError, TypeError):
            return
        new_meals = max(0, new_meals)
        self._sum_meals.setText(str(new_meals))
        meal_rate = _n(self._rate_meal.text(), self.MEAL_RATE)
        new_meal_cost = new_meals * meal_rate
        self._confirmed["meals"] = new_meals
        self._confirmed["meal_cost"] = new_meal_cost
        self._cost_meals.setText(f"${new_meal_cost:,.2f}")
        self._recalc_cost_totals()

    def _recalc_ase_total(self):
        if not self._ase_enabled.isChecked():
            self._ase_total_edit.setText("—")
            return
        idx   = self._ase_tier_combo.currentIndex()
        label, price = _ASE_TIERS[idx]
        if price is None:
            try: price = float(self._ase_custom_edit.text() or "0")
            except: price = 0.0
        try:
            ase_margin = float(self._ase_margin_lbl.text() or "0") / 100.0
        except:
            ase_margin = 0.0
        ase_grand = price * (1.0 + ase_margin)
        self._ase_total_edit.setText(f"${ase_grand:,.2f}")

    # ── ASE helpers ───────────────────────────────────────────────────────────

    def _on_ase_toggled(self, checked: bool):
        self._ase_body.setVisible(checked)
        self._recalc_ase_total()

    def _on_ase_tier_changed(self, idx: int):
        _, price = _ASE_TIERS[idx]
        is_custom = price is None
        self._ase_custom_lbl.setVisible(is_custom)
        self._ase_custom_edit.setVisible(is_custom)
        self._recalc_ase_total()

    def _adj_margin(self, delta: int):
        try: v = float(self._margin_lbl.text() or "0")
        except: v = 0.0
        v = max(0.0, min(99.9, v + delta))
        self._margin_lbl.setText(f"{v:.1f}")

    def _adj_ase_margin(self, delta: int):
        try: v = float(self._ase_margin_lbl.text() or "0")
        except: v = 0.0
        v = max(0.0, min(99.9, v + delta))
        self._ase_margin_lbl.setText(f"{v:.1f}")

    # ── Serialisation ─────────────────────────────────────────────────────────

    def get_data(self) -> dict:
        hotel_states = []
        if self._schedule and self._hotel_checks:
            for t in range(self._num_techs):
                ts = []
                for r in range(len(self._schedule)):
                    cb = (self._hotel_checks[t][r]
                          if t < len(self._hotel_checks) and r < len(self._hotel_checks[t])
                          else None)
                    ts.append(bool(cb.isChecked()) if cb else False)
                hotel_states.append(ts)
        return {
            "num_techs":         self._num_techs,
            "hours_per_day":     self._hours_edit.text(),
            "working_days":      self._days_edit.text(),
            "tech_names":        [r["name"].text()        for r in self._tech_rows],
            "tech_travel":       [r["travel"].text()      for r in self._tech_rows],
            "tech_modes":        [r["mode"].currentText() for r in self._tech_rows],
            "tech_flight_costs": [r["flight_cost"].text() for r in self._tech_rows],
            "tech_mileages":     [r["mileage"].text()     for r in self._tech_rows],
            "work_days":         [i for i, cb in enumerate(self._day_checks) if cb.isChecked()],
            "schedule":          [(dl, act, iw, dh) for dl, act, iw, dh in self._schedule],
            "hotel_states":      hotel_states,
            "confirmed":         self._confirmed,
            "margin":            self._margin_lbl.text(),
            "ase_enabled":       self._ase_enabled.isChecked(),
            "ase_tier":          self._ase_tier_combo.currentIndex(),
            "ase_custom":        self._ase_custom_edit.text(),
            "ase_margin":        self._ase_margin_lbl.text(),
            # Rate overrides
            "rate_labor":   self._rate_labor.text(),
            "rate_travel":  self._rate_travel.text(),
            "rate_hotel":   self._rate_hotel.text(),
            "rate_meal":    self._rate_meal.text(),
            "rate_mile":    self._rate_mile.text(),
            # Project info
            "ir_project":  self._ir_project.text()  if self._ir_project  else "",
            "ir_proposal": self._ir_proposal.text() if self._ir_proposal else "",
            "ir_customer": self._ir_customer.text() if self._ir_customer else "",
            "ir_location": self._ir_location.text() if self._ir_location else "",
            "ir_scope":    self._ir_scope.text()    if self._ir_scope    else "",
            "ir_contact":  self._ir_contact.toPlainText() if self._ir_contact else "",
            "ir_picture":  (self._ir_picture.image_path or "") if self._ir_picture else "",
            # Version control
            "version_major":   self._version_major_spin.value() if self._version_major_spin else 1,
            "version_minor":   self._version_minor_spin.value() if self._version_minor_spin else 0,
            "version_history": ([self._version_history.item(i).text()
                                 for i in range(self._version_history.count())]
                                if self._version_history else []),
        }

    def restore_data(self, d: dict):
        if not d: return
        self._days_edit.setText(str(d.get("working_days", "")))
        self._hours_edit.setText(str(d.get("hours_per_day", "8")))
        self._num_techs = d.get("num_techs", 1)
        self._tech_count_lbl.setText(str(self._num_techs))
        self._rebuild_tech_rows()
        for i, nm in enumerate(d.get("tech_names", [])):
            if i < len(self._tech_rows): self._tech_rows[i]["name"].setText(nm)
        for i, tv in enumerate(d.get("tech_travel", [])):
            if i < len(self._tech_rows): self._tech_rows[i]["travel"].setText(tv)
        for i, mode in enumerate(d.get("tech_modes", [])):
            if i < len(self._tech_rows):
                combo = self._tech_rows[i]["mode"]
                idx   = combo.findText(mode)
                if idx >= 0:
                    combo.blockSignals(True)
                    combo.setCurrentIndex(idx)
                    combo.blockSignals(False)
                    if "_toggle" in self._tech_rows[i]:
                        self._tech_rows[i]["_toggle"](mode)
        for i, fc in enumerate(d.get("tech_flight_costs", [])):
            if i < len(self._tech_rows): self._tech_rows[i]["flight_cost"].setText(fc)
        for i, mi in enumerate(d.get("tech_mileages", [])):
            if i < len(self._tech_rows): self._tech_rows[i]["mileage"].setText(mi)
        wd = set(d.get("work_days", [0,1,2,3,4]))
        for i, cb in enumerate(self._day_checks):
            cb.setChecked(i in wd)
        self._margin_lbl.setText(d.get("margin", "0.0"))
        self._ase_enabled.setChecked(d.get("ase_enabled", False))
        idx = d.get("ase_tier", 0)
        if 0 <= idx < self._ase_tier_combo.count():
            self._ase_tier_combo.setCurrentIndex(idx)
        self._ase_custom_edit.setText(d.get("ase_custom", ""))
        self._ase_margin_lbl.setText(d.get("ase_margin", "0.0"))
        self._rate_labor.setText(d.get("rate_labor",  str(self.LABOR_RATE)))
        self._rate_travel.setText(d.get("rate_travel", str(self.TRAVEL_RATE)))
        self._rate_hotel.setText(d.get("rate_hotel",  str(self.HOTEL_RATE)))
        self._rate_meal.setText(d.get("rate_meal",    str(self.MEAL_RATE)))
        self._rate_mile.setText(d.get("rate_mile",    str(self.MILE_RATE)))
        self._confirmed = d.get("confirmed", {})
        if self._confirmed:
            c = self._confirmed
            self._sum_days.setText(str(c.get("work_days", "—")))
            self._sum_hotel.setText(str(c.get("hotel_nights", "—")))
            self._sum_meals.setText(str(c.get("meals", "—")))
            self._sum_flights.setText(str(c.get("flights", "—")))
            self._sum_miles.setText(f"{c.get('drive_miles', 0):,.0f}")
            self._cost_labor.setText(f"${c.get('labor_cost', 0):,.2f}")
            self._cost_travel.setText(f"${c.get('travel_cost', 0):,.2f}")
            self._cost_hotel.setText(f"${c.get('hotel_cost', 0):,.2f}")
            self._cost_meals.setText(f"${c.get('meal_cost', 0):,.2f}")
            self._cost_miles.setText(f"${c.get('mile_cost', 0):,.2f}")
            self._cost_flights.setText(f"${c.get('flight_total', 0):,.2f}")
            self._summary_card.setVisible(True)
            self._recalc_cost_totals()
        # Project info
        if self._ir_project:  self._ir_project.setText(d.get("ir_project", ""))
        if self._ir_proposal: self._ir_proposal.setText(d.get("ir_proposal", ""))
        if self._ir_customer: self._ir_customer.setText(d.get("ir_customer", ""))
        if self._ir_location: self._ir_location.setText(d.get("ir_location", ""))
        if self._ir_scope:    self._ir_scope.setText(d.get("ir_scope", ""))
        if self._ir_contact:  self._ir_contact.setPlainText(d.get("ir_contact", ""))
        if self._ir_picture and d.get("ir_picture"): self._ir_picture.set_image(d["ir_picture"])
        # Version control
        if self._version_major_spin is not None:
            self._version_major_spin.setValue(int(d.get("version_major", 1)))
            self._version_minor_spin.setValue(int(d.get("version_minor", 0)))
            self._update_version_label()
            self._version_history.clear()
            for _entry in d.get("version_history", []):
                self._version_history.addItem(_entry)

    # ── Version control ───────────────────────────────────────────────────────

    def _version_str(self) -> str:
        if self._version_major_spin is None:
            return "V1.0"
        return f"V{self._version_major_spin.value()}.{self._version_minor_spin.value()}"

    def _update_version_label(self):
        if self._version_badge is not None:
            self._version_badge.setText(self._version_str())

    def _toggle_history_panel(self):
        self._history_collapsed = not self._history_collapsed
        self._history_panel.setVisible(not self._history_collapsed)
        self._btn_collapse_hist.setText(
            "▶  Change History" if self._history_collapsed
            else "▼  Change History")

    def _add_version_entry(self):
        from PyQt6.QtWidgets import (QDialog as _QDlg, QVBoxLayout as _QVL2,
                                     QDialogButtonBox as _QBB)
        v = self._version_str()
        dlg = _QDlg(self)
        dlg.setWindowTitle(f"Add Change Entry — {v}")
        dlg.setFixedWidth(420)
        dlg.setStyleSheet(
            "QDialog { background:#ffffff; }"
            "QLabel  { color:#1a0509; font-size:12px; }"
            "QLineEdit { border:1px solid #d6c0c5; border-radius:4px;"
            "  padding:6px 10px; font-size:12px; color:#1a0509; background:#fafbfd; }"
            "QLineEdit:focus { border-color:#920d2e; }"
            "QPushButton { background:#f4f6fa; color:#1a0509;"
            "  border:1px solid #d6c0c5; border-radius:4px;"
            "  padding:6px 18px; min-width:72px; font-size:12px; }"
            "QPushButton:hover { background:#e0e6f0; }")
        vbox = _QVL2(dlg)
        vbox.setSpacing(10); vbox.setContentsMargins(16, 16, 16, 16)
        from PyQt6.QtWidgets import QLabel as _QL2, QLineEdit as _QLE2
        vbox.addWidget(_QL2(f"<b>Version:</b>  {v}"))
        desc = _QLE2(); desc.setPlaceholderText("Describe this change…")
        vbox.addWidget(desc)
        btns = _QBB(_QBB.StandardButton.Ok | _QBB.StandardButton.Cancel)
        btns.accepted.connect(dlg.accept); btns.rejected.connect(dlg.reject)
        vbox.addWidget(btns)
        if dlg.exec() == _QDlg.DialogCode.Accepted and desc.text().strip():
            today = _date.today().strftime("%m/%d/%Y")
            self._version_history.addItem(
                f"{v}  —  {desc.text().strip()}  ({today})")

    def _remove_version_entry(self):
        if self._version_history is None:
            return
        row = self._version_history.currentRow()
        if row >= 0:
            self._version_history.takeItem(row)

    # ── File I/O ──────────────────────────────────────────────────────────────

    def save_data(self):
        import re as _re
        desktop  = os.path.join(os.path.expanduser("~"), "Desktop")
        proposal = (self._ir_proposal.text().strip() if self._ir_proposal else "") or "thermal"
        safe     = _re.sub(r'[\\/:*?"<>|]', "-", proposal)
        _ver     = f" {self._version_str()}"
        path, _ = QFileDialog.getSaveFileName(
            self, "Save Thermal Imaging Project",
            os.path.join(desktop, f"{safe}{_ver}.mcmxt"),
            "Thermal Project (*.mcmxt)")
        if not path: return
        try:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(self.get_data(), f, indent=2, ensure_ascii=False)
            QMessageBox.information(self, "Saved", f"\u2714  Saved:\n{path}")
        except Exception as e:
            QMessageBox.critical(self, "Save Failed", str(e))

    def load_data(self):
        desktop = os.path.join(os.path.expanduser("~"), "Desktop")
        path, _ = QFileDialog.getOpenFileName(
            self, "Load Thermal Imaging Project", desktop,
            "Thermal Project (*.mcmxt)")
        if not path: return
        try:
            with open(path, "r", encoding="utf-8") as f:
                self.restore_data(json.load(f))
        except Exception as e:
            QMessageBox.critical(self, "Load Failed", str(e))


# ── ThermalCard — embeds ThermalImagingWidget inside a collapsible cost-estimator card ──

class ThermalCard(_Card):
    """
    Collapsible card version of the Thermal Imaging tool for embedding in the
    Cost Estimator widget.  An 'Active' checkbox in the header enables/disables
    the section; when inactive the body stays hidden.
    """

    def __init__(self, parent=None):
        super().__init__("Thermal Imaging", parent)  # sets self._hdr_row via _Card.__init__

        # Active checkbox in card header
        self._enable_cb = QCheckBox("Active")
        self._enable_cb.setChecked(False)
        self._enable_cb.setStyleSheet(
            "QCheckBox { font-size:12px; color:#920d2e; font-weight:700;"
            "  border:none; padding:0; }"
            "QCheckBox::indicator { width:14px; height:14px; }"
            "QCheckBox::indicator:unchecked {"
            "  border:2px solid #d6c0c5; background:#ffffff; border-radius:3px; }"
            "QCheckBox::indicator:checked {"
            "  border:2px solid #920d2e; background:#920d2e; border-radius:3px; }"
        )
        self._enable_cb.toggled.connect(self._on_enable_changed)
        self._hdr_row.addWidget(self._enable_cb)

        # Compact ThermalImagingWidget (no scroll wrapper, no bottom bar)
        self._thermal = ThermalImagingWidget(compact=True)
        self._body_layout.addWidget(self._thermal)

        # Start inactive / collapsed
        self._body.setVisible(False)
        self._collapsed = True
        self._toggle_btn.setText("►")

    def _on_enable_changed(self, checked: bool):
        self._body.setVisible(checked)
        self._collapsed = not checked
        self._toggle_btn.setText("▼" if checked else "►")

    def export_report(self):
        data = self.get_data()
        desktop = os.path.join(os.path.expanduser("~"), "Desktop")
        proposal = data.get("ir_proposal", "THERMAL-REPORT").strip() or "THERMAL-REPORT"
        path, _ = QFileDialog.getSaveFileName(
            self, "Export Thermal Imaging Report",
            os.path.join(desktop, f"{proposal}.docx"),
            "Word Document (*.docx)")
        if not path: return
        try:
            generate_thermal_doc(data, path)
            QMessageBox.information(self, "Exported", f"Report saved to:\n{path}")
        except Exception as e:
            import traceback
            QMessageBox.critical(self, "Export Failed", f"{e}\n\n{traceback.format_exc()}")

    def get_data(self) -> dict:
        return {"active": self._enable_cb.isChecked(), **self._thermal.get_data()}

    def restore_data(self, d: dict):
        self._thermal.restore_data(d)
        active = d.get("active", False)
        self._enable_cb.blockSignals(True)
        self._enable_cb.setChecked(active)
        self._enable_cb.blockSignals(False)
        self._on_enable_changed(active)


# ── generate_thermal_doc ──────────────────────────────────────────────────────

def generate_thermal_doc(data: dict, output_path: str):
    """Generate thermal imaging report via docxtpl then python-docx pricing table."""
    import os as _os, sys as _sys
    from docxtpl import DocxTemplate, InlineImage
    from docx.shared import Mm
    from datetime import date as _today_dt

    # Use sys._MEIPASS when running from a PyInstaller bundle so the
    # template is found inside the extracted _internal folder.
    _here = getattr(_sys, '_MEIPASS', _os.path.dirname(_os.path.abspath(__file__)))
    candidates = [
        _os.path.join(_here, "thermal_template.docx"),
    ]
    template = next((c for c in candidates if _os.path.exists(c)), None)
    if not template:
        raise FileNotFoundError(
            "Thermal template not found. Expected 'thermal_template.docx' in the scripts folder.")

    today_str  = _ordinal_date(_today_dt.today())
    proposal   = data.get("ir_proposal", "")
    customer   = data.get("ir_customer", "")
    location   = data.get("ir_location", "")
    project    = data.get("ir_project",  "Thermal Imaging Study and Report")
    scope_desc = data.get("ir_scope",    "")
    pic_path   = data.get("ir_picture",  "")

    contact_lines = (data.get("ir_contact", "") or "").splitlines()
    def _cl(i): return contact_lines[i].strip() if i < len(contact_lines) else ""
    pres_name  = _cl(0); pres_email = _cl(2)
    acct_name  = _cl(4); acct_email = _cl(6)

    # Compute pricing totals
    confirmed = data.get("confirmed", {})
    try:
        margin_pct = float(data.get("margin", "0") or "0") / 100.0
    except Exception:
        margin_pct = 0.0

    subtotal    = (confirmed.get("labor_cost", 0) + confirmed.get("travel_cost", 0) +
                   confirmed.get("hotel_cost", 0) + confirmed.get("meal_cost", 0) +
                   confirmed.get("mile_cost", 0)  + confirmed.get("flight_total", 0))
    labor_grand = subtotal * (1.0 + margin_pct)

    ase_grand = 0.0
    if data.get("ase_enabled"):
        tier_idx = data.get("ase_tier", 0)
        if 0 <= tier_idx < len(_ASE_TIERS):
            _, tier_price = _ASE_TIERS[tier_idx]
            if tier_price is None:
                try: tier_price = float(data.get("ase_custom", "0") or "0")
                except: tier_price = 0.0
            try:
                ase_margin = float(data.get("ase_margin", "0") or "0") / 100.0
            except:
                ase_margin = 0.0
            ase_grand = tier_price * (1.0 + ase_margin)

    grand_total = labor_grand + ase_grand

    # Build scanning description lines
    num_techs = data.get("num_techs", 1)
    work_days = confirmed.get("work_days", 0)
    hours_pd  = float(data.get("hours_per_day", "8") or "8")

    days_desc  = (f"Up to {work_days} day{'s' if work_days != 1 else ''} of on-site support "
                  f"({hours_pd:.0f} hours per day)") if work_days > 0 else ""
    techs_desc = (f"for {num_techs} technician{'s' if num_techs != 1 else ''}"
                  "\nWork to occur during standard business hours (M-F, 8-5)")

    extras = []
    if confirmed.get("hotel_nights", 0) > 0:
        extras.append(f"Hotel ({confirmed['hotel_nights']} nights)")
    if confirmed.get("meals", 0) > 0:
        extras.append(f"Meals ({confirmed['meals']})")
    if confirmed.get("drive_miles", 0) > 0:
        extras.append(f"Mileage ({confirmed['drive_miles']:,.0f} mi)")
    if confirmed.get("flights", 0) > 0:
        extras.append(f"Flights ({confirmed['flights']})")
    expenses_desc = ("Includes: " + ", ".join(extras)) if extras else ""

    # Phase 1 — render Jinja2 placeholders (including customer picture)
    tpl = DocxTemplate(template)
    customer_image = ""
    if pic_path and _os.path.exists(pic_path):
        try:
            customer_image = InlineImage(tpl, pic_path, width=Mm(50))
        except Exception:
            customer_image = ""

    tpl.render({
        "customer_name":        customer,
        "customer_location":    location,
        "proposal_number":      proposal,
        "today_date":           today_str,
        "pres_name":            pres_name,
        "pres_email":           pres_email,
        "acct_name":            acct_name,
        "acct_email":           acct_email,
        "project_name":         project,
        "scope_description":    scope_desc,
        "scanning_price":       f"{labor_grand:,.0f}",
        "ase_price":            f"${ase_grand:,.0f}" if ase_grand > 0 else "—",
        "total_price":          f"{grand_total:,.0f}",
        "scanning_days_desc":   days_desc,
        "scanning_techs_desc":  techs_desc,
        "scanning_expenses_desc": expenses_desc,
        "customer_picture":     customer_image,
    })
    tpl.save(output_path)

    # ── Phase 2: post-process for any placeholders docxtpl couldn't reach ───────
    # docxtpl renders the body but may miss text inside text boxes on the first
    # page, VML shapes, and — most critically — Word headers/footers which live
    # in separate XML parts.  Open the saved file with python-docx and replace
    # any remaining {{ ... }} markers in ALL parts of the document.
    try:
        from docx import Document as _Doc
        from docx.oxml.ns import qn as _qn

        doc2 = _Doc(output_path)
        _changed = [False]

        _repl = {
            "customer_name":          customer,
            "customer_location":      location,
            "proposal_number":        proposal,
            "today_date":             today_str,
            "pres_name":              pres_name,
            "pres_email":             pres_email,
            "acct_name":              acct_name,
            "acct_email":             acct_email,
            "project_name":           project,
            "scope_description":      scope_desc,
            "scanning_price":         f"{labor_grand:,.0f}",
            "ase_price":              f"${ase_grand:,.0f}" if ase_grand > 0 else "—",
            "total_price":            f"{grand_total:,.0f}",
            "scanning_days_desc":     days_desc,
            "scanning_techs_desc":    techs_desc,
            "scanning_expenses_desc": expenses_desc,
        }

        def _fix_t(t_elem):
            txt = t_elem.text or ""
            if "{{" not in txt:
                return
            for key, val in _repl.items():
                for marker in (f"{{{{{key}}}}}", f"{{{{ {key} }}}}"):
                    if marker in txt:
                        txt = txt.replace(marker, str(val))
                        _changed[0] = True
            t_elem.text = txt

        # Collect XML roots to scan: body + all headers and footers
        xml_roots = [doc2.element]
        for sec in doc2.sections:
            for part in (sec.header, sec.first_page_header, sec.even_page_header,
                         sec.footer,  sec.first_page_footer,  sec.even_page_footer):
                try:
                    if part is not None and part._element is not None:
                        xml_roots.append(part._element)
                except Exception:
                    pass

        for root in xml_roots:
            for t_elem in root.iter(_qn("w:t")):
                _fix_t(t_elem)

        # Picture: find any remaining {{ customer_picture }} marker anywhere
        if pic_path and _os.path.exists(pic_path):
            from docx.shared import Inches
            for root in xml_roots:
                for t_elem in root.iter(_qn("w:t")):
                    txt = t_elem.text or ""
                    if "customer_picture" not in txt:
                        continue
                    t_elem.text = ""
                    r_elem = t_elem.getparent()
                    tmp_doc = _Doc()
                    run = tmp_doc.paragraphs[0].add_run()
                    run.add_picture(pic_path, width=Inches(1.75))
                    r_elem.addprevious(run._r)
                    _changed[0] = True
                    break
                else:
                    continue
                break

        # Layout fix: the first paragraph in the template contains the customer
        # picture placeholder plus two <w:br/> line-break runs.  When no picture
        # is provided those line breaks add dead vertical space that shifts the
        # cover-page content downward.  Remove them so the paragraph is invisible.
        if not (pic_path and _os.path.exists(pic_path)):
            try:
                body = doc2.element.body
                first_p = next(c for c in body if c.tag == _qn("w:p"))
                # Remove all <w:r> children whose only content is <w:br/>
                for r_elem in list(first_p):
                    if r_elem.tag != _qn("w:r"):
                        continue
                    children = [c for c in r_elem if not c.tag.endswith("}rPr")]
                    if len(children) == 1 and children[0].tag == _qn("w:br"):
                        first_p.remove(r_elem)
                        _changed[0] = True
            except Exception:
                pass

        if _changed[0]:
            doc2.save(output_path)
    except Exception:
        pass  # post-processing is best-effort; the base export still succeeds
